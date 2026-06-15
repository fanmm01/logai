# /// script
# dependencies = [
#     "flask",
#     "requests",
#     "pillow",
#     "openai",
#     "python-docx",
#     "PyPDF2",
#     "pymupdf",
#     "websockets",
# ]
# ///

# LogAI 4.4.0 - TRPG Log Analysis and Illustration Server
# 原作者：Air, Gemini
# 改编：fanmm @fanmm01, github copilot
# logutil段大量参考与摘抄了 @chaye2333的fwlog项目的设计和实现，感谢其开源贡献！

import os
import sys
import json
import time
import base64
import zlib
import re
import uuid
import socket
import platform
import ctypes
import threading
import subprocess
import tempfile
import sqlite3
import asyncio
import argparse
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from queue import Full, Queue
from typing import Any, Dict, Optional
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from flask import Flask, request, send_file, jsonify, abort
from PIL import Image, ImageDraw, ImageFont
from openai import OpenAI
from xml.etree import ElementTree as ET
import PyPDF2
import urllib.parse
import zipfile
import shutil
import hashlib
import datetime
import atexit
from docx import Document
from websockets.legacy.client import connect as ws_connect

# ================= 配置区域 =================
AI_API_KEY = "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxx" # 请填入你的 OpenAI API Key 或 DeepSeek API Key
AI_BASE_URL = "https://api.deepseek.com" # 或 https://api.openai.com/v1
AI_MODEL = "deepseek-v4-flash"
AI_MODEL_PRO = "deepseek-v4-pro"

# --- 绘图专用配置 (NovelAI) ---
# 请填入 NovelAI 提供的 API Key (通常是以 pst- 开头的一长串字符)
IMAGE_API_KEY = "pst-zzzzzzzzzzzzzzzzzzzz"

# 推荐使用最新的 V3 模型，这是目前 NovelAI 画二次元最好的模型
IMAGE_MODEL = "nai-diffusion-4-5-full" 
# ===========================================

PRO_SYSTEM_PROMPT = """
            你是一位毒舌但极其专业的 TRPG 跑团鉴赏家（KP/DM）。请阅读以下跑团 Log，生成一份简报。
            请严格按照以下格式输出（不要用 Markdown，不要加粗，直接分行）：

           【总体评分】：(0-100分)
           (请给出理由)
           【剧情概要】：
           (简述发生了什么，500字内)
           【高光时刻】：
           (找出1-3个最精彩或最搞的一幕)
           【主要槽点】：
           (吐槽逻辑漏洞、糟糕的RP或离谱的操作)
           【KP寄语】：
           (一句话总结)

           风格要求：幽默、犀利、像老练的调查员在写结案报告。当日志内容是DND时，将KP寄语替换成DM寄语。
    """

KIND_SYSTEM_PROMPT = """
            你是一位温柔但极其专业的 TRPG 跑团鉴赏家（KP/DM）。请阅读以下跑团 Log，生成一份简报。
            请严格按照以下格式输出（不要用 Markdown，不要加粗，直接分行）：

           【总体评分】：(0-100分)
           (请给出理由)
           【剧情概要】：
           (简述发生了什么，500字内)
           【高光时刻】：
           (找出1-3个最精彩或最搞的一幕)
           【主要槽点】：
           (吐槽逻辑漏洞、糟糕的RP或离谱的操作)
           【KP寄语】：
           (一句话总结)

           风格要求：给予适当的鼓励以及表演，表现的更加体贴人。当日志内容是DND时，将KP寄语替换成DM寄语。
    """

DEFAULT_SYSTEM_PROMPT = """
            你是一位毒舌但极其专业的 TRPG 跑团鉴赏家（KP/DM）。请阅读以下跑团 Log，生成一份简报。
            请严格按照以下格式输出（不要用 Markdown，不要加粗，直接分行）：

           【总体评分】：(0-100分)
           (请给出理由)
           【剧情概要】：
           (简述发生了什么，500字内)
           【高光时刻】：
           (找出1-3个最精彩或最搞的一幕)
           【主要槽点】：
           (吐槽逻辑漏洞、糟糕的RP或离谱的操作)
           【KP寄语】：
           (一句话总结)

           风格要求：幽默、犀利、像老练的调查员在写结案报告。当日志内容是DND时，将KP寄语替换成DM寄语。
    """

# 字体路径
FONT_PATH = "./fonts/GB2312.ttf" 

# 最大处理条目数
MAX_LOG_ENTRIES = 20000
# 发送给 AI 的最大字符数
MAX_AI_CHARS = 3000000

URL_RE = re.compile(r"https?://[^\s\]\"']+")

# ====== fwlog 移植：发言识别正则 ======
ANGLE_SPEAKER_RE = re.compile(
    r"^\s*<(?P<name>[^>\n]+)>\s*:\s*(?P<content>.*?)\s*$"
)
PLAIN_SPEAKER_RE = re.compile(
    r"^\s*(?P<name>[^:<>\[\]【】\n][^:<>\[\]【】\n]{0,79}?)\s*:\s*(?P<content>.*?)\s*$"
)
MULTILINE_ANGLE_OPEN_RE = re.compile(r"^\s*<\s*$")
MULTILINE_ANGLE_CLOSE_RE = re.compile(r"^\s*>\s*:\s*(?P<content>.*?)\s*$")
LOG_METADATA_LINE_RE = re.compile(
    r"^\s*(?:(?P<date>\d{4}/\d{1,2}/\d{1,2}))?(?:\s*(?P<clock>\d{1,2}:\d{2}(?::\d{2})?))?\s*$"
)
TIMESTAMPED_ANGLE_SPEAKER_RE = re.compile(
    r"^\s*(?:(?P<date>\d{4}/\d{1,2}/\d{1,2})\s+)?(?:(?P<clock>\d{1,2}:\d{2}(?::\d{2})?)\s+)?<(?P<name>[^>\n]+)>\s*:\s*(?P<content>.*?)\s*$"
)
TIMESTAMPED_PLAIN_SPEAKER_RE = re.compile(
    r"^\s*(?:(?P<date>\d{4}/\d{1,2}/\d{1,2})\s+)?(?:(?P<clock>\d{1,2}:\d{2}(?::\d{2})?)\s+)?(?P<name>[^:<>\[\]【】][^:<>\[\]【】]{0,79}?)\s*:\s*(?P<content>.*?)\s*$"
)
SELF_LOG_IGNORE_RE = re.compile(r"已(?:经)?记录\s*\d+\s*条消息")

# === 新增：方括号时间戳 + 管道分隔符格式 ===
# 格式 a/b: [YYYY-MM-DD HH:MM:SS] <角色名|玩家昵称> text
BRACKET_PIPE_SPEAKER_RE = re.compile(
    r"^\s*\[(?P<date>\d{4}-\d{2}-\d{2})\s+(?P<clock>\d{2}:\d{2}:\d{2})\]\s*<(?P<name>[^|>]+)\|(?P<player>[^>]+)>\s*(?P<content>.*?)\s*$"
)

# 格式 c: [YYYY-MM-DD HH:MM:SS] * 角色名|玩家昵称 text
ASTERISK_PIPE_SPEAKER_RE = re.compile(
    r"^\s*\[(?P<date>\d{4}-\d{2}-\d{2})\s+(?P<clock>\d{2}:\d{2}:\d{2})\]\s*\*\s*(?P<name>[^|]+)\|(?P<player>[^\s]+)\s*(?P<content>.*?)\s*$"
)

# --- 百度网盘 OpenAPI 配置 ---
BAIDU_APP_KEY = "kkk"
BAIDU_SECRET_KEY = "sss"
# 首次使用前，请在浏览器访问以下链接（将其中的【你的AppKey】替换成实际的AppKey）：
# http://openapi.baidu.com/oauth/2.0/authorize?response_type=code&client_id=【你的AppKey】&redirect_uri=oob&scope=basic,netdisk
# 同意授权后，网页会显示一段 Authorization Code，将其复制到下方：
BAIDU_AUTH_CODE = "aaa"


BAIDU_TARGET_DIR = "/coc_20260220_041522" # 指定的搜索目录
BAIDU_TOKEN_FILE = "baidu_token.json"    # 用于持久化保存token的文件

# --- 服务与 NapCat 桥接配置 ---
LOGAI_HOST = os.getenv("LOGAI_HOST", "0.0.0.0")
LOGAI_PORT = int(os.getenv("LOGAI_PORT", "8000"))


def get_lan_ip():
    """Detect the server's LAN IP address. Falls back to 127.0.0.1."""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.settimeout(0.1)
        s.connect(('8.8.8.8', 53))
        ip = s.getsockname()[0]
        s.close()
        if ip and not ip.startswith('127.'):
            return ip
    except Exception:
        pass
    # Fallback: iterate interfaces
    try:
        hostname = socket.gethostname()
        for info in socket.getaddrinfo(hostname, None, socket.AF_INET):
            ip = info[4][0]
            ip_str = str(ip)                # 强制转 str，消除类型歧义
            if ip_str and not ip_str.startswith('127.'):
                return ip_str
    except Exception:
        pass
    return '127.0.0.1'


def sanitize_body_for_log(body):
    """Replace base64 blobs with truncated markers for safe logging."""
    if not isinstance(body, dict):
        return str(body)
    safe = {}
    for k, v in body.items():
        s = str(v)
        if s.startswith('base64://') and len(s) > 200:
            safe[k] = f'base64://<...{len(s)-9} chars truncated...>'
        else:
            safe[k] = v
    return safe


BRIDGE_TOKEN = os.getenv("BRIDGE_TOKEN", "")
BRIDGE_TTL_SEC = int(os.getenv("BRIDGE_TTL_SEC", "86400"))
MAX_BRIDGE_FILES_PER_GROUP = int(os.getenv("MAX_BRIDGE_FILES_PER_GROUP", "10"))
BRIDGE_CACHE_DIR = os.getenv("BRIDGE_CACHE_DIR", os.path.join(os.path.dirname(os.path.abspath(__file__)), "instance", "napcat_file_bridge"))
BRIDGE_PUBLIC_BASE = os.getenv("BRIDGE_PUBLIC_BASE", f"http://{get_lan_ip()}:{LOGAI_PORT}").rstrip("/")
NAPCAT_API_BASE = os.getenv("NAPCAT_API_BASE", "http://127.0.0.1:8084").rstrip("/")
NAPCAT_API_BASES_RAW = os.getenv("NAPCAT_API_BASES", "").strip()
NAPCAT_TOKEN = os.getenv("NAPCAT_TOKEN", "1")
DOWNLOAD_TIMEOUT_SEC = int(os.getenv("DOWNLOAD_TIMEOUT_SEC", "180"))
MAX_FILE_MB = int(os.getenv("MAX_FILE_MB", "150"))
MAX_FILE_BYTES = MAX_FILE_MB * 1024 * 1024
PULL_LATEST_ON_EMPTY = os.getenv("PULL_LATEST_ON_EMPTY", "1") == "1"
BRIDGE_QUEUE_SIZE = max(1, int(os.getenv("BRIDGE_QUEUE_SIZE", "128")))
REFRESH_LATEST_ON_READ = os.getenv("REFRESH_LATEST_ON_READ", "1") == "1"
BRIDGE_POLL_INTERVAL_SEC = int(os.getenv("BRIDGE_POLL_INTERVAL_SEC", "5"))
LOGUTIL_POLL_INTERVAL_SEC = 1  # 1s during logutil recording

# --- NapCat 文件桥接模式 ---
# 0 = WS 实时推送模式: 依赖 WebSocket 监听 group_upload 事件，收到即处理
# 1 = 轮询模式 (备用): 每 BRIDGE_POLL_INTERVAL_SEC 秒主动拉取最新文件
NC_FILE_BRIDGE_MODE = int(os.getenv("NC_FILE_BRIDGE_MODE", "0"))
MAX_BRIDGE_FILES_PER_GROUP = max(1, int(os.getenv("MAX_BRIDGE_FILES_PER_GROUP", "20")))

# --- NapCat WebSocket 实时连接配置 (logutil recording) ---
NAPCAT_WS_URL = os.getenv("NAPCAT_WS_URL", "ws://127.0.0.1:3001")
NAPCAT_WS_TOKEN = os.getenv("NAPCAT_WS_TOKEN") or NAPCAT_TOKEN or ""
LOGUTIL_WS_ENABLED = os.getenv("LOGUTIL_WS_ENABLED", "1") == "1"
LOGUTIL_MILESTONE_INTERVAL = 1000  # Notify every 1000 items

# ===========================================

# --- 每日全局省流缓存池 ---
DAILY_CACHE = {}

def get_daily_cache(hash_key):
    """获取今日的缓存图片，跨天自动清空内存"""
    today = datetime.date.today().isoformat()
    if today not in DAILY_CACHE:
        DAILY_CACHE.clear()
        DAILY_CACHE[today] = {}
    return DAILY_CACHE[today].get(hash_key)

def set_daily_cache(hash_key, images_list):
    """将生成的图片字节流写入今日缓存池"""
    today = datetime.date.today().isoformat()
    if today not in DAILY_CACHE:
        DAILY_CACHE.clear()
        DAILY_CACHE[today] = {}
    DAILY_CACHE[today][hash_key] = images_list

app = Flask(__name__)
SERVICE_VERSION = "4.5.3"
_openai_client = None

def get_openai_client():
    """Lazy-init OpenAI client so CLI/env overrides to API_KEY/BASE_URL take effect."""
    global _openai_client
    if _openai_client is None:
        _openai_client = OpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL)
    return _openai_client

# 任务队列与缓存
executor = ThreadPoolExecutor(max_workers=4) # 允许同时处理4个分析任务
JOB_CACHE = {} # 存储任务状态和结果

LATEST_FILES: Dict[int, list] = {}  # group_id -> list of cached items, index 0 = oldest
CONTENT_INDEX: Dict[str, str] = {}
LINK_CACHE: Dict[int, list] = {}  # group_id -> list of link items, index 0 = oldest
MAX_BRIDGE_LINKS_PER_GROUP = 30
LOG_IMPORTED_FILES: Dict[str, set] = {}  # log_id_str -> set of file_id strings already imported
LAST_ERROR_BY_GROUP: Dict[int, str] = {}
LAST_EVENT_SUMMARY: Dict[str, Any] = {}
LAST_NAPCAT_BASE = ""
LAST_NAPCAT_ERROR = ""
UPLOAD_STATES: Dict[int, Dict[str, Any]] = {}
STATE_LOCK = threading.RLock()
UPLOAD_QUEUE = Queue(maxsize=BRIDGE_QUEUE_SIZE)
UPLOAD_WORKER: Optional[threading.Thread] = None

# v4.4.0: Job cancellation support
CANCEL_FLAGS: Dict[str, threading.Event] = {}

# v4.4.0: History system — preserves evicted bridge items on disk
HISTORY = []  # global history list, index 0 = newest
MAX_HISTORY_ITEMS = 50
HISTORY_FILE = os.path.join(BRIDGE_CACHE_DIR, "history.json")

# --- WebSocket 实时监听 (logutil recording) ---
# Create a dummy event loop just to allocate the queue; the real loop replaces it in run_ws_event_loop.
_ws_init_loop = asyncio.new_event_loop()
WS_MESSAGE_QUEUE: "asyncio.Queue" = asyncio.Queue()  # type: ignore[abstract]
_ws_init_loop.close()
WS_CLIENT: Optional[Any] = None  # LogutilBotClient instance, set in run_ws_event_loop
WS_WORKER: Optional[threading.Thread] = None

PAINTER_SERVERS = [
    'https://s02.trpgbot.com/s/',
    'https://s03.trpgbot.com/models/',
    'https://api.dice.center/dicelogger/'
]
KOKONA_BASE_URL = "https://dicelogger.s3-accelerate.amazonaws.com/"
LOGUTIL_DB_FILE = os.path.join(os.path.dirname(__file__), "logutil.db")
STORY_PAINTER_UPLOAD_URL = os.getenv("LOGAI_STORY_PAINTER_UPLOAD_URL", "https://weizaima.com/dice/api/log")
STORY_PAINTER_TOKEN = os.getenv("LOGAI_STORY_PAINTER_TOKEN", "")
STORY_PAINTER_TIMEOUT_SEC = int(os.getenv("LOGAI_STORY_PAINTER_TIMEOUT_SEC", "60"))
STORY_PAINTER_VERSION = 101
STORY_PAINTER_CLIENT = "SealDice"

# ====== 文件去重状态 ======
RECENT_FILE_CAPTURES = {}
RECENT_FILE_DEDUPE_WINDOW_SEC = 1
RECENT_FILE_CAPTURE_TTL_SEC = 600

# Upload-level dedup: prevent same (group_id, file_id) from flooding the queue
RECENT_UPLOAD_EVENTS = {}

# Polling: groups that should be periodically checked for new NapCat files
BRIDGE_POLL_GROUPS = set()  # set of int group_id

# --- 系统工具 ---
def disable_quick_edit():
    """禁用Windows快速编辑模式防挂起"""
    if platform.system() == "Windows":
        try:
            kernel32 = ctypes.windll.kernel32
            hInput = kernel32.GetStdHandle(-10)
            mode = ctypes.c_ulong()
            kernel32.GetConsoleMode(hInput, ctypes.byref(mode))
            mode.value &= ~0x0040
            mode.value &= ~0x0020
            kernel32.SetConsoleMode(hInput, mode)
        except: pass

def get_session():
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Connection": "keep-alive"
    })
    retries = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
    session.mount("https://", HTTPAdapter(max_retries=retries))
    return session

def safe_decode(byte_content):
    if not byte_content: return ""
    for encoding in ['utf-8', 'gb18030', 'big5']:
        try: return byte_content.decode(encoding)
        except: pass
    return byte_content.decode('utf-8', errors='ignore')


def get_token_usage_suffix(resp):
    usage = getattr(resp, 'usage', None)
    total_tokens = getattr(usage, 'total_tokens', None)
    return f" | Tokens: {total_tokens}" if total_tokens is not None else ""


def get_pymupdf_blocks(text_dict):
    if not isinstance(text_dict, dict):
        return []
    blocks = text_dict.get('blocks', [])
    return blocks if isinstance(blocks, list) else []

# --- 数据获取函数 (复用之前的逻辑) ---
WEIZAIMA_LOG_API = "https://weizaima.com/dice/api/load_data"
DICE_ZONE_LOG_API = "https://log-api.dice.zone/api/dice/load_data"

def normalize_query_key_source(source):
    value = str(source or '').strip().lower()
    if value in ('dice_zone', 'dicezone', 'dice.zone', 'log.dice.zone', 'log-api.dice.zone'):
        return 'dice_zone'
    if value in ('weizaima', 'weizaima.com', 'log.weizaima.com'):
        return 'weizaima'
    return value

def infer_query_key_source_from_host(hostname):
    host = str(hostname or '').strip().lower()
    if not host:
        return 'weizaima'
    if host.endswith('dice.zone'):
        return 'dice_zone'
    if host.endswith('weizaima.com'):
        return 'weizaima'
    return 'weizaima'

def fetch_query_key_log(api_url, key, password=None):
    try:
        resp = get_session().get(api_url, params={"key": key, "password": password}, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            if 'data' in data:
                return json.loads(zlib.decompress(base64.b64decode(data['data'])).decode('utf-8'))
    except Exception as e: print(f"KeyLog Error ({api_url}): {e}")
    return None

def fetch_weizaima(key, password=None):
    return fetch_query_key_log(WEIZAIMA_LOG_API, key, password)

def fetch_dice_zone(key, password=None):
    return fetch_query_key_log(DICE_ZONE_LOG_API, key, password)

def format_weizaima_text(log_obj):
    if not log_obj: return ""
    items = []
    if isinstance(log_obj, dict):
        items = log_obj.get('items') or (log_obj.get('data') or {}).get('items') or log_obj.get('messages') or []
    if not isinstance(items, list):
        return ""
    lines = []
    for item in items[:MAX_LOG_ENTRIES]:
        if not isinstance(item, dict):
            continue
        nick = item.get('nickname') or item.get('nick') or item.get('sender') or item.get('user') or "?"
        im_userid = item.get('IMUserId') or item.get('uniformId') or item.get('user_id') or nick
        msg = item.get('message') or item.get('content') or item.get('text') or item.get('msg') or ""
        item_time = item.get('time')
        # fallback: image or attachment indicators
        if not msg:
            if str(item.get('type', '')).lower() == 'image' or item.get('file'):
                msg = '[图片]'
            else:
                raw = item.get('raw') or item.get('data') or {}
                if isinstance(raw, dict) and (raw.get('type') == 'image' or raw.get('file')):
                    msg = '[图片]'
        if not msg:
            continue
        # replace CQ image tags with placeholder and strip HTML
        msg = re.sub(r'\[CQ:image[^\]]*\]', '[图片]', str(msg))
        msg = re.sub(r'<[^>]+>', '', msg).strip()
        if not msg:
            continue
        # v4.4.5: preserve IMUserId and time via bracket pipe format
        if item_time and im_userid:
            ts_str = datetime.datetime.fromtimestamp(safe_int(item_time, 0)).strftime('%Y-%m-%d %H:%M:%S')
            if im_userid == nick:
                # 格式 b: 玩家本人发言 → <玩家昵称|游戏外>
                lines.append(f"[{ts_str}] <{nick}|游戏外> {msg}")
            else:
                # 格式 a: 角色发言 → <角色名|玩家昵称>
                lines.append(f"[{ts_str}] <{nick}|{im_userid}> {msg}")
        elif item_time:
            ts_str = datetime.datetime.fromtimestamp(safe_int(item_time, 0)).strftime('%Y-%m-%d %H:%M:%S')
            lines.append(f"[{ts_str}] <{nick}|{nick}> {msg}")
        else:
            lines.append(f"<{nick}|{nick}> {msg}")
    return "\n".join(lines)

def fetch_trpgbot(full_id):
    try:
        sid, log_id = full_id.split('-', 1)
        base_url = PAINTER_SERVERS[int(sid)]
        sess = get_session()
        sess.headers.update({"Referer": "https://logpainter.trpgbot.com/"})
        meta = sess.get(f"{base_url}logReader.php", params={"m": "metaData", "id": log_id, "r": 0.1}, timeout=20).json()
        dl_url = meta.get('redirectDownloadUrl') or f"{base_url}logReader.php?m=rawData&id={log_id}"
        return safe_decode(sess.get(dl_url, timeout=90).content)
    except Exception as e: print(f"TRPGBot Error: {e}"); return None

def fetch_kokona(s3_key):
    try:
        resp = get_session().get(f"{KOKONA_BASE_URL}{s3_key}", timeout=60)
        return safe_decode(resp.content) if resp.status_code == 200 else None
    except Exception as e: print(f"Kokona Error: {e}"); return None

def fetch_raw_url(url):
    try:
        resp = get_session().get(url, timeout=120)
        return safe_decode(resp.content) if resp.status_code == 200 else None
    except Exception as e: print(f"RawURL Error: {e}"); return None

def expand_short_alias(raw):
    """v4.4.4: 短别名展開 — F14→[file]-14, L0→[link]-0, H23→[history]-23.
    v4.4.4.1: 也支持跨群访问 F14-123456→[file]-14-123456。
    与前端 expandShortAlias 完全一致。"""
    s = str(raw or '').strip()
    # 跨群短别名: F14-123456, L0-999888
    m = re.match(r'^([FLH])(\d+)-(\d+)$', s, re.I)
    if m:
        prefix = m.group(1).upper()
        num = m.group(2)
        gid = m.group(3)
        if prefix == 'F': return f'[file]-{num}-{gid}'
        if prefix == 'L': return f'[link]-{num}-{gid}'
        if prefix == 'H': return f'[history]-{num}-{gid}'
    # 短别名: F14, L0, H23
    m = re.match(r'^([FLH])(\d+)$', s, re.I)
    if m:
        prefix = m.group(1).upper()
        num = m.group(2)
        if prefix == 'F': return f'[file]-{num}'
        if prefix == 'L': return f'[link]-{num}'
        if prefix == 'H': return f'[history]-{num}'
    # v4.4.4: SealDice 去括号修复 — file-0→[file]-0
    m = re.match(r'^(file|link|history)-(\d+)$', s, re.I)
    if m:
        return f'[{m.group(1).lower()}]-{m.group(2)}'
    return raw


def parse_log_target_entry(raw, password=None, source=None):
    """Parse a raw log reference into a target dict {key, source, password}.
    Matches fwlog's parse_log_target_entry exactly for non-bridge sources,
    with bridge_file / bridge_file_name / bridge_content extensions."""
    value = str(raw or '').strip()
    if not value:
        return None

    # v4.4.4: expand short aliases (F14→[file]-14, L0→[link]-0, H23→[history]-23)
    value = expand_short_alias(value)

    url_match = URL_RE.search(value)
    if url_match:
        value = url_match.group(0)

    # --- logai bridge extensions (not in fwlog) ---
    # Direct bridge content URL (raw text cached under /bridge/content/...)
    if '/bridge/content/' in value:
        return {'key': value, 'source': 'raw_url', 'password': ''}

    # [file]-N pattern: reference bridge-cached files by index (0=oldest, higher=newer)
    # v4.4.5: also match optional cross-group suffix [file]-N-GID
    file_idx_match = re.match(r'^\[file\]-(\d+)(?:-(\d+))?(?:\s|$)', value, re.IGNORECASE)
    if file_idx_match:
        result = {'key': file_idx_match.group(1), 'source': 'bridge_file', 'password': ''}
        if file_idx_match.group(2):
            result['cross_group_id'] = file_idx_match.group(2)
        return result

    # v4.4.0: [link]-N pattern: reference bridge-cached link text by index
    link_idx_match = re.match(r'^\[link\]-(\d+)(?:-(\d+))?(?:\s|$)', value, re.IGNORECASE)
    if link_idx_match:
        result = {'key': link_idx_match.group(1), 'source': 'bridge_link', 'password': ''}
        if link_idx_match.group(2):
            result['cross_group_id'] = link_idx_match.group(2)
        return result

    # v4.4.3: [history]-N pattern: reference evicted bridge items by index
    history_idx_match = re.match(r'^\[history\]-(\d+)(?:-(\d+))?(?:\s|$)', value, re.IGNORECASE)
    if history_idx_match:
        result = {'key': history_idx_match.group(1), 'source': 'bridge_history', 'password': ''}
        if history_idx_match.group(2):
            result['cross_group_id'] = history_idx_match.group(2)
        return result

    # Bare file name: look like filenames (not URLs, not known key patterns)
    # Examples: "[2026-06-11_11-25]8月23日营地.txt", "8月23日营地.txt", "8月23日营地"
    if not re.match(r'^https?://', value, re.IGNORECASE) and '://' not in value:
        if '=' not in value and '#' not in value and not value.isdigit():
            looks_like_filename = (
                value.startswith('[') or
                bool(re.search(r'\.\w{2,5}$', value)) or
                bool(re.search(r'^\[\d{4}-\d{2}-\d{2}[_\s]', value)) or
                # Bare name without extension: Chinese chars, short enough to be a filename (not a sentence)
                (bool(re.search(r'[一-鿿]', value)) and len(value) <= 40 and not re.search(r'[，。！？；：、\n]', value))
            )
            if looks_like_filename:
                return {'key': value, 'source': 'bridge_file_name', 'password': ''}

    # --- fwlog-compatible URL parsing ---
    parsed = urllib.parse.urlparse(value)
    query = urllib.parse.parse_qs(parsed.query)
    key = ''
    resolved_source = normalize_query_key_source(source)
    resolved_password = str(password or '').strip()

    if query.get('s3'):
        key = query['s3'][0]
        if not resolved_source:
            resolved_source = 'kokona'
    elif query.get('key'):
        key = query['key'][0]
        if not resolved_source:
            resolved_source = infer_query_key_source_from_host(parsed.hostname)
        if parsed.fragment and not resolved_password:
            resolved_password = parsed.fragment
    elif parsed.fragment:
        fragment_key = re.sub(r'[^a-zA-Z0-9-_]', '', parsed.fragment)
        if fragment_key:
            key = fragment_key
            if '-' in key and not resolved_source:
                resolved_source = 'trpgbot'
    else:
        key = value

    if not key:
        return None

    if not resolved_source:
        resolved_source = infer_source_by_key(key)

    return {'key': key, 'source': resolved_source, 'password': resolved_password}

def infer_source_by_key(key):
    value = str(key or '').strip()
    # Filenames with bracket timestamps: [YYYY-MM-DD_HH-MM]xxx.ext
    if value.startswith('[') and re.search(r'^\s*\[\d{4}-\d{2}-\d{2}[_\s]', value):
        return 'bridge_file_name'
    # Filenames with extensions that aren't URLs (v4.3: 匹配任意文本类扩展名)
    if '.' in value and not value.startswith(('http://', 'https://')):
        ext_match = re.search(r'\.(\w{1,10})$', value, re.IGNORECASE)
        if ext_match:
            return 'bridge_file_name'
    if value.startswith(('http://', 'https://')):
        parsed = urllib.parse.urlparse(value)
        query = urllib.parse.parse_qs(parsed.query)
        if '/bridge/content/' in value:
            return 'raw_url'
        if query.get('s3'):
            return 'kokona'
        if query.get('key'):
            return infer_query_key_source_from_host(parsed.hostname)
        if parsed.hostname and parsed.hostname.endswith('dice.zone'):
            return 'dice_zone'
        if parsed.fragment:
            fragment_key = re.sub(r'[^a-zA-Z0-9-_]', '', parsed.fragment)
            if fragment_key and '-' in fragment_key:
                return 'trpgbot'
        return 'raw_url'
    if value and '-' in value and value.split('-')[0].isdigit():
        return "trpgbot"
    if value and ('_' in value or len(value) > 20):
        return "kokona"
    return "weizaima"

def fetch_log_text_by_source(key, password=None, source=None, group_id=None):
    """Fetch log text from the appropriate source.  Matches fwlog's behaviour
    for weizaima / dice_zone / kokona / trpgbot / raw_url, plus logai bridge
    extensions (bridge_file / bridge_file_name)."""
    target = parse_log_target_entry(key, password=password, source=source)
    if not target:
        return ""

    resolved_key = target['key']
    resolved_password = target.get('password') or ''
    resolved_source = normalize_query_key_source(target.get('source') or infer_source_by_key(resolved_key))
    # v4.4.5: cross-group override for bridge sources
    cross_gid_raw = target.get('cross_group_id')
    bridge_gid = safe_int(cross_gid_raw, 0) if cross_gid_raw else 0

    # --- fwlog-compatible sources ---
    # v4.4.0: cache fetched link text to bridge
    gid_for_cache = safe_int(group_id, 0) if group_id else 0
    if resolved_source == "kokona":
        raw_text = fetch_kokona(resolved_key)
        result = format_raw_text(raw_text)
        if result and gid_for_cache > 0:
            write_link_cache(gid_for_cache, resolved_key, result)
        return result
    if resolved_source == "trpgbot":
        raw_text = fetch_trpgbot(resolved_key)
        result = format_raw_text(raw_text)
        if result and gid_for_cache > 0:
            write_link_cache(gid_for_cache, resolved_key, result)
        return result
    if resolved_source == "raw_url":
        raw_text = fetch_raw_url(resolved_key)
        result = format_raw_text(raw_text)
        if result and gid_for_cache > 0:
            write_link_cache(gid_for_cache, resolved_key, result)
        return result
    if resolved_source == "dice_zone":
        raw_text = fetch_dice_zone(resolved_key, resolved_password)
        result = format_weizaima_text(raw_text)
        if result and gid_for_cache > 0:
            write_link_cache(gid_for_cache, resolved_key, result)
        return result
    if resolved_source == "weizaima":
        raw_text = fetch_weizaima(resolved_key, resolved_password)
        result = format_weizaima_text(raw_text)
        if result and gid_for_cache > 0:
            write_link_cache(gid_for_cache, resolved_key, result)
        return result

    # --- logai bridge extensions (not present in fwlog) ---
    if resolved_source == "bridge_file_name":
        gid = bridge_gid or (safe_int(group_id, 0) if group_id else 0)
        if gid > 0:
            with STATE_LOCK:
                file_list = list(LATEST_FILES.get(gid, []))
            normalized_search = re.sub(r'\s+', '', str(resolved_key))
            # Strategy 1: exact normalized match
            for item in file_list:
                item_name = str(item.get('name', ''))
                normalized_item_name = re.sub(r'\s+', '', item_name)
                if normalized_item_name and normalized_search and normalized_item_name == normalized_search:
                    ck = item.get('content_key', '')
                    with STATE_LOCK:
                        path = CONTENT_INDEX.get(ck, '')
                    if path and os.path.exists(path):
                        with open(path, 'r', encoding='utf-8') as f:
                            return format_raw_text(f.read())
            # Strategy 2: strip extensions from cached names, compare base names
            for item in file_list:
                item_name = str(item.get('name', ''))
                item_base = re.sub(r'\.\w{2,5}$', '', item_name)
                normalized_base = re.sub(r'\s+', '', item_base)
                if normalized_base and normalized_search and normalized_base == normalized_search:
                    ck = item.get('content_key', '')
                    with STATE_LOCK:
                        path = CONTENT_INDEX.get(ck, '')
                    if path and os.path.exists(path):
                        with open(path, 'r', encoding='utf-8') as f:
                            return format_raw_text(f.read())
            # Strategy 3: substring containment (search term inside cached name)
            if normalized_search:
                for item in file_list:
                    item_name = str(item.get('name', ''))
                    normalized_item_name = re.sub(r'\s+', '', item_name)
                    if normalized_search in normalized_item_name:
                        ck = item.get('content_key', '')
                        with STATE_LOCK:
                            path = CONTENT_INDEX.get(ck, '')
                        if path and os.path.exists(path):
                            with open(path, 'r', encoding='utf-8') as f:
                                return format_raw_text(f.read())
        return ""
    if resolved_source == "bridge_file":
        idx = safe_int(resolved_key, 0)
        gid = bridge_gid or (safe_int(group_id, 0) if group_id else 0)
        if gid > 0:
            with STATE_LOCK:
                file_list = list(LATEST_FILES.get(gid, []))
            if 0 <= idx < len(file_list):
                item = file_list[idx]
                ck = item.get('content_key', '')
                with STATE_LOCK:
                    path = CONTENT_INDEX.get(ck, '')
                if path and os.path.exists(path):
                    with open(path, 'r', encoding='utf-8') as f:
                        return format_raw_text(f.read())
        return ""

    # v4.4.0: [link]-N support
    if resolved_source == "bridge_link":
        idx = safe_int(resolved_key, 0)
        gid = bridge_gid or (safe_int(group_id, 0) if group_id else 0)
        if gid > 0:
            with STATE_LOCK:
                link_list = list(LINK_CACHE.get(gid, []))
            if 0 <= idx < len(link_list):
                item = link_list[idx]
                ck = item.get('content_key', '')
                with STATE_LOCK:
                    path = CONTENT_INDEX.get(ck, '')
                if path and os.path.exists(path):
                    with open(path, 'r', encoding='utf-8') as f:
                        return format_raw_text(f.read())
        return ""

    # v4.4.3: [history]-N support
    if resolved_source == "bridge_history":
        idx = safe_int(resolved_key, 0)
        gid = bridge_gid or (safe_int(group_id, 0) if group_id else 0)
        with STATE_LOCK:
            hist_list = [h for h in HISTORY if safe_int(h.get('group_id', 0), 0) == gid] if gid > 0 else list(HISTORY)
        if 0 <= idx < len(hist_list):
            item = hist_list[idx]
            ck = item.get('content_key', '')
            with STATE_LOCK:
                path = CONTENT_INDEX.get(ck, '')
            if path and os.path.exists(path):
                with open(path, 'r', encoding='utf-8') as f:
                    return format_raw_text(f.read())
        return ""

    return format_weizaima_text(fetch_weizaima(resolved_key, resolved_password))

def format_raw_text(raw_text):
    if not raw_text: return ""
    lines = raw_text.split('\n')
    # Only convert simple <Name> text, NOT <Name|Player> text (bracket-pipe format)
    pattern = re.compile(r'<([^|>]+)>\s*(.*)')
    clean = []
    for line in lines[:MAX_LOG_ENTRIES]:
        line = line.strip()
        if not line: continue
        m = pattern.search(line)
        if m: clean.append(f"{m.group(1)}: {m.group(2).strip()}")
        else:
            clean.append(line)
    return "\n".join(clean)

# --- v4.2: 直接文本处理（用于.ai无文件模式）---
def background_process_direct_text(job_id, direct_text, is_pro=False, is_kind=False, mode='analyze', persona="", custom_prompt="", theme='default', get_text=False, group_id=0):
    """后台线程：直接使用提供的文本进行AI分析（无需抓取URL）"""
    print(f"[{job_id}] 开始直接文本分析... 文本长度: {len(direct_text)}")
    try:
        log_text = str(direct_text or '')
        if not log_text.strip():
            raise Exception("文本内容为空")

        # 智能截断
        if len(log_text) > MAX_AI_CHARS:
            part = int(MAX_AI_CHARS * 0.4)
            mid = log_text[part:-part].split('\n')
            step = max(1, int(len(mid)/100))
            log_text_ai = f"{log_text[:part]}\n...[略]...\n{chr(10).join(mid[::step])}\n{log_text[-part:]}"
        else:
            log_text_ai = log_text

        # Prompt选择 (v4.3: .ai 直接文本模式使用中性prompt，不混入log评分模板)
        report_title = "AI 分析结果"
        if custom_prompt:
            system_prompt = custom_prompt
            report_title = "TRPG 自定义分析报告"
        elif mode == 'recap':
            report_title = "TRPG 跑团前文回顾"
            system_prompt = """你是一个专业且细致的 TRPG 跑团记录员（书记）。请阅读以下跑团 Log，为 KP 和玩家梳理一份详细的【前文回顾】..."""
        else:
            # .ai 无文件模式 / 通用文本分析：使用中性系统提示
            system_prompt = "你是一个专业且全面的AI助手。请仔细阅读用户提供的内容，给出详尽、准确的分析和回答。如有需要，请使用【分页符】进行内容分页。"

        if persona:
            system_prompt += f"\n\n【极其重要的扮演指令】：\n在生成上述所有评价和梳理内容时，请你完全带入以下角色人设来进行语气和口吻的渲染。...\n{persona}"

        if theme == 'default':
            system_prompt += "\n\n【排版指令】：你可以根据当前内容的故事氛围，在回复的【最开头】加上标签以控制最终生成的图片风格。支持的标签有：【主题：经典】、【主题：克苏鲁】、【主题：赛博】、【主题：历史】、【主题：废土】、【主题：二次元】、【主题：终端】。如果你觉得不需要特殊风格，可不写此标签。"

        print(f"[{job_id}] 开始请求 LLM...")
        # v4.4.0: 检查取消标志
        if CANCEL_FLAGS.get(job_id) and CANCEL_FLAGS[job_id].is_set():
            JOB_CACHE[job_id]['status'] = 'cancelled'
            JOB_CACHE[job_id]['text'] = '任务已被取消。'
            return
        model = AI_MODEL_PRO if is_pro else AI_MODEL
        resp = get_openai_client().chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": log_text_ai}],
            temperature=1.0, max_tokens=65535
        )
        result_text = resp.choices[0].message.content or ""
        token_usage = get_token_usage_suffix(resp)

        result_text, final_theme = extract_theme_from_text(result_text, theme)

        images_list = text_to_images(result_text, f"DirectText:{job_id[:8]}", report_title, final_theme, token_usage)
        JOB_CACHE[job_id]['status'] = 'done'
        JOB_CACHE[job_id]['images'] = images_list
        JOB_CACHE[job_id]['text'] = result_text
        print(f"[{job_id}] 直接文本处理完成")

        # get_text模式：保存文本文件
        if get_text:
            try:
                ensure_bridge_cache_dir()
                text_key = uuid.uuid4().hex
                text_path = os.path.join(BRIDGE_CACHE_DIR, f"{text_key}.txt")
                safe_filename = re.sub(r'[\\/*?:"<>|]', '_', f"ai_analysis_{job_id[:8]}.txt")
                with open(text_path, 'w', encoding='utf-8') as fw:
                    fw.write(result_text)
                with STATE_LOCK:
                    CONTENT_INDEX[text_key] = text_path
                public_base = resolve_public_base_or_fallback()
                JOB_CACHE[job_id]['text_key'] = text_key
                JOB_CACHE[job_id]['text_filename'] = safe_filename
                print(f"[{job_id}] get_text 文件已保存: {safe_filename}")
                # 自动上传到群（仿 send_log_via_napcat 机制）
                if group_id > 0:
                    try:
                        fs, _ = napcat_upload_group_file(group_id, text_path, safe_filename)
                        JOB_CACHE[job_id]['text_file_sent'] = fs
                        if fs:
                            print(f"[{job_id}] get_text 文件已上传到群 {group_id}")
                    except Exception as ue:
                        print(f"[{job_id}] get_text 上传失败: {ue}")
            except Exception as e:
                print(f"[{job_id}] get_text 保存失败: {e}")

    except Exception as e:
        print(f"[{job_id}] 直接文本处理失败: {e}")
        err_img_bytes = text_to_images(f"AI处理失败：\n{str(e)}", "Error")[0]
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['images'] = [err_img_bytes]
        JOB_CACHE[job_id]['text'] = f"处理失败：{str(e)}"

# --- 核心处理任务 ---
def background_process(job_id, key, password, source, is_pro=False, is_kind=False, mode='analyze', persona="", custom_prompt="", theme='default', group_id=0, get_text=False):
    """后台线程：执行 Log 下载、分析、绘图"""
    print(f"[{job_id}] 开始处理Log... Source: {source}, Mode: {mode}")
    try:
        # ================= 1. 尝试触发省流缓存 =================
        hash_key = None
        key_for_label = "multi"
        if not is_pro:
            # 只有普通模式参与缓存，确保同一个Log和同样的提示配置拥有唯一签名
            hash_str = f"url_log_{key}_{mode}_{is_kind}_{persona}_{custom_prompt}_{theme}"
            hash_key = hashlib.md5(hash_str.encode('utf-8')).hexdigest()
            cached_images = get_daily_cache(hash_key)
            if cached_images:
                print(f"[{job_id}] 命中今日缓存库！省流模式启动，秒回历史图片。")
                JOB_CACHE[job_id]['status'] = 'done'
                JOB_CACHE[job_id]['images'] = cached_images
                return
        # ========================================================
        log_text = ""
        if isinstance(key, list):
            parts = []
            for i, one_key in enumerate(key):
                one_key = str(one_key or '').strip()
                if not one_key:
                    continue
                one_source = source[i] if isinstance(source, list) and i < len(source) else (source if isinstance(source, str) else None)
                one_pwd = password[i] if isinstance(password, list) and i < len(password) else (password if isinstance(password, str) else None)
                print(f"[{job_id}] fetching segment {i+1}: key={one_key!r} source_hint={one_source!r} pwd_present={'yes' if one_pwd else 'no'}")
                one_text = fetch_log_text_by_source(one_key, one_pwd, one_source, group_id=group_id)
                if not one_text:
                    print(f"[{job_id}] segment {i+1} fetch empty for key={one_key!r} source_hint={one_source!r}")
                    continue
                print(f"[{job_id}] segment {i+1} fetched length={len(one_text)} preview={one_text[:200]!r}")
                source_name = one_source or infer_source_by_key(one_key)
                parts.append(f"【第{i+1}段日志 | 来源:{source_name} | 标识:{one_key[:20]}】\n{one_text}")

            if parts:
                log_text = "\n\n===== 多段日志拼接分隔线 =====\n\n".join(parts)
            key_for_label = f"multi:{len(parts)}"
        else:
            print(f"[{job_id}] fetching single key: key={key!r} source_hint={source!r} pwd_present={'yes' if password else 'no'}")
            log_text = fetch_log_text_by_source(key, password, source, group_id=group_id)
            print(f"[{job_id}] single fetch result length={len(log_text) if log_text else 0}")
            key_for_label = str(key or 'single')
        
        if not log_text:
            raise Exception("日志内容获取失败或为空")

        # 智能截断防爆 Token
        if len(log_text) > MAX_AI_CHARS:
            part = int(MAX_AI_CHARS * 0.4)
            mid = log_text[part:-part].split('\n')
            step = max(1, int(len(mid)/100))
            log_text_ai = f"{log_text[:part]}\n...[略]...\n{chr(10).join(mid[::step])}\n{log_text[-part:]}"
        else:
            log_text_ai = log_text

        # 核心：根据不同模式分配对应的 Prompt
        report_title = "TRPG 跑团日志分析"
        # 【新增】：如果有自定义提示词，强行覆盖，并把标题改为自定义
        if custom_prompt:
            system_prompt = custom_prompt
            report_title = "TRPG 自定义分析报告"
        
        elif mode == 'recap':
            report_title = "TRPG 跑团前文回顾"
            system_prompt = """你是一个专业且细致的 TRPG 跑团记录员（书记）。请阅读以下跑团 Log，为 KP 和玩家梳理一份详细的【前文回顾】，帮助大家快速找回跑团记忆。
请严格按照以下 4 个板块输出，并且在输出每个大板块之前，必须使用“【分页符】”这四个字单起一行作为分隔标识（不要用Markdown，不要加粗）：

【分页符】
【一、当前剧情进度总览】：
（详细说明截至目前的故事进度，大家在哪，正在面临什么状况，遇到了什么危机或主线推进到了哪一步）
【分页符】
【二、PC行动轨迹与状态梳理】：
（尽可能详细分条列出每位主要玩家角色/PC近期做了什么举动，达成了什么目的，或处于什么特殊状态/受到什么伤害）
【分页符】
【三、当前已获线索与道具盘点】：
（总结当前大家掌握的所有情报、未解之谜、NPC给出的重要信息以及拿到的关键道具）
【分页符】
【四、下一步推进方向提示】：
（基于当前局势，客观给出2-3个可供调查员们继续推进剧情的可能方向或需要立刻解决的问题）"""

        else: # 默认的 analyze 评分分析
            report_title = "TRPG 跑团日志评分"
            if is_kind: system_prompt = KIND_SYSTEM_PROMPT
            elif is_pro: system_prompt = PRO_SYSTEM_PROMPT
            else: system_prompt = DEFAULT_SYSTEM_PROMPT
        
        # 核心：人设系统劫持（强制带入骰娘语气且防止格式崩溃）
        if persona:
            system_prompt += f"\n\n【极其重要的扮演指令】：\n在生成上述所有评价和梳理内容时，请你完全带入以下角色人设来进行语气和口吻的渲染。你可以自称、吐槽或撒娇，让输出充满该人设的个性。\n（绝对警告：你必须严格保留前文要求的【板块标题】和【分页符】等格式标识符，千万不能省略或修改它们，只能改变正文部分的说话语气！）：\n{persona}"
        
        # 【新增】：如果用户没指定主题，赋予大模型绝对的主题控制权！
        if theme == 'default':
            system_prompt += "\n\n【排版指令】：你可以根据当前内容的故事氛围，在回复的【最开头】加上标签以控制最终生成的图片风格。支持的标签有：【主题：经典】、【主题：克苏鲁】、【主题：赛博】、【主题：历史】、【主题：废土】、【主题：二次元】、【主题：终端】。如果你觉得不需要特殊风格，可不写此标签。"

        # 5. 请求 AI
        print(f"[{job_id}] 未命中缓存，开始请求 LLM 消耗 Token...")
        # v4.4.0: 检查取消标志
        if CANCEL_FLAGS.get(job_id) and CANCEL_FLAGS[job_id].is_set():
            JOB_CACHE[job_id]['status'] = 'cancelled'
            JOB_CACHE[job_id]['text'] = '任务已被取消。'
            return
        model = AI_MODEL_PRO if is_pro else AI_MODEL
        resp = get_openai_client().chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": log_text_ai}],
            temperature=1.0, max_tokens=65535
        )
        result_text = resp.choices[0].message.content or ""
        token_usage = get_token_usage_suffix(resp)
        
        result_text, final_theme = extract_theme_from_text(result_text, theme)

        # 6. 绘图与返回
        images_list = text_to_images(result_text, f"Key:{key_for_label[:8]}", report_title, final_theme, token_usage)
        JOB_CACHE[job_id]['images'] = images_list
        JOB_CACHE[job_id]['text'] = result_text  # 保存原始文本供 get_text 使用
        print(f"[{job_id}] 渲染处理完成")

        # 7b. 如果 get_text 模式，将文本保存为桥接文件
        if get_text:
            try:
                ensure_bridge_cache_dir()
                text_key = uuid.uuid4().hex
                text_path = os.path.join(BRIDGE_CACHE_DIR, f"{text_key}.txt")
                safe_filename = re.sub(r'[\\/*?:"<>|]', '_', f"ai_analysis_{job_id[:8]}.txt")
                with open(text_path, 'w', encoding='utf-8') as fw:
                    fw.write(result_text)
                with STATE_LOCK:
                    CONTENT_INDEX[text_key] = text_path
                public_base = resolve_public_base_or_fallback()
                JOB_CACHE[job_id]['text_key'] = text_key
                JOB_CACHE[job_id]['text_filename'] = safe_filename
                print(f"[{job_id}] get_text 文件已保存: {safe_filename}")
                # 自动上传到群（仿 send_log_via_napcat 机制）
                if group_id > 0:
                    try:
                        fs, _ = napcat_upload_group_file(group_id, text_path, safe_filename)
                        JOB_CACHE[job_id]['text_file_sent'] = fs
                        if fs:
                            print(f"[{job_id}] get_text 文件已上传到群 {group_id}")
                    except Exception as ue:
                        print(f"[{job_id}] get_text 上传失败: {ue}")
            except Exception as e:
                print(f"[{job_id}] get_text 保存失败: {e}")

        # v4.4.4.1: set status done AFTER get_text to avoid race condition
        JOB_CACHE[job_id]['status'] = 'done'

        # ================= 7. 写入省流缓存 =================
        if not is_pro and hash_key:
            set_daily_cache(hash_key, images_list)
            print(f"[{job_id}] 结果已存入今日缓存库。")

    except Exception as e:
        print(f"[{job_id}] 失败: {e}")
        err_img_bytes = text_to_images(f"Log处理失败：\n{str(e)}", "Error")[0]
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['images'] = [err_img_bytes]
        JOB_CACHE[job_id]['text'] = f"处理失败：{str(e)}"

def extract_text_from_file(file_content, filename):
    """根据文件扩展名提取文本，增强容错能力"""
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    
    try:
        file_stream = BytesIO(file_content)
        
        if ext in ['.txt', '.md', '.json', '.yaml', '.yml']:
            text = safe_decode(file_content)
            
        elif ext == '.docx':
            doc = Document(file_stream)
            text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text and para.text.strip()])
            
        elif ext == '.pdf':
            # 改用 pymupdf (fitz) 读取PDF，容错率极高，会自动忽略纯图片
            import fitz
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            pages_text = []
            # 限制读取前300页防撑爆内存
            for page_num in range(min(len(pdf_document), 300)):
                page_text = str(pdf_document[page_num].get_text() or '').strip()
                if page_text:
                    pages_text.append(page_text)
            text = "\n".join(pages_text)
            pdf_document.close()
            
        else:
            # v4.3.4: 未知扩展名直接当作文本解码（.py, .js, .c 等）
            text = safe_decode(file_content)
            if not text or not text.strip():
                return f"[ParseError]不支持的文件格式: {ext}（且无法作文本解码）"
            
    except Exception as e:
        # 加上特殊前缀，方便外层精准拦截
        return f"[ParseError]文件读取损坏 ({str(e)})\n可能是文件过大或本身已损坏。"
        
    return text


def ensure_bridge_cache_dir():
    os.makedirs(BRIDGE_CACHE_DIR, exist_ok=True)


def now_ts():
    return int(time.time())


def bridge_log(tag, message):
    print(f"[bridge][{time.strftime('%Y-%m-%d %H:%M:%S')}] {tag}: {message}")


def snapshot_item(item):
    if not item:
        return None
    return dict(item)


def safe_int(value, default=0):
    try:
        return int(value)
    except Exception:
        return default


# ====== fwlog 移植：发言识别辅助函数 ======

def make_log_item(nickname, im_userid, ts, message, raw_msg_id):
    """构建标准 log_item dict"""
    return {
        "nickname": nickname or "Unknown",
        "im_userid": str(im_userid or ""),
        "time": safe_int(ts, int(time.time())),
        "message": str(message or ""),
        "raw_msg_id": str(raw_msg_id or ""),
    }


def looks_like_speaker_name(name):
    """验证候选发言者名是否有效"""
    candidate = str(name or "").strip()
    if not candidate or len(candidate) > 80:
        return False
    lowered = candidate.lower()
    if lowered.startswith("http") or "://" in candidate:
        return False
    if candidate.startswith("CQ:") or "/" in candidate or "\\" in candidate:
        return False
    if re.fullmatch(r"\d+", candidate):
        return False
    return True


def parse_datetime_parts(date_text, clock_text, fallback_ts):
    """
    解析日期时间字符串为 Unix 时间戳。
    支持 YYYY/MM/DD（fwlog原格式）和 YYYY-MM-DD（新增方括号格式）。
    """
    if not date_text or not clock_text:
        return safe_int(fallback_ts, int(time.time()))
    # YYYY/MM/DD 格式
    slash_fmts = ["%Y/%m/%d %H:%M:%S", "%Y/%m/%d %H:%M"]
    # YYYY-MM-DD 格式（新增）
    dash_fmts = ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M"]
    combined = f"{date_text} {clock_text}"
    for fmt in dash_fmts + slash_fmts:
        try:
            return int(time.mktime(time.strptime(combined, fmt)))
        except Exception:
            continue
    return safe_int(fallback_ts, int(time.time()))


def split_name_and_user_id(name):
    """从 Name(12345) 格式中分离名称和 QQ 号"""
    candidate = str(name or "").strip()
    match = re.match(r"^(?P<name>.*?)(?:\((?P<user_id>\d+)\))?$", candidate)
    if not match:
        return candidate, ""
    plain_name = str(match.group("name") or "").strip()
    user_id = str(match.group("user_id") or "").strip()
    return plain_name, user_id


def build_speaker_match(name, content, fallback_ts, date_text="", clock_text="",
                        bracket_pipe_player=None, is_bracket_pipe=False, is_asterisk=False):
    """
    构建发言匹配结果 dict。

    新增参数：
      bracket_pipe_player: 方括号管道格式中 | 右侧的玩家昵称（格式 a/b）
      is_bracket_pipe: 是否为方括号管道格式 [time] <name|player> text
      is_asterisk:     是否为星号格式 [time] * name|player text
    """
    if is_bracket_pipe and bracket_pipe_player is not None:
        # 格式 a/b: [time] <角色名|玩家昵称> text
        player = str(bracket_pipe_player or "").strip()
        char_name = str(name or "").strip()
        if not looks_like_speaker_name(char_name):
            return None
        ts_val = parse_datetime_parts(date_text, clock_text, fallback_ts)
        if player == "游戏外":
            # 格式 b: <玩家昵称|游戏外> — 玩家昵称同时作为 nickname 和 IMUserID, message 包裹中文括号
            content_str = str(content or "").strip()
            return {
                "name": char_name,
                "user_id": char_name,
                "time": ts_val,
                "content": f"（{content_str}）" if content_str else "",
            }
        else:
            # 格式 a: 角色名→nickname, 玩家昵称→IMUserID（整个字符串）
            return {
                "name": char_name,
                "user_id": player,
                "time": ts_val,
                "content": str(content or "").strip(),
            }

    if is_asterisk and bracket_pipe_player is not None:
        # 格式 c: [time] * 角色名|玩家昵称 text
        player = str(bracket_pipe_player or "").strip()
        char_name = str(name or "").strip()
        if not looks_like_speaker_name(char_name):
            return None
        ts_val = parse_datetime_parts(date_text, clock_text, fallback_ts)
        content_str = str(content or "").strip()
        return {
            "name": char_name,
            "user_id": player,
            "time": ts_val,
            "content": f"{char_name}{content_str}" if content_str else "",
        }

    # 原 fwlog 逻辑
    plain_name, user_id = split_name_and_user_id(name)
    if not looks_like_speaker_name(plain_name):
        return None
    return {
        "name": plain_name,
        "user_id": user_id,
        "time": parse_datetime_parts(date_text, clock_text, fallback_ts),
        "content": str(content or "").strip(),
    }


def parse_metadata_only_line(line, fallback_ts):
    """识别纯时间戳行（无发言者），返回 pending_meta dict 或 None"""
    text = str(line or "").strip()
    if not text:
        return None
    match = LOG_METADATA_LINE_RE.match(text)
    if not match:
        return None
    date_text = str(match.group("date") or "").strip()
    clock_text = str(match.group("clock") or "").strip()
    if not date_text and not clock_text:
        return None
    return {
        "raw": text,
        "date": date_text,
        "clock": clock_text,
        "time": parse_datetime_parts(date_text, clock_text, fallback_ts),
    }


def match_speaker_line(line, fallback_ts=None):
    """
    单行发言匹配。
    匹配顺序：新格式 a/b → 新格式 c → fwlog时间戳尖括号 → fwlog时间戳冒号
    → fwlog尖括号 → fwlog冒号
    """
    text = str(line or "").strip()
    if not text:
        return None

    fallback_value = safe_int(fallback_ts, int(time.time()))

    # 1) 方括号管道格式 [time] <name|player> text  (格式 a/b)
    match = BRACKET_PIPE_SPEAKER_RE.match(text)
    if match:
        return build_speaker_match(
            match.group("name"),
            match.group("content"),
            fallback_value,
            str(match.group("date") or "").strip(),
            str(match.group("clock") or "").strip(),
            bracket_pipe_player=match.group("player"),
            is_bracket_pipe=True,
        )

    # 2) 星号格式 [time] * name|player text  (格式 c)
    match = ASTERISK_PIPE_SPEAKER_RE.match(text)
    if match:
        return build_speaker_match(
            match.group("name"),
            match.group("content"),
            fallback_value,
            str(match.group("date") or "").strip(),
            str(match.group("clock") or "").strip(),
            bracket_pipe_player=match.group("player"),
            is_asterisk=True,
        )

    # 3-4) fwlog 时间戳格式
    for pattern in (TIMESTAMPED_ANGLE_SPEAKER_RE, TIMESTAMPED_PLAIN_SPEAKER_RE):
        match = pattern.match(text)
        if match:
            speaker = build_speaker_match(
                match.group("name"),
                match.group("content"),
                fallback_value,
                str(match.group("date") or "").strip(),
                str(match.group("clock") or "").strip(),
            )
            if speaker:
                return speaker

    # 5) fwlog 尖括号格式 <Name> : content
    match = ANGLE_SPEAKER_RE.match(text)
    if match:
        return build_speaker_match(match.group("name"), match.group("content"), fallback_value)

    # 6) fwlog 冒号格式 Name: content
    match = PLAIN_SPEAKER_RE.match(text)
    if not match:
        return None
    return build_speaker_match(match.group("name"), match.group("content"), fallback_value)


def match_multiline_angle_speaker(lines, start_index, fallback_ts, pending_meta=None):
    """
    多行尖括号块匹配：
    <
    角色名（可跨多行）
    >: 对白
    """
    if start_index >= len(lines):
        return None
    if not MULTILINE_ANGLE_OPEN_RE.match(str(lines[start_index] or "")):
        return None

    name_parts = []
    index = start_index + 1
    while index < len(lines):
        text = str(lines[index] or "").strip()
        close_match = MULTILINE_ANGLE_CLOSE_RE.match(text)
        if close_match:
            if not name_parts:
                return None
            name = " ".join(name_parts).strip()
            speaker = build_speaker_match(
                name,
                close_match.group("content"),
                pending_meta["time"] if pending_meta else fallback_ts,
                pending_meta.get("date", "") if pending_meta else "",
                pending_meta.get("clock", "") if pending_meta else "",
            )
            if not speaker:
                return None
            return speaker, index + 1

        if not text or text.startswith("<"):
            return None

        name_parts.append(text)
        index += 1

    return None


def extract_source_ts(item):
    if not isinstance(item, dict):
        return 0

    for key in (
        "source_ts",
        "upload_time",
        "modify_time",
        "create_time",
        "uploadTime",
        "modifyTime",
        "ts",
    ):
        value = item.get(key)
        if value in (None, ""):
            continue
        try:
            return int(value)
        except Exception:
            continue

    return 0


def extract_cached_ts(item):
    if not isinstance(item, dict):
        return 0

    for key in ("cached_ts", "ts"):
        value = item.get(key)
        if value in (None, ""):
            continue
        try:
            return int(value)
        except Exception:
            continue

    return 0


def build_upload_state(info, status="queued", last_error=""):
    source_ts = extract_source_ts(info)
    return {
        "group_id": safe_int(info.get("group_id", 0), 0),
        "user_id": safe_int(info.get("user_id", 0), 0),
        "file_id": str(info.get("file_id", "")),
        "name": str(info.get("name", "未知文件")),
        "busid": safe_int(info.get("busid", 0), 0),
        "url": str(info.get("url") or ""),
        "content_url": "",
        "text_filename": "",
        "text_chars": 0,
        "text_bytes": 0,
        "ts": source_ts if source_ts > 0 else now_ts(),
        "source_ts": source_ts,
        "cached_ts": now_ts(),
        "status": status,
        "processing": status in ("queued", "processing"),
        "last_error": last_error,
    }


def should_refresh_cached_item(remote_info, cached_item):
    if not remote_info:
        return False

    if not cached_item:
        return True

    remote_file_id = str(remote_info.get("file_id", ""))
    cached_file_id = str(cached_item.get("file_id", ""))
    if remote_file_id and cached_file_id and remote_file_id == cached_file_id:
        return False

    remote_ts = extract_source_ts(remote_info)
    cached_ts = extract_source_ts(cached_item)

    if remote_ts > 0 and cached_ts > 0:
        if remote_ts > cached_ts:
            return True
        if remote_ts < cached_ts:
            return False

    return remote_file_id != cached_file_id


def build_napcat_base_candidates():
    candidates = []

    if NAPCAT_API_BASES_RAW:
        for item in NAPCAT_API_BASES_RAW.split(","):
            value = item.strip().rstrip("/")
            if value:
                candidates.append(value)

    if NAPCAT_API_BASE:
        candidates.append(NAPCAT_API_BASE)

    candidates.extend([
        "http://127.0.0.1:8084",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:34567",
        "http://127.0.0.1:6099",
    ])

    deduped = []
    seen = set()
    for value in candidates:
        if value in seen:
            continue
        seen.add(value)
        deduped.append(value)
    return deduped


NAPCAT_BASE_CANDIDATES = build_napcat_base_candidates()


def check_auth(req):
    if not BRIDGE_TOKEN:
        return True
    return req.headers.get("Authorization", "") == BRIDGE_TOKEN


def check_auth_with_query(req):
    if check_auth(req) or not BRIDGE_TOKEN:
        return True
    return str(req.args.get("token", "")) == BRIDGE_TOKEN


def upload_worker_loop():
    bridge_log("worker", "background upload worker started")

    while True:
        info = UPLOAD_QUEUE.get()
        group_id = safe_int(info.get("group_id", 0), 0)
        file_id = str(info.get("file_id", ""))

        try:
            with STATE_LOCK:
                current = UPLOAD_STATES.get(group_id)
                if current and str(current.get("file_id", "")) == file_id:
                    current["status"] = "processing"
                    current["processing"] = True

            cleanup_expired()
            stored = process_group_upload(info)

            # Dedup: if same name + same chars already cached, discard the new duplicate
            duplicate = find_duplicate_by_name_and_chars(group_id, stored)
            if duplicate:
                remove_bridge_cache_item(group_id, stored)
                bridge_log("store", f"dedup skipped (same name+chars): group={group_id} file={stored.get('name', '')} chars={stored.get('text_chars', 0)}")
            else:
                with STATE_LOCK:
                    LAST_ERROR_BY_GROUP.pop(stored["group_id"], None)
                    current = UPLOAD_STATES.get(stored["group_id"])
                    if current and str(current.get("file_id", "")) == str(stored["file_id"]):
                        UPLOAD_STATES.pop(stored["group_id"], None)

                bridge_log("store", f"ok group={stored['group_id']} file={stored['name']} chars={stored['text_chars']}")

                # Auto-import extracted text into logutil if recording
                try:
                    auto_import_bridge_file_to_logutil(stored)
                except Exception as auto_exc:
                    bridge_log("autoimport", f"failed group={stored['group_id']} err={auto_exc}")
        except Exception as exc:
            err = str(exc)
            with STATE_LOCK:
                if group_id > 0:
                    LAST_ERROR_BY_GROUP[group_id] = err
                    current = UPLOAD_STATES.get(group_id)
                    if current and str(current.get("file_id", "")) == file_id:
                        current["status"] = "error"
                        current["processing"] = False
                        current["last_error"] = err

            bridge_log("error", f"group={group_id} {err}")
        finally:
            UPLOAD_QUEUE.task_done()


BRIDGE_POLL_WORKER: Optional[threading.Thread] = None


def get_effective_poll_interval(group_id):
    """Return the poll interval for a group.
    Priority: 1) logutil recording=1s  2) custom override  3) default."""
    try:
        state = ensure_logutil_group_state(str(group_id))
        if state.get('recording'):
            return LOGUTIL_POLL_INTERVAL_SEC
    except Exception:
        pass
    override = BRIDGE_POLL_INTERVAL_OVERRIDE.get(group_id)
    if override is not None and override > 0:
        return override
    return BRIDGE_POLL_INTERVAL_SEC


def bridge_poll_worker_loop():
    """Periodically poll NapCat for new files in tracked groups.
    Uses per-group timing: 1s granularity base, each group polled at its own interval."""
    bridge_log("poll", f"background poll worker started, default_interval={BRIDGE_POLL_INTERVAL_SEC}s, logutil_interval={LOGUTIL_POLL_INTERVAL_SEC}s")
    last_poll_time: Dict[int, float] = {}

    while True:
        time.sleep(1)  # base granularity for dynamic intervals

        try:
            with STATE_LOCK:
                groups = list(BRIDGE_POLL_GROUPS)
            if not groups:
                continue

            now = time.time()
            for group_id in groups:
                effective_interval = get_effective_poll_interval(group_id)
                last = last_poll_time.get(group_id, 0)
                if now - last < effective_interval:
                    continue  # not time for this group yet

                last_poll_time[group_id] = now
                try:
                    remote_info = get_latest_group_file_info(group_id)
                    if not remote_info:
                        continue
                    remote_file_id = str(remote_info.get('file_id', ''))
                    # Skip if this file_id already exists anywhere in the cache
                    with STATE_LOCK:
                        file_list = LATEST_FILES.get(group_id, [])
                    already_cached = any(
                        str(item.get('file_id', '')) == remote_file_id
                        for item in file_list
                    )
                    if already_cached:
                        continue  # silent skip — no need to re-download
                    remote_info['public_base'] = ''
                    pulled = process_group_upload(remote_info)
                    # Dedup: if same name + same chars already cached, discard the new duplicate
                    duplicate = find_duplicate_by_name_and_chars(group_id, pulled)
                    if duplicate:
                        remove_bridge_cache_item(group_id, pulled)
                        bridge_log("poll", f"dedup skipped (same name+chars): group={group_id} file={pulled.get('name', '')} chars={pulled.get('text_chars', 0)}")
                    else:
                        with STATE_LOCK:
                            LAST_ERROR_BY_GROUP.pop(group_id, None)
                        bridge_log("poll", f"new file detected group={group_id} file={pulled.get('name', '')} chars={pulled.get('text_chars', 0)} interval={effective_interval}s")
                        # Auto-import to logutil if recording
                        try:
                            auto_import_bridge_file_to_logutil(pulled)
                        except Exception:
                            pass
                except Exception as exc:
                    err_str = str(exc)
                    # Suppress "fileUUID not found" — file expired from NapCat, already cached
                    if "fileUUID" in err_str or "real fileUUID" in err_str:
                        bridge_log("poll", f"poll group={group_id} file expired (already cached)")
                    else:
                        bridge_log("poll", f"poll group={group_id} err={exc}")
        except Exception as exc:
            bridge_log("poll", f"poll loop err={exc}")


def ensure_poll_worker_started():
    # v4.3.4: WS模式(0)下不启动HTTP轮询，仅依赖WS实时推送
    if NC_FILE_BRIDGE_MODE == 0:
        return
    global BRIDGE_POLL_WORKER

    with STATE_LOCK:
        if BRIDGE_POLL_WORKER and BRIDGE_POLL_WORKER.is_alive():
            return

        BRIDGE_POLL_WORKER = threading.Thread(
            target=bridge_poll_worker_loop,
            name="bridge-poll-worker",
            daemon=True,
        )
        BRIDGE_POLL_WORKER.start()


def ensure_worker_started():
    global UPLOAD_WORKER

    with STATE_LOCK:
        if UPLOAD_WORKER and UPLOAD_WORKER.is_alive():
            return

        UPLOAD_WORKER = threading.Thread(
            target=upload_worker_loop,
            name="napcat-upload-worker",
            daemon=True,
        )
        UPLOAD_WORKER.start()


def enqueue_group_upload(info):
    ensure_worker_started()

    state = build_upload_state(info)
    group_id = safe_int(state.get("group_id", 0), 0)
    file_id = str(state.get("file_id", ""))

    with STATE_LOCK:
        LAST_ERROR_BY_GROUP.pop(group_id, None)
        UPLOAD_STATES[group_id] = state

    try:
        UPLOAD_QUEUE.put_nowait(dict(info))
    except Full:
        message = f"上传队列已满({BRIDGE_QUEUE_SIZE})，请稍后重试"
        with STATE_LOCK:
            LAST_ERROR_BY_GROUP[group_id] = message
            current = UPLOAD_STATES.get(group_id)
            if current and str(current.get("file_id", "")) == file_id:
                current["status"] = "error"
                current["processing"] = False
                current["last_error"] = message
        raise RuntimeError(message)

    queue_size = UPLOAD_QUEUE.qsize()
    with STATE_LOCK:
        current = UPLOAD_STATES.get(group_id)
        if current and str(current.get("file_id", "")) == file_id:
            current["queue_size"] = queue_size
            return dict(current)

    state["queue_size"] = queue_size
    return state


def extract_group_upload(payload):
    post_type = str(payload.get("post_type", "")).lower()
    notice_type = str(payload.get("notice_type", "")).lower()

    if post_type != "notice" or notice_type != "group_upload":
        if post_type == "message":
            group_id = payload.get("group_id")
            user_id = payload.get("user_id")
            message = payload.get("message")

            if isinstance(message, list):
                for seg in message:
                    if not isinstance(seg, dict):
                        continue
                    if str(seg.get("type", "")).lower() != "file":
                        continue
                    data = seg.get("data") or {}
                    file_id = data.get("id") or data.get("file_id")
                    file_name = data.get("name") or data.get("file") or "未知文件"
                    busid = safe_int(data.get("busid", 0), 0)
                    file_url = data.get("url") or ""
                    if group_id and file_id:
                        return {
                            "group_id": safe_int(group_id, 0),
                            "user_id": safe_int(user_id, 0),
                            "file_id": str(file_id),
                            "name": str(file_name),
                            "busid": busid,
                            "url": str(file_url),
                            "event_type": "message.file",
                        }
        return None

    group_id = payload.get("group_id")
    sender_id = payload.get("user_id")
    file_info = payload.get("file") or {}

    file_id = file_info.get("id") or file_info.get("file_id")
    file_name = file_info.get("name") or file_info.get("file_name") or "未知文件"
    busid = safe_int(file_info.get("busid", 0), 0)

    if not group_id or not file_id:
        return None

    return {
        "group_id": safe_int(group_id, 0),
        "user_id": safe_int(sender_id, 0),
        "file_id": str(file_id),
        "name": str(file_name),
        "busid": busid,
        "url": str(file_info.get("url") or ""),
        "source_ts": extract_source_ts(file_info) or extract_source_ts(payload),
        "event_type": "notice.group_upload",
    }


def napcat_json_post(path, body, timeout_sec):
    global LAST_NAPCAT_BASE, LAST_NAPCAT_ERROR

    last_err = None
    for base in NAPCAT_BASE_CANDIDATES:
        url = f"{base.rstrip('/')}/{path.lstrip('/')}"
        session = get_session()
        if NAPCAT_TOKEN:
            session.headers["Authorization"] = NAPCAT_TOKEN
        try:
            resp = session.post(url, json=body, timeout=timeout_sec)
            resp.raise_for_status()
            data = resp.json()
            LAST_NAPCAT_BASE = base
            LAST_NAPCAT_ERROR = ""
            bridge_log("napcat", f"ok base={base} path=/{path.lstrip('/')} body={sanitize_body_for_log(body)}")
            return data
        except Exception as exc:
            last_err = exc
            LAST_NAPCAT_ERROR = f"base={base} err={exc}"
            bridge_log("napcat", f"fail base={base} path=/{path.lstrip('/')} err={exc}")

    raise RuntimeError(f"所有NapCat地址均连接失败: {last_err}")


def get_group_file_url(group_id, file_id, busid):
    data = napcat_json_post(
        "/get_group_file_url",
        {
            "group_id": group_id,
            "file_id": file_id,
            "busid": busid,
        },
        timeout_sec=30,
    )

    file_url = ((data or {}).get("data") or {}).get("url", "")
    if not file_url:
        raise RuntimeError(f"get_group_file_url失败: {data}")
    return str(file_url)


def list_group_root_files(group_id):
    data = napcat_json_post("/get_group_root_files", {"group_id": group_id}, timeout_sec=30)
    files = ((data or {}).get("data") or {}).get("files") or []
    return files if isinstance(files, list) else []


def pick_latest_file(files):
    valid = [item for item in files if isinstance(item, dict) and (item.get("file_id") or item.get("id"))]
    if not valid:
        return None
    valid.sort(key=extract_source_ts, reverse=True)
    return valid[0]


def get_latest_group_file_info(group_id):
    files = list_group_root_files(group_id)
    bridge_log("pull", f"group={group_id} root_files={len(files)}")
    latest = pick_latest_file(files)
    if not latest:
        return None

    info = {
        "group_id": group_id,
        "user_id": safe_int(latest.get("uploader", 0), 0),
        "file_id": str(latest.get("file_id") or latest.get("id") or ""),
        "name": str(latest.get("file_name") or latest.get("name") or "未知文件"),
        "busid": safe_int(latest.get("busid", 0), 0),
        "url": str(latest.get("url") or ""),
        "source_ts": extract_source_ts(latest),
        "event_type": "pull.group_root_files",
    }
    if not info["file_id"]:
        return None

    bridge_log(
        "pull",
        f"latest file_id={info['file_id']} name={info['name']} busid={info['busid']} source_ts={info.get('source_ts', 0)}",
    )
    return info


def pull_latest_from_napcat(group_id, public_base=""):
    info = get_latest_group_file_info(group_id)
    if not info:
        return None
    info["public_base"] = public_base
    return process_group_upload(info)


def download_file_bytes(url):
    session = get_session()
    resp = session.get(url, timeout=DOWNLOAD_TIMEOUT_SEC, stream=True)
    resp.raise_for_status()

    total = 0
    chunks = []
    for buf in resp.iter_content(chunk_size=1024 * 256):
        if not buf:
            continue
        total += len(buf)
        if total > MAX_FILE_BYTES:
            raise RuntimeError(f"文件过大，超过 {MAX_FILE_MB}MB 上限")
        chunks.append(buf)
    return b"".join(chunks)


def extract_doc_legacy_text(data):
    ensure_bridge_cache_dir()

    for command_name in ("antiword", "catdoc"):
        command_path = shutil.which(command_name)
        if not command_path:
            continue

        temp_path = os.path.join(BRIDGE_CACHE_DIR, f"tmp_{uuid.uuid4().hex}.doc")
        try:
            with open(temp_path, "wb") as fw:
                fw.write(data)
            proc = subprocess.run(
                [command_path, temp_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=DOWNLOAD_TIMEOUT_SEC,
                check=False,
            )
            text = proc.stdout.decode("utf-8", errors="replace")
            if text.strip():
                return text
        except Exception:
            continue
        finally:
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception:
                pass

    raise RuntimeError("DOC解析失败：请安装 antiword 或 catdoc")


def extract_text_from_group_file(filename, data):
    ext = os.path.splitext((filename or "").lower())[1]

    if ext in (".txt", ".log", ".json", ".csv", ".md", ".xml", ".yaml", ".yml"):
        return safe_decode(data)

    if ext == ".docx":
        try:
            with zipfile.ZipFile(BytesIO(data)) as zf:
                targets = []
                for name in zf.namelist():
                    lowered = name.lower()
                    if lowered.startswith("word/") and any(x in lowered for x in ("document.xml", "header", "footer", "footnotes", "endnotes", "comments")):
                        targets.append(name)

                texts = []
                for path in targets:
                    raw = zf.read(path)
                    root = ET.fromstring(raw)
                    # v4.5.3: aggregate text runs per paragraph (not per node)
                    for p in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                        para_text = ''.join(
                            t.text or '' for t in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        ).strip()
                        if para_text:
                            texts.append(para_text)
                if texts:
                    return "\n".join(texts)
        except Exception:
            pass

    if ext in (".docx", ".pdf"):
        text = extract_text_from_file(data, filename)
        if text.startswith("[ParseError]"):
            raise RuntimeError(text.replace("[ParseError]", ""))
        return text

    if ext == ".doc":
        return extract_doc_legacy_text(data)

    return safe_decode(data)


def is_loopback_base(url):
    parsed = urllib.parse.urlparse(str(url or "").strip())
    host = str(parsed.hostname or "").strip().lower()
    return host in ("", "127.0.0.1", "localhost", "::1")


def resolve_public_base(preferred_base=""):
    configured = str(BRIDGE_PUBLIC_BASE or "").strip().rstrip("/")
    preferred = str(preferred_base or "").strip().rstrip("/")

    if configured and not is_loopback_base(configured):
        return configured
    if preferred:
        return preferred
    return configured or preferred


def build_content_url(content_key, public_base=""):
    base = resolve_public_base(public_base)
    content_url = f"{base}/bridge/content/{content_key}"
    if BRIDGE_TOKEN:
        content_url = f"{content_url}?token={urllib.parse.quote(BRIDGE_TOKEN, safe='')}"
    return content_url


def get_content_preview(content_key, chars=12):
    """Read first `chars` characters from a cached file for link preview."""
    if not content_key:
        return ''
    with STATE_LOCK:
        path = CONTENT_INDEX.get(content_key, '')
    if not path or not os.path.exists(path):
        return ''
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read(chars).replace('\n', ' ').replace('\r', '')
    except Exception:
        return ''


def hydrate_bridge_item(item, public_base=""):
    if not item:
        return None

    hydrated = dict(item)
    content_key = str(hydrated.get("content_key", "")).strip()
    if content_key:
        hydrated["content_url"] = build_content_url(content_key, public_base=public_base)
    return hydrated


# ====== v4.4.0: History System ======

def _evict_to_history(item, item_type='file'):
    """将即将被淘汰的桥接项移入 HISTORY 列表（而非直接删除）。
    item_type: 'file' | 'link'
    HISTORY 索引 0 = 最新。超过 MAX_HISTORY_ITEMS 时淘汰最旧的并删除磁盘文件。"""
    if not item:
        return
    entry = dict(item)
    entry['_type'] = item_type
    entry['_evicted_ts'] = now_ts()
    with STATE_LOCK:
        HISTORY.insert(0, entry)
        while len(HISTORY) > MAX_HISTORY_ITEMS:
            removed = HISTORY.pop()
            removed_key = str(removed.get("content_key", ""))
            removed_path = CONTENT_INDEX.pop(removed_key, "")
            if removed_path and os.path.exists(removed_path):
                try:
                    os.remove(removed_path)
                except Exception:
                    pass


def _serialize_history_item(item):
    """将 history item 转为可 JSON 序列化的字典。"""
    result = {}
    for k, v in item.items():
        if isinstance(v, (str, int, float, bool, type(None))):
            result[k] = v
        else:
            result[k] = str(v)
    return result


def save_history():
    """将 HISTORY 持久化到磁盘（HISTORY_FILE）。"""
    try:
        ensure_bridge_cache_dir()
        with STATE_LOCK:
            data = [_serialize_history_item(item) for item in HISTORY]
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        bridge_log("history", f"saved {len(data)} items to {HISTORY_FILE}")
    except Exception as e:
        print(f"[history] save failed: {e}")


def load_history():
    """从磁盘恢复 HISTORY 列表。会验证磁盘文件是否存在，缺失的跳过。"""
    global HISTORY
    if not os.path.exists(HISTORY_FILE):
        return
    try:
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        restored = []
        for entry in data:
            ck = str(entry.get('content_key', ''))
            # Reconstruct disk path from content_key
            path = os.path.join(BRIDGE_CACHE_DIR, f"{ck}.txt")
            if ck and os.path.exists(path):
                with STATE_LOCK:
                    CONTENT_INDEX[ck] = path
                entry['_type'] = entry.get('_type', 'file')
                restored.append(entry)
            else:
                print(f"[history] skip missing: ck={ck}")
        with STATE_LOCK:
            HISTORY = restored
        bridge_log("history", f"loaded {len(restored)} items from {HISTORY_FILE}")
        # Trim to MAX_HISTORY_ITEMS
        with STATE_LOCK:
            while len(HISTORY) > MAX_HISTORY_ITEMS:
                removed = HISTORY.pop()
                removed_key = str(removed.get("content_key", ""))
                removed_path = CONTENT_INDEX.pop(removed_key, "")
                if removed_path and os.path.exists(removed_path):
                    try:
                        os.remove(removed_path)
                    except Exception:
                        pass
    except Exception as e:
        print(f"[history] load failed: {e}")
        HISTORY = []


def _flush_all_to_history():
    """关闭前：将所有 LATEST_FILES 和 LINK_CACHE 的内容移入 HISTORY。"""
    with STATE_LOCK:
        for gid, file_list in list(LATEST_FILES.items()):
            for item in reversed(file_list):
                _evict_to_history(item, 'file')
            file_list.clear()
        for gid, link_list in list(LINK_CACHE.items()):
            for item in reversed(link_list):
                _evict_to_history(item, 'link')
            link_list.clear()
    save_history()


def shutdown_handler():
    """进程退出时的清理回调：持久化历史记录。"""
    print("[history] shutdown — flushing all caches to history...")
    try:
        _flush_all_to_history()
    except Exception as e:
        print(f"[history] shutdown error: {e}")


# ========================================


def write_text_cache(group_id, original_name, text, file_id, busid, user_id, source_ts=0, public_base=""):
    ensure_bridge_cache_dir()

    key = uuid.uuid4().hex
    base_name = os.path.splitext(os.path.basename(original_name or ""))[0] or "log_content"
    text_filename = f"{base_name}.txt"
    path = os.path.join(BRIDGE_CACHE_DIR, f"{key}.txt")
    cached_ts = now_ts()
    item_ts = source_ts if source_ts > 0 else cached_ts

    with open(path, "w", encoding="utf-8") as fw:
        fw.write(text)

    item = {
        "group_id": group_id,
        "file_id": file_id,
        "busid": busid,
        "name": original_name,
        "user_id": user_id,
        "ts": item_ts,
        "source_ts": item_ts,
        "cached_ts": cached_ts,
        "content_key": key,
        "content_url": build_content_url(key, public_base=public_base),
        "text_filename": text_filename,
        "text_chars": len(text),
        "text_bytes": os.path.getsize(path),
        "_type": "file",
    }

    with STATE_LOCK:
        if group_id not in LATEST_FILES:
            LATEST_FILES[group_id] = []
        file_list = LATEST_FILES[group_id]

        # Dedup: if same file_id already exists, remove old entry
        for i, old_item in enumerate(file_list):
            if str(old_item.get("file_id", "")) == str(file_id):
                old_key = str(old_item.get("content_key", ""))
                old_path = CONTENT_INDEX.pop(old_key, "")
                if old_path and os.path.exists(old_path):
                    try:
                        os.remove(old_path)
                    except Exception:
                        pass
                file_list.pop(i)
                break

        # Append new item (index -1 = latest, index 0 = oldest)
        file_list.append(item)
        CONTENT_INDEX[key] = path

        # Trim to max — v4.4.0: evict overflow to history instead of deleting
        while len(file_list) > MAX_BRIDGE_FILES_PER_GROUP:
            removed = file_list.pop(0)
            _evict_to_history(removed, 'file')

    return item


def write_link_cache(group_id, url, fetched_text):
    """v4.4.0: 将着色器链接的纯文本内容保存到桥接缓存。
    使用 LINK_CACHE 存储（与 LATEST_FILES 分离），编号规则 [link]-N。"""
    if not fetched_text or not str(fetched_text).strip():
        return None
    ensure_bridge_cache_dir()
    key = uuid.uuid4().hex
    path = os.path.join(BRIDGE_CACHE_DIR, f"{key}.txt")
    cached_ts = now_ts()

    with open(path, "w", encoding="utf-8") as fw:
        fw.write(str(fetched_text))

    # Truncate URL for display name
    url_display = str(url or "")[:80]
    item = {
        "group_id": group_id,
        "url": str(url or ""),
        "name": url_display,
        "ts": cached_ts,
        "source_ts": cached_ts,
        "cached_ts": cached_ts,
        "content_key": key,
        "content_url": build_content_url(key, public_base=resolve_public_base_or_fallback()),
        "text_chars": len(str(fetched_text)),
        "text_bytes": os.path.getsize(path),
        "_type": "link",
    }

    with STATE_LOCK:
        if group_id not in LINK_CACHE:
            LINK_CACHE[group_id] = []
        link_list = LINK_CACHE[group_id]
        # Dedup: skip if same URL already cached
        for i, old_item in enumerate(link_list):
            if str(old_item.get("url", "")) == str(url):
                old_key = str(old_item.get("content_key", ""))
                old_path = CONTENT_INDEX.pop(old_key, "")
                if old_path and os.path.exists(old_path):
                    try:
                        os.remove(old_path)
                    except Exception:
                        pass
                link_list.pop(i)
                break
        link_list.append(item)
        CONTENT_INDEX[key] = path
        # Trim to max — evict overflow to history
        while len(link_list) > MAX_BRIDGE_LINKS_PER_GROUP:
            removed = link_list.pop(0)
            _evict_to_history(removed, 'link')

    return item


def process_group_upload(info):
    group_id = safe_int(info.get("group_id", 0), 0)
    file_id = str(info.get("file_id", ""))
    busid = safe_int(info.get("busid", 0), 0)
    filename = str(info.get("name", "未知文件"))
    user_id = safe_int(info.get("user_id", 0), 0)
    source_ts = extract_source_ts(info)
    public_base = str(info.get("public_base") or "")

    if group_id <= 0 or not file_id:
        raise RuntimeError("无效的group_id/file_id")

    direct_url = str(info.get("url") or "").strip()
    download_url = direct_url if direct_url else get_group_file_url(group_id, file_id, busid)
    bridge_log("download", f"group={group_id} file={filename} file_id={file_id} busid={busid} url={download_url[:120]}")
    data = download_file_bytes(download_url)
    bridge_log("download", f"group={group_id} downloaded_bytes={len(data)}")

    text = extract_text_from_group_file(filename, data)
    if not text.strip():
        raise RuntimeError("文件中未提取到可用文本")

    bridge_log("extract", f"group={group_id} chars={len(text)} name={filename}")
    return write_text_cache(group_id, filename, text, file_id, busid, user_id, source_ts=source_ts, public_base=public_base)


def find_duplicate_by_name_and_chars(group_id, new_item):
    """Return the existing cached item if a duplicate (same name AND same text_chars) exists.
    Returns None if no duplicate is found."""
    with STATE_LOCK:
        file_list = list(LATEST_FILES.get(group_id, []))
    new_name = str(new_item.get("name", ""))
    new_chars = safe_int(new_item.get("text_chars", -1), -1)
    new_key = str(new_item.get("content_key", ""))
    if not new_name or new_chars < 0:
        return None
    for item in file_list:
        if str(item.get("content_key", "")) == new_key:
            continue  # skip self
        if str(item.get("name", "")) == new_name and safe_int(item.get("text_chars", 0), 0) == new_chars:
            return item
    return None


def remove_bridge_cache_item(group_id, item):
    """Remove a single cached item from LATEST_FILES and its on-disk cache file."""
    ck = str(item.get("content_key", ""))
    with STATE_LOCK:
        file_list = LATEST_FILES.get(group_id, [])
        if item in file_list:
            file_list.remove(item)
        old_path = CONTENT_INDEX.pop(ck, "")
    if old_path and os.path.exists(old_path):
        try:
            os.remove(old_path)
        except Exception:
            pass


def cleanup_expired():
    if BRIDGE_TTL_SEC <= 0:
        return

    now = now_ts()
    removed_paths = []

    with STATE_LOCK:
        expired_group_ids = []
        for gid, file_list in list(LATEST_FILES.items()):
            kept = []
            for item in file_list:
                ts = extract_cached_ts(item)
                if ts > 0 and (now - ts) <= BRIDGE_TTL_SEC:
                    kept.append(item)
                else:
                    # v4.4.0: evict expired items to history instead of deleting
                    _evict_to_history(item, 'file')
            if kept:
                LATEST_FILES[gid] = kept
            else:
                expired_group_ids.append(gid)

        for gid in expired_group_ids:
            LATEST_FILES.pop(gid, None)

        expired_states = []
        for gid, item in list(UPLOAD_STATES.items()):
            ts = extract_cached_ts(item)
            if ts <= 0 or (now - ts) > BRIDGE_TTL_SEC:
                expired_states.append(gid)

        for gid in expired_states:
            UPLOAD_STATES.pop(gid, None)

        for gid in set(expired_group_ids + expired_states):
            if gid not in LATEST_FILES and gid not in UPLOAD_STATES:
                LAST_ERROR_BY_GROUP.pop(gid, None)

    for path in removed_paths:
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass

def background_file_process(job_id, file_url, filename, mode='analyze', is_pro=False, is_kind=False, persona="", custom_prompt="", theme='default', get_text=False, group_id=0):
    """后台任务：下载文件并根据模式进行分析，支持多模态原生文档阅读与输出多图"""
    print(f"[{job_id}] 开始处理文件: {filename}, Mode: {mode}")
    try:
        # 1. 下载文件（支持单文件与多文件）
        if isinstance(file_url, list):
            urls = [str(u or '').strip() for u in file_url if str(u or '').strip()]
        else:
            urls = [str(file_url or '').strip()]
        if isinstance(filename, list):
            names = [str(n or '').strip() for n in filename]
        else:
            names = [str(filename or '').strip()]

        session = get_session()
        is_multi_file = len(urls) > 1
        resolved_filename = names[0] if names else str(filename or 'unknown')

        content = b""
        downloaded = 0
        user_content = None
        text_ai = ""

        if is_multi_file:
            merged_parts = []
            content_hash_chunks = []
            for idx, one_url in enumerate(urls):
                one_name = names[idx] if idx < len(names) and names[idx] else f"file_{idx+1}.txt"
                resp = session.get(one_url, timeout=120, stream=True)
                resp.raise_for_status()

                one_content = b""
                one_downloaded = 0
                for chunk in resp.iter_content(chunk_size=65536):
                    if chunk:
                        one_content += chunk
                        one_downloaded += len(chunk)
                        downloaded += len(chunk)
                        if downloaded > 150 * 1024 * 1024:
                            print(f"[{job_id}] 警告：总文件体积超过 50MB，已被安全截断！")
                            break

                content_hash_chunks.append(one_content)
                raw_text = extract_text_from_file(one_content, one_name)
                if raw_text.startswith("[ParseError]"):
                    raise Exception(raw_text.replace("[ParseError]", ""))
                if not raw_text or len(raw_text.strip()) < 1:
                    continue
                merged_parts.append(f"【第{idx+1}段文件：{one_name}】\n{raw_text}")

            if not merged_parts:
                raise Exception("多文件提取后内容为空")

            raw_text = "\n\n===== 多文件拼接分隔线 =====\n\n".join(merged_parts)
            text_ai = raw_text
            if len(raw_text) > MAX_AI_CHARS:
                part = int(MAX_AI_CHARS * 0.4)
                mid = raw_text[part:-part].split('\n')
                step = max(1, int(len(mid)/100))
                text_ai = f"{raw_text[:part]}\n...[中间部分略]...\n{chr(10).join(mid[::step])}\n{raw_text[-part:]}"

            user_content = f"文件名：多文件拼接({len(merged_parts)}段)\n内容如下：\n{text_ai}"
            content = b"".join(content_hash_chunks)
            resolved_filename = "multi_files.txt"
        else:
            resp = session.get(urls[0], timeout=120, stream=True)
            resp.raise_for_status()
            for chunk in resp.iter_content(chunk_size=65536):
                if chunk:
                    content += chunk
                    downloaded += len(chunk)
                    # 限制最大下载 50MB，防止内存爆炸
                    if downloaded > 50 * 1024 * 1024:
                        print(f"[{job_id}] 警告：文件超过 50MB，已被安全截断！")
                        break
        
        # 2. 核心：判断是否启用 LLM 的原生多模态视觉/文档阅读能力
        ext = os.path.splitext(resolved_filename)[1].lower()

        if (not is_multi_file) and ext == '.pdf' and downloaded <= 150 * 1024 * 1024:
            # 【原生 PDF 阅读模式】(限制在40MB内防代理服务器 Nginx 报 413 Payload Too Large)
            print(f"[{job_id}] 启用 LLM 原生 PDF 阅读模式 (大小: {downloaded/1024/1024:.2f}MB)")
            base64_pdf = base64.b64encode(content).decode('utf-8')
            user_content = [
                {"type": "text", "text": f"文件名：{resolved_filename}\n请仔细阅读这份 PDF 模组文档（包含其排版和图像），并严格按照系统设定的板块与要求进行分析。"},
                {"type": "image_url", "image_url": {"url": f"data:application/pdf;base64,{base64_pdf}"}}
            ]
            
        elif (not is_multi_file) and ext in ['.png', '.jpg', '.jpeg', '.webp'] and downloaded <= 20 * 1024 * 1024:
            # 【原生图片阅读模式】
            print(f"[{job_id}] 启用 LLM 原生图片阅读模式")
            mime_type = "image/jpeg" if ext in ['.jpg', '.jpeg'] else f"image/{ext[1:]}"
            base64_img = base64.b64encode(content).decode('utf-8')
            user_content = [
                {"type": "text", "text": f"文件名：{resolved_filename}\n请仔细观察这张图片/设定图，并严格按照系统设定的板块与要求进行分析。"},
                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_img}"}}
            ]
            
        else:
            # 【文本提取回退模式】(非视觉格式，或文件超大)
            print(f"[{job_id}] 启用文本本地提取模式")
            raw_text = extract_text_from_file(content, resolved_filename)
            
            if raw_text.startswith("[ParseError]"):
                raise Exception(raw_text.replace("[ParseError]", ""))
                
            if not raw_text or len(raw_text.strip()) < 10:
                raise Exception("文件内容为空或提取不到文字。(如果模组全是扫描版图片且文件过大，AI暂无法阅读)")

            # 智能压缩文本防爆 Token
            if len(raw_text) > MAX_AI_CHARS:
                part = int(MAX_AI_CHARS * 0.4)
                mid = raw_text[part:-part].split('\n')
                step = max(1, int(len(mid)/100))
                text_ai = f"{raw_text[:part]}\n...[中间部分略]...\n{chr(10).join(mid[::step])}\n{raw_text[-part:]}"
            else:
                text_ai = raw_text

            user_content = f"文件名：{resolved_filename}\n内容如下：\n{text_ai}"

        # ================= 1. 尝试触发文件省流缓存 =================
        # 注意：由于群文件链接 file_url 经常变，我们只能通过哈希“文件的真实数据内容”来确认是不是同一个文件
        hash_key = None
        if not is_pro:
            if isinstance(user_content, list): 
                content_hash = hashlib.md5(content).hexdigest()
            else: 
                content_hash = hashlib.md5(text_ai.encode('utf-8')).hexdigest()
                
            hash_str = f"file_log_{content_hash}_{mode}_{is_kind}_{persona}_{custom_prompt}_{theme}"
            hash_key = hashlib.md5(hash_str.encode('utf-8')).hexdigest()
            
            cached_images = get_daily_cache(hash_key)
            if cached_images:
                print(f"[{job_id}] 命中今日文件内容缓存！省流模式启动，秒回历史图片。")
                JOB_CACHE[job_id]['status'] = 'done'
                JOB_CACHE[job_id]['images'] = cached_images
                return
        # ==========================================================

        # 3. 根据不同模式分配 Prompt 与 绘图标题
        report_title = "TRPG 模组解析报告"
        
        # 【新增】：检测并覆盖
        if custom_prompt:
            report_title = "TRPG 自定义分析报告"
            system_prompt = custom_prompt
        
        elif mode == 'log_analyze':
            report_title = "TRPG 跑团日志评分"
            if is_kind: system_prompt = KIND_SYSTEM_PROMPT
            elif is_pro: system_prompt = PRO_SYSTEM_PROMPT
            else: system_prompt = DEFAULT_SYSTEM_PROMPT
        
        elif mode == 'log_recap':
            report_title = "TRPG 跑团前文回顾"
            system_prompt = """你是一个专业且细致的 TRPG 跑团记录员（书记）。请阅读以下跑团 Log，为 KP 和玩家梳理一份详细的【前文回顾】，帮助大家快速找回跑团记忆。
请严格按照以下 4 个板块输出，并且在输出每个大板块之前，必须使用“【分页符】”这四个字单起一行作为分隔标识（不要用Markdown，不要加粗）：

【分页符】
【一、当前剧情进度总览】：
（详细说明截至目前的故事进度，大家在哪，正在面临什么状况，遇到了什么危机或主线推进到了哪一步）
【分页符】
【二、PC行动轨迹与状态梳理】：
（尽可能详细分条列出每位主要玩家角色/PC近期做了什么举动，达成了什么目的，或处于什么特殊状态/受到什么伤害）
【分页符】
【三、当前已获线索与道具盘点】：
（总结当前大家掌握的所有情报、未解之谜、NPC给出的重要信息以及拿到的关键道具）
【分页符】
【四、下一步推进方向提示】：
（基于当前局势，客观给出2-3个可供调查员们继续推进剧情的可能方向或需要立刻解决的问题）"""

        elif mode == 'prepare':
            report_title = "TRPG 备团资料梳理"
            system_prompt = """你是一个资深的TRPG跑团KP/DM，请阅读以下模组文档内容，为带团准备一份详尽的备团参考。
请严格提供以下5个板块的内容，并且在输出每个大板块之前，必须使用“【分页符】”这四个字单起一行作为分隔标识（不要用Markdown，不要加粗）：

【分页符】
【一、模组背景】：
（阐述事件真相、幕后黑手动机、历史遗留问题，让KP掌握全局）
【分页符】
【二、故事梗概】：
（按时间线或事件发展顺序，简述调查员将经历的主要剧情节点）
【分页符】
【三、人物关系】：
（列出核心NPC的表面身份、真实身份、动机及相互关系）
【分页符】
【四、地图与场景梳理】：
（罗列关键场景及可获取的线索或触发的事件）
【分页符】
【五、建议流程】：
（带团节奏建议，指出哪里需重点渲染，哪里容易卡关需暗中提示）"""

        elif mode == 'refine':
            report_title = "TRPG 模组润色与审查"
            system_prompt = """你是一个资深的TRPG剧本医生/编辑，请阅读以下模组，评估目前的写作状态并给出修改建议。
请严格提供以下5个板块的内容，并且在输出每个大板块之前，必须使用“【分页符】”这四个字单起一行作为分隔标识（不要用Markdown，不要加粗）：

【分页符】
【一、完成度预估】：
（评估书写进度百分比，例如：完成度60%，并简述理由）
【分页符】
【二、当前进度点评】：
（客观评价已写好的部分，指出亮点与明显缺失的核心要素）
【分页符】
【三、写作建议】：
（针对薄弱环节，提供具体的构思方向或剧情补充建议）
【分页符】
【四、需要调整的地方】：
（指出逻辑漏洞、规则应用错误、或排版行文生硬之处）
【分页符】
【五、具体修改示例】：
（选取文中某段落或缺失的设定，给出一个经你润色补充的具体文本范例）"""

        else: # 默认 analyze
            system_prompt = """
        你是一个专业的TRPG模组锐评大师，请阅读以下COC/DND模组文档，生成一份评测简报。要求语言风格严谨犀利，不轻易给高分。
        请严格按照以下格式输出（不要用Markdown，不要加粗，直接分行）：
        【文件标题】：(文件名)
        【模组/文档类型】：（模组类型）
        【总体评分】：（0-100分）
        （请给出理由）
        【核心内容概要】：
        （简述模组里面的内容以及逻辑情况，分析调查员可能的行动方向以及对应结果）
        【亮点/特色】：
        （找出几个模组中描写或逻辑最佳的地方，如果没有则忽略不写）
        【问题/槽点】：
        （找出并吐槽模组中的逻辑漏洞，文笔硬伤，忽视规则书等等行为）
        【专家总结】：
        （用一到两句话来总结模组）
        """

        # 核心：人设系统劫持
        if persona:
            system_prompt += f"\n\n【极其重要的扮演指令】：\n在生成上述所有评价和梳理内容时，请你完全带入以下角色人设来进行语气和口吻的渲染。你可以自称、吐槽或撒娇，让输出充满该人设的个性。\n（绝对警告：你必须严格保留前文要求的【板块标题】和【分页符】等格式标识符，千万不能省略或修改它们，只能改变正文部分的说话语气！）：\n{persona}"
        
        # 【新增】：如果用户没指定主题，赋予大模型绝对的主题控制权！
        if theme == 'default':
            system_prompt += "\n\n【排版指令】：你可以根据当前内容的故事氛围，在回复的【最开头】加上标签以控制最终生成的图片风格。支持的标签有：【主题：经典】、【主题：克苏鲁】、【主题：赛博】、【主题：历史】、【主题：废土】、【主题：二次元】、【主题：终端】。如果你觉得不需要特殊风格，可不写此标签。"
        
        # 4. 请求 AI
        # 3. 请求 AI
        print(f"[{job_id}] 未命中缓存，开始请求 LLM 消耗 Token...")
        resp = get_openai_client().chat.completions.create(
            model = AI_MODEL_PRO if is_pro else AI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt}, 
                {"role": "user", "content": user_content}
            ],
            temperature=1.0, max_tokens=65535
        )
        result_text = resp.choices[0].message.content or ""
        token_usage = get_token_usage_suffix(resp)
        
        result_text, final_theme = extract_theme_from_text(result_text, theme)

        # 4. 多图渲染与保存
        images_list = text_to_images(result_text, resolved_filename, report_title, final_theme, token_usage)
        JOB_CACHE[job_id]['status'] = 'done'
        JOB_CACHE[job_id]['images'] = images_list
        JOB_CACHE[job_id]['text'] = result_text  # 保存原始文本供 get_text 使用
        print(f"[{job_id}] 文件分析完成，共生成 {len(images_list)} 张图")

        # 5b. 如果 get_text 模式，将文本保存为桥接文件
        if get_text:
            try:
                ensure_bridge_cache_dir()
                text_key = uuid.uuid4().hex
                text_path = os.path.join(BRIDGE_CACHE_DIR, f"{text_key}.txt")
                safe_filename = re.sub(r'[\\/*?:"<>|]', '_', f"ai_analysis_{job_id[:8]}.txt")
                with open(text_path, 'w', encoding='utf-8') as fw:
                    fw.write(result_text)
                with STATE_LOCK:
                    CONTENT_INDEX[text_key] = text_path
                public_base = resolve_public_base_or_fallback()
                text_url = build_content_url(text_key, public_base=public_base)
                JOB_CACHE[job_id]['text_url'] = text_url
                JOB_CACHE[job_id]['text_key'] = text_key
                JOB_CACHE[job_id]['text_filename'] = safe_filename
                print(f"[{job_id}] get_text 文件已保存: {text_url}")
                # 自动上传到群
                if group_id > 0:
                    try:
                        fs, _ = napcat_upload_group_file(group_id, text_path, safe_filename)
                        JOB_CACHE[job_id]['text_file_sent'] = fs
                        if fs:
                            print(f"[{job_id}] get_text 文件已上传到群 {group_id}")
                    except Exception as ue:
                        print(f"[{job_id}] get_text 上传失败: {ue}")
            except Exception as e:
                print(f"[{job_id}] get_text 保存失败: {e}")

        # ================= 5. 写入省流缓存 =================
        if not is_pro and hash_key:
            set_daily_cache(hash_key, images_list)
            print(f"[{job_id}] 文件处理结果已存入今日缓存库。")

    except Exception as e:
        print(f"[{job_id}] 文件处理失败: {e}")
        err_img_bytes = text_to_images(f"文件处理失败：\n{str(e)}", filename)[0]
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['images'] =[err_img_bytes]
        JOB_CACHE[job_id]['text'] = f"处理失败：{str(e)}"

TRANSLATE_SYSTEM_PROMPT = "你是一个专业的翻译助手。请准确翻译用户提供的文本，保留原文格式，只返回翻译结果，不要添加任何解释或评论。"

# v4.4.1: TextDB.online 云数据库 (用于 goal-ALL 翻译)
# API: GET/POST https://textdb.online/update/?key=...&value=...
# 注意: POST body (data=) 不被接受，必须使用 URL query params 发送
# 读取: GET https://textdb.online/{key} (始终返回200，不存在则内容为空)
# 成功判定: 响应 JSON 中 status==1 (不能依赖 HTTP 状态码)
# URL 长度限制: 约12KB (Cloudflare 414)，value 需截断
TEXTDB_UPDATE_URL = "https://textdb.online/update/"
TEXTDB_MAX_VALUE_CHARS = 2500  # URL-safe 上限 (~22KB URL)，实测 414 出现在 ~3000 chars (~27KB)


def textdb_upload(key, value):
    """上传文本到 TextDB.online。使用 GET + query params 发送。
    值超过 TEXTDB_MAX_VALUE_CHARS 时自动截断（保留头部+尾部标记）。"""
    try:
        v = str(value)
        if len(v) > TEXTDB_MAX_VALUE_CHARS:
            head = v[:TEXTDB_MAX_VALUE_CHARS // 2]
            tail = v[-(TEXTDB_MAX_VALUE_CHARS // 2 - 20):]
            v = f"{head}\n\n... [中间省略 {len(value) - TEXTDB_MAX_VALUE_CHARS} 字符] ...\n\n{tail}"
        # 使用 POST + query params（文档建议 POST 方法；POST body 不可用）
        resp = requests.post(TEXTDB_UPDATE_URL, params={
            'key': str(key),
            'value': v
        }, timeout=30)
        # TextDB 永远返回 200 (成功) 或 414 (URL过长)；检测 JSON 中 status==1
        data = {}
        if resp.text and resp.text.strip():
            try:
                data = resp.json()
            except Exception:
                data = {}
        ok = isinstance(data, dict) and data.get('status') == 1
        if not ok:
            print(f"[textdb] write failed: http={resp.status_code} status={data.get('status')} key={key[:20]}")
        return ok
    except Exception as e:
        print(f"[textdb] upload failed: {e}")
        return False


def textdb_get_url(key):
    """返回 TextDB 在线查看链接：https://textdb.online/{key}"""
    return f"https://textdb.online/{key}"


def textdb_read(key):
    """从 TextDB.online 读取翻译进度（用于验证）。始终返回200，不存在则内容为空。"""
    try:
        resp = requests.get(f"https://textdb.online/{key}", timeout=30)
        return resp.text if resp.status_code == 200 else ''
    except Exception as e:
        print(f"[textdb] read failed: {e}")
        return ''


def chunk_text_by_sentences(text, target_chars=2000):
    """将文本按句子边界切分为 ~target_chars 大小的块。"""
    chunks = []
    current = ""
    # 按句子边界分割
    sentences = re.split(r'(?<=[。！？\n])(?=[^。！？\n])', text)
    for sent in sentences:
        if len(current) + len(sent) > target_chars and current:
            chunks.append(current.strip())
            current = sent
        else:
            current += sent
    if current.strip():
        chunks.append(current.strip())
    return chunks if chunks else [text]


@app.route('/api/translate', methods=['GET'])
def translate_task():
    """翻译文件任务 (v4.4.0: 支持 goal=all 模式)"""
    if len(JOB_CACHE) > 100: JOB_CACHE.clear()

    file_url = request.args.get('url')
    filename = request.args.get('filename', 'unknown')
    target_lang = request.args.get('lang', 'zh-CN')
    is_pro = request.args.get('pro', 'false').lower() == 'true'
    group_id = safe_int(request.args.get('group_id', 0), 0)
    goal_all = request.args.get('goal', '').lower() == 'all'

    if not file_url:
        return jsonify({'status': 'error', 'msg': '缺少文件URL'})

    job_id = str(uuid.uuid4())
    CANCEL_FLAGS[job_id] = threading.Event()
    JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time(), 'group_id': group_id}

    if goal_all:
        textdb_key = f"logai-trans-{uuid.uuid4().hex[:12]}"
        textdb_url = textdb_get_url(textdb_key)
        JOB_CACHE[job_id]['mode'] = 'goal-all'
        JOB_CACHE[job_id]['textdb_key'] = textdb_key
        JOB_CACHE[job_id]['textdb_url'] = textdb_url
        executor.submit(background_translate_goal_all, job_id, file_url, filename, target_lang, group_id, textdb_key)
        return jsonify({'status': 'ok', 'id': job_id, 'mode': 'goal-all',
                        'textdb_key': textdb_key, 'textdb_url': textdb_url,
                        'msg': f'开始 goal-ALL 翻译为 {target_lang}...\n在线查看进度: {textdb_url}'})
    else:
        executor.submit(background_translate_process, job_id, file_url, filename, target_lang, is_pro, group_id)
        return jsonify({'status': 'ok', 'id': job_id, 'msg': f'正在翻译为 {target_lang}...'})

def background_translate_process(job_id, file_url, filename, target_lang='zh-CN', is_pro=False, group_id=0):
    """后台线程：下载并翻译文件"""
    print(f"[{job_id}] 开始翻译文件: {filename} -> {target_lang}")
    try:
        sess = get_session()
        resp = sess.get(file_url, timeout=60)
        if resp.status_code != 200:
            raise Exception(f"文件下载失败: {resp.status_code}")
        
        file_content = safe_decode(resp.content)
        
        # 检测语言
        lang_hint = ""
        if any('\u4e00' <= c <= '\u9fff' for c in file_content[:500]):
            lang_hint = "原文是中文"
        elif any('\u3040' <= c <= '\u309f' or '\u30a0' <= c <= '\u30ff' for c in file_content[:500]):
            lang_hint = "原文是日文"
        elif any('\uac00' <= c <= '\ud7af' for c in file_content[:500]):
            lang_hint = "原文是韩文"
        
        model = AI_MODEL_PRO if is_pro else AI_MODEL
        
        # 截断过长的内容
        file_text = file_content[:MAX_AI_CHARS] if len(file_content) > MAX_AI_CHARS else file_content
        
        translate_prompt = f"{lang_hint}\n请将以下文本翻译成{target_lang}：\n\n{file_text}"
        
        resp = get_openai_client().chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": TRANSLATE_SYSTEM_PROMPT},
                {"role": "user", "content": translate_prompt}
            ],
            temperature=0.5, max_tokens=4000
        )
        result_text = resp.choices[0].message.content
        
        # 保存翻译结果到桥接缓存（用于上传到群）
        ensure_bridge_cache_dir()
        trans_key = uuid.uuid4().hex
        trans_path = os.path.join(BRIDGE_CACHE_DIR, f"{trans_key}.txt")
        base_name = os.path.splitext(str(filename))[0]
        safe_filename = re.sub(r'[\\/*?:"<>|]', '_', f"翻译_{target_lang}_{base_name}.txt")
        with open(trans_path, 'w', encoding='utf-8') as fw:
            fw.write(result_text or '')
        with STATE_LOCK:
            CONTENT_INDEX[trans_key] = trans_path

        JOB_CACHE[job_id]['status'] = 'done'
        JOB_CACHE[job_id]['text'] = result_text
        JOB_CACHE[job_id]['text_key'] = trans_key
        JOB_CACHE[job_id]['text_filename'] = safe_filename
        JOB_CACHE[job_id]['original_filename'] = filename
        print(f"[{job_id}] 文件翻译完成: {filename}")
        # 自动上传到群
        if group_id > 0:
            try:
                fs, _ = napcat_upload_group_file(group_id, trans_path, safe_filename)
                JOB_CACHE[job_id]['text_file_sent'] = fs
                if fs:
                    print(f"[{job_id}] 翻译文件已上传到群 {group_id}")
            except Exception as ue:
                print(f"[{job_id}] 翻译上传失败: {ue}")

    except Exception as e:
        print(f"[{job_id}] 文件翻译失败: {e}")
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['text'] = f"翻译失败：{str(e)}"


# v4.4.0: goal-ALL 分块翻译
def background_translate_goal_all(job_id, file_url, filename, target_lang, group_id, textdb_key=None):
    """后台线程：分块翻译文件，每10秒上传到 TextDB.online。"""
    print(f"[{job_id}] 开始 goal-ALL 翻译: {filename} -> {target_lang}")
    cancel_event = CANCEL_FLAGS.get(job_id)
    try:
        sess = get_session()
        resp = sess.get(file_url, timeout=120)
        if resp.status_code != 200:
            raise Exception(f"文件下载失败: {resp.status_code}")

        source_text = safe_decode(resp.content)
        if len(source_text) > MAX_AI_CHARS:
            source_text = source_text[:MAX_AI_CHARS]

        chunks = chunk_text_by_sentences(source_text, target_chars=2000)
        if not textdb_key:
            textdb_key = f"logai-trans-{uuid.uuid4().hex[:12]}"
        textdb_url = textdb_get_url(textdb_key)

        JOB_CACHE[job_id]['mode'] = 'goal-all'
        JOB_CACHE[job_id]['textdb_key'] = textdb_key
        JOB_CACHE[job_id]['textdb_url'] = textdb_url
        JOB_CACHE[job_id]['total_chunks'] = len(chunks)
        JOB_CACHE[job_id]['completed_chunks'] = 0

        accumulated = ""
        last_upload_ts = time.time()

        model = AI_MODEL
        for i, chunk in enumerate(chunks):
            if cancel_event and cancel_event.is_set():
                JOB_CACHE[job_id]['status'] = 'cancelled'
                JOB_CACHE[job_id]['text'] = f'翻译已被停止。已完成 {i}/{len(chunks)} 段。'
                return

            resp = get_openai_client().chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": TRANSLATE_SYSTEM_PROMPT},
                    {"role": "user", "content": f"请将以下文本翻译成{target_lang}：\n\n{chunk}"}
                ],
                temperature=0.5, max_tokens=4000
            )
            translated = resp.choices[0].message.content or ""
            accumulated += translated + "\n\n"

            JOB_CACHE[job_id]['completed_chunks'] = i + 1
            JOB_CACHE[job_id]['accumulated_chars'] = len(accumulated)

            # Upload to TextDB every 10 seconds or on last chunk
            now = time.time()
            if now - last_upload_ts >= 10 or i == len(chunks) - 1:
                textdb_upload(textdb_key, accumulated)
                last_upload_ts = now
                print(f"[{job_id}] goal-ALL progress: {i+1}/{len(chunks)}, uploaded to TextDB")

        # Final upload
        textdb_upload(textdb_key, accumulated)

        # Also save locally in bridge cache
        ensure_bridge_cache_dir()
        final_key = uuid.uuid4().hex
        final_path = os.path.join(BRIDGE_CACHE_DIR, f"{final_key}.txt")
        with open(final_path, 'w', encoding='utf-8') as fw:
            fw.write(accumulated)
        with STATE_LOCK:
            CONTENT_INDEX[final_key] = final_path

        JOB_CACHE[job_id]['status'] = 'done'
        JOB_CACHE[job_id]['text'] = accumulated
        JOB_CACHE[job_id]['text_key'] = final_key
        JOB_CACHE[job_id]['textdb_url'] = textdb_url
        print(f"[{job_id}] goal-ALL 翻译完成: {textdb_url}")

        # Auto-upload to group
        if group_id > 0:
            try:
                safe_name = re.sub(r'[\\/*?:"<>|]', '_', f"翻译_{target_lang}_{filename}.txt")
                fs, _ = napcat_upload_group_file(group_id, final_path, safe_name)
                JOB_CACHE[job_id]['text_file_sent'] = fs
            except Exception as ue:
                print(f"[{job_id}] goal-ALL 上传失败: {ue}")

    except Exception as e:
        print(f"[{job_id}] goal-ALL 翻译失败: {e}")
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['text'] = f"goal-ALL翻译失败：{str(e)}"


@app.route('/api/translate_result', methods=['GET'])
def get_translate_result():
    """获取翻译结果"""
    job_id = request.args.get('id')
    job = JOB_CACHE.get(job_id)
    if not job or 'text' not in job:
        return jsonify({'status': 'not_found'})
    resp = {'status': job['status'], 'text': job.get('text', ''), 'filename': job.get('original_filename', ''), 'text_key': job.get('text_key', ''), 'text_filename': job.get('text_filename', '')}
    # v4.4.0: goal-ALL fields
    if job.get('mode') == 'goal-all':
        resp['mode'] = 'goal-all'
        resp['textdb_key'] = job.get('textdb_key', '')
        resp['textdb_url'] = job.get('textdb_url', '')
        resp['total_chunks'] = job.get('total_chunks', 0)
        resp['completed_chunks'] = job.get('completed_chunks', 0)
        resp['accumulated_chars'] = job.get('accumulated_chars', 0)
    return jsonify(resp)

@app.route('/api/translate_and_upload', methods=['GET'])
def translate_and_upload():
    """翻译并上传到群文件"""
    if len(JOB_CACHE) > 100: JOB_CACHE.clear()
    
    file_url = request.args.get('url')
    filename = request.args.get('filename', 'unknown')
    target_lang = request.args.get('lang', 'zh-CN')
    group_id = request.args.get('group_id', '')
    is_pro = request.args.get('pro', 'false').lower() == 'true'
    overwrite = request.args.get('overwrite', 'false').lower() == 'true'
    upload_baseurl = request.args.get('upload_url', '')
    
    if not file_url:
        return jsonify({'status': 'error', 'msg': '缺少文件URL'})
    
    if not group_id or not upload_baseurl:
        return jsonify({'status': 'error', 'msg': '缺少群号或上传地址'})
    
    job_id = str(uuid.uuid4())
    JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time()}
    
    executor.submit(background_translate_and_upload, job_id, file_url, filename, target_lang, group_id, upload_baseurl, is_pro, overwrite)

    mode_msg = "覆盖模式" if overwrite else "注释模式"
    return jsonify({'status': 'ok', 'id': job_id, 'msg': f'正在翻译并上传到群文件...({mode_msg})'})


def translate_and_save_file(job_id, file_bytes, original_filename, target_lang, is_pro=False, overwrite=False):
    """根据文件类型翻译并保存"""
    import os
    
    sess = get_session()
    ext = os.path.splitext(original_filename)[1].lower()
    name_without_ext = os.path.splitext(original_filename)[0]

    mode_desc = "覆盖模式" if overwrite else "注释模式"
    print(f"[{job_id}] 翻译模式: {mode_desc}")
    
    # 检测语言
    check_text = ""
    if ext == '.pdf':
        try:
            from io import BytesIO
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            check_text = "\n".join(text_parts)
        except:
            check_text = safe_decode(file_bytes)[:1000]
    elif ext == '.docx':
        try:
            from io import BytesIO
            import docx
            doc = docx.Document(BytesIO(file_bytes))
            check_text = "\n".join([p.text for p in doc.paragraphs])
        except:
            check_text = safe_decode(file_bytes)[:1000]
    else:
        check_text = safe_decode(file_bytes)[:1000]
    
    lang_hint = ""
    if any('\u4e00' <= c <= '\u9fff' for c in check_text):
        lang_hint = "原文是中文"
    elif any('\u3040' <= c <= '\u309f' or '\u30a0' <= c <= '\u30ff' for c in check_text):
        lang_hint = "原文是日文"
    elif any('\uac00' <= c <= '\ud7af' for c in check_text):
        lang_hint = "原文是韩文"
    
    model = AI_MODEL_PRO if is_pro else AI_MODEL
    
    # 根据不同文件类型处理
    if ext == '.pdf':
        # PDF处理 - 使用pymupdf保留原有格式
        try:
            from io import BytesIO
            import fitz  # pymupdf
            
            print(f"[{job_id}] 使用pymupdf处理PDF，保留原有格式")
            
            # 打开原始PDF
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            
            # 提取所有文本用于翻译
            all_texts = []
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text_dict = page.get_text("dict")
                for block in get_pymupdf_blocks(text_dict):
                    if block.get("type") == 0:  # 文本块
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                text = span.get("text", "").strip()
                                if text and len(text) > 2:
                                    all_texts.append(text)
            
            if not all_texts:
                raise Exception("PDF中没有可提取的文本")
            
            print(f"[{job_id}] 提取到 {len(all_texts)} 个文本片段")
            
            # 翻译所有文本
            translated_map = {}
            batch_size = 30
            for i in range(0, len(all_texts), batch_size):
                batch = all_texts[i:i+batch_size]
                batch_text = "\n---\n".join(batch)
                prompt = f"{lang_hint}\n请将以下文本翻译成{target_lang}，保持每行对应（用---分隔）：\n{batch_text}"
                
                resp = get_openai_client().chat.completions.create(
                    model=model,
                    messages=[{"role": "system", "content": TRANSLATE_SYSTEM_PROMPT}, {"role": "user", "content": prompt}],
                    temperature=0.5, max_tokens=4000
                )

                translated_text = resp.choices[0].message.content or ""
                translated_batch = translated_text.split('\n---\n')
                for orig, trans in zip(batch, translated_batch):
                    translated_map[orig] = trans.strip()
            
            # 在原始PDF中替换文本
            # 尝试加载中文字体 - 使用pymupdf的方式
            chinese_font_path = None
            font_paths = [
                ("C:/Windows/Fonts/msyh.ttc", "china-ss"),  # 微软雅黑
                ("C:/Windows/Fonts/simhei.ttf", "china-ss"),  # 黑体
                ("C:/Windows/Fonts/simsun.ttc", "china-ss"),  # 宋体
                ("C:/Windows/Fonts/msgothic.ttc", "japan-ss"),  # MS Gothic (日文)
                ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "china-ss"),  # Linux 文泉驿
                ("/System/Library/Fonts/PingFang.ttc", "china-ss"),  # macOS 苹方
            ]
            
            font_name = "china-ss"  # 默认字体名称
            for font_path, font_alias in font_paths:
                if os.path.exists(font_path):
                    chinese_font_path = font_path
                    font_name = font_alias
                    print(f"[{job_id}] 使用字体: {font_path}")
                    break
            
            if not chinese_font_path:
                print(f"[{job_id}] 警告: 未找到中文字体，中文可能显示为方框")
            
            # 按页面收集所有需要替换的文本位置
            replacements = []  # (page_num, bbox, original_text, translated_text)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text_dict = page.get_text("dict")
                
                for block in get_pymupdf_blocks(text_dict):
                    if block.get("type") == 0:  # 文本块
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                original_text = span.get("text", "").strip()
                                if original_text:
                                    # 查找是否有对应的翻译
                                    for orig, trans in translated_map.items():
                                        if orig in original_text or original_text in orig:
                                            bbox = span.get("bbox")
                                            if bbox:
                                                replacements.append((page_num, bbox, original_text, trans))
                                                break
            
            # 按页面分组处理，从后往前处理（避免位置变化影响）
            from collections import defaultdict
            page_replacements = defaultdict(list)
            for page_num, bbox, orig, trans in replacements:
                page_replacements[page_num].append((bbox, orig, trans))
            
            for page_num in page_replacements:
                page = pdf_document[page_num]
                # 按y坐标从大到小排序（从下到上处理）
                items = sorted(page_replacements[page_num], key=lambda x: x[0][1], reverse=True)
                
                for bbox, original_text, translated_text in items:
                    x0, y0, x1, y1 = bbox
                    
                    try:
                        if overwrite:
                            # 覆盖模式：用白色矩形覆盖原文，然后插入翻译
                            # 1. 添加白色背景矩形覆盖原文
                            white_rect = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
                            page.draw_rect(white_rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                            
                            # 2. 插入翻译文本（黑色，原字体大小）
                            if chinese_font_path:
                                # 嵌入字体并插入文本
                                font_buffer = open(chinese_font_path, "rb").read()
                                page.insert_font(fontname="MyCJK", fontbuffer=font_buffer)
                                page.insert_textbox(
                                    fitz.Rect(x0, y0, x1 + 50, y1 + 20),
                                    translated_text[:200],
                                    fontsize=9,
                                    color=(0, 0, 0),  # 黑色
                                    overlay=True,
                                    fontname="MyCJK"
                                )
                            else:
                                page.insert_textbox(
                                    fitz.Rect(x0, y0, x1 + 50, y1 + 20),
                                    translated_text[:200],
                                    fontsize=9,
                                    color=(0, 0, 0),
                                    overlay=True
                                )
                        else:
                            # 注释模式：在原文下方插入红色翻译
                            text_box = fitz.Rect(x0, y1, x1 + 100, y1 + 30)
                            
                            if chinese_font_path:
                                # 嵌入字体并插入文本
                                font_buffer = open(chinese_font_path, "rb").read()
                                page.insert_font(fontname="MyCJK", fontbuffer=font_buffer)
                                page.insert_textbox(
                                    text_box,
                                    translated_text[:100],
                                    fontsize=8,
                                    color=(1, 0, 0),  # 红色
                                    overlay=True,
                                    fontname="MyCJK"
                                )
                            else:
                                page.insert_textbox(
                                    text_box,
                                    translated_text[:100],
                                    fontsize=8,
                                    color=(1, 0, 0),
                                    overlay=True
                                )
                    except Exception as e:
                        print(f"[{job_id}] 处理文本失败: {e}")
            
            # 保存新PDF
            new_filename = f"翻译_{target_lang}_{name_without_ext}.pdf"
            fd, temp_file_path = tempfile.mkstemp(suffix='.pdf')
            os.close(fd)
            pdf_document.save(temp_file_path)
            pdf_document.close()
            
            print(f"[{job_id}] PDF处理完成，保留原有格式")
            return temp_file_path, new_filename
            
        except ImportError as e:
            print(f"[{job_id}] 缺少pymupdf库，回退到txt: {e}")
            ext = '.txt'
        except Exception as e:
            print(f"[{job_id}] PDF处理失败: {e}")
            ext = '.txt'
    
    if ext == '.docx':
        # DOCX处理
        try:
            from io import BytesIO
            import docx
            
            doc = docx.Document(BytesIO(file_bytes))
            original_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 翻译
            translated_paragraphs = []
            batch_size = 10
            for i in range(0, len(original_paragraphs), batch_size):
                batch = original_paragraphs[i:i+batch_size]
                batch_text = "\n".join(batch)
                prompt = f"{lang_hint}\n请翻译以下内容成{target_lang}：\n{batch_text}"
                resp = get_openai_client().chat.completions.create(
                    model=model,
                    messages=[{"role": "system", "content": TRANSLATE_SYSTEM_PROMPT}, {"role": "user", "content": prompt}],
                    temperature=0.5, max_tokens=65535
                )
                translated_text = resp.choices[0].message.content or ""
                translated_paragraphs.extend(translated_text.split('\n'))
            
            # 创建新docx
            new_doc = docx.Document()
            for orig, trans in zip(original_paragraphs, translated_paragraphs):
                if orig.strip():
                    new_doc.add_paragraph(orig)
                if trans.strip():
                    new_doc.add_paragraph(trans)
            
            new_filename = f"翻译_{target_lang}_{name_without_ext}.docx"
            fd, temp_file_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            new_doc.save(temp_file_path)
            
            return temp_file_path, new_filename
            
        except ImportError as e:
            print(f"[{job_id}] 缺少python-docx库，回退到txt: {e}")
            ext = '.txt'
    
    # TXT和其他格式
    file_content = safe_decode(file_bytes)
    if not file_content or len(file_content.strip()) < 10:
        raise Exception("文件内容为空或无法读取")
    
    file_text = file_content[:MAX_AI_CHARS] if len(file_content) > MAX_AI_CHARS else file_content
    prompt = f"{lang_hint}\n请将以下文本翻译成{target_lang}，保持原有格式：\n{file_text}"
    
    resp = get_openai_client().chat.completions.create(
        model=model,
        messages=[{"role": "system", "content": TRANSLATE_SYSTEM_PROMPT}, {"role": "user", "content": prompt}],
        temperature=0.5, max_tokens=65535
    )
    result_text = resp.choices[0].message.content or ""
    
    new_filename = f"翻译_{target_lang}_{name_without_ext}.txt"
    fd, temp_file_path = tempfile.mkstemp(suffix='.txt', text=True)
    os.close(fd)
    with open(temp_file_path, 'w', encoding='utf-8') as f:
        f.write(result_text)
    
    return temp_file_path, new_filename

def background_translate_and_upload(job_id, file_url, filename, target_lang, group_id, upload_baseurl, is_pro=False, overwrite=False):
    """后台线程：下载、翻译并上传文件"""
    print(f"[{job_id}] 开始翻译上传: {filename} -> {target_lang} -> group {group_id}, 覆盖模式: {overwrite}")
    
    try:
        sess = get_session()
        
        # 下载文件
        resp = sess.get(file_url, timeout=60)
        if resp.status_code != 200:
            raise Exception(f"文件下载失败: {resp.status_code}")
        
        file_bytes = resp.content
            
        # 翻译并生成新文件
        temp_file_path, new_filename = translate_and_save_file(job_id, file_bytes, filename, target_lang, is_pro, overwrite)
        
        try:
            # 等待确保文件关闭
            time.sleep(0.5)
            
            # 上传到群文件
            upload_url = f"{upload_baseurl}/upload_group_file?group_id={group_id}&file=file://{temp_file_path}&name={new_filename}"
            upload_resp = sess.get(upload_url, timeout=60)
            upload_result = upload_resp.json()
            
            if upload_result.get('status') == 'ok':
                JOB_CACHE[job_id]['status'] = 'done'
                JOB_CACHE[job_id]['msg'] = f'翻译完成并已上传到群文件: {new_filename}'
                print(f"[{job_id}] 翻译并上传成功: {new_filename}")
            else:
                raise Exception(f"上传失败: {upload_result}")
        finally:
            time.sleep(0.5)
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
            except Exception as e:
                print(f"[{job_id}] 清理临时文件失败: {e}")
        
    except Exception as e:
        print(f"[{job_id}] 翻译上传失败: {e}")
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['msg'] = f"翻译上传失败：{str(e)}"

import re

def get_theme_config(theme_name):
    """获取不同主题的色彩和字体配置 (对比度增强版)"""
    if theme_name == 'cyberpunk': # 赛博朋克
        return {'bg': (13, 17, 23), 'header_bg': (255, 0, 85), 'header_text': (0, 255, 255), 'text': (0, 235, 170), 'title': (255, 255, 255), 'highlight': (255, 215, 0), 'quote': (150, 150, 150), 'border': (0, 150, 150), 'italic': (255, 100, 200), 'font_type': 'sans'}
    elif theme_name == 'historical': # 历史/古风 (加深了墨水黑和朱砂红，提升对比度)
        return {'bg': (230, 218, 195), 'header_bg': (94, 64, 52), 'header_text': (240, 230, 210), 'text': (30, 20, 15), 'title': (20, 10, 5), 'highlight': (160, 30, 30), 'quote': (90, 60, 50), 'border': (150, 120, 100), 'italic': (40, 70, 120), 'font_type': 'serif'}
    elif theme_name == 'cthulhu': # 克苏鲁深海风 (大幅提亮了文字颜色，暗底更清晰)
        return {'bg': (10, 15, 12), 'header_bg': (20, 45, 35), 'header_text': (180, 220, 190), 'text': (180, 220, 200), 'title': (150, 240, 170), 'highlight': (220, 70, 70), 'quote': (110, 150, 130), 'border': (60, 120, 90), 'italic': (170, 100, 220), 'font_type': 'serif'}
    elif theme_name == 'wasteland': # 废土/末日风
        return {'bg': (40, 35, 30), 'header_bg': (85, 45, 20), 'header_text': (220, 200, 180), 'text': (210, 190, 160), 'title': (255, 160, 50), 'highlight': (255, 90, 20), 'quote': (140, 120, 100), 'border': (120, 80, 40), 'italic': (180, 200, 100), 'font_type': 'sans'}
    elif theme_name == 'anime': # 二次元/软萌风
        return {'bg': (255, 248, 250), 'header_bg': (255, 182, 193), 'header_text': (255, 255, 255), 'text': (90, 70, 80), 'title': (220, 100, 140), 'highlight': (255, 100, 120), 'quote': (180, 150, 160), 'border': (255, 200, 220), 'italic': (120, 160, 255), 'font_type': 'sans'}
    elif theme_name == 'terminal': # 终端黑客风
        return {'bg': (0, 0, 0), 'header_bg': (0, 40, 0), 'header_text': (0, 255, 0), 'text': (0, 200, 0), 'title': (0, 255, 0), 'highlight': (0, 255, 0), 'quote': (0, 100, 0), 'border': (0, 150, 0), 'italic': (0, 255, 0), 'font_type': 'sans'}
    elif theme_name == 'classic': # 经典风 (完美复刻原版的深蓝灰+米黄背景)
        return {'bg': (242, 241, 237), 'header_bg': (52, 73, 94), 'header_text': (255, 255, 255), 'text': (40, 40, 40), 'title': (20, 20, 20), 'highlight': (192, 57, 43), 'quote': (100, 100, 100), 'border': (180, 180, 180), 'italic': (41, 128, 185), 'font_type': 'sans'}
    else: # default 简约风 (纯白底黑字)
        return {'bg': (255, 255, 255), 'header_bg': (240, 240, 240), 'header_text': (50, 50, 50), 'text': (30, 30, 30), 'title': (0, 0, 0), 'highlight': (200, 50, 50), 'quote': (120, 120, 120), 'border': (200, 200, 200), 'italic': (50, 100, 200), 'font_type': 'sans'}

def extract_theme_from_text(result_text, current_theme):
    """嗅探大模型输出的主题标签，动态改变主题"""
    final_theme = current_theme
    match = re.search(r'【主题[：:](.*?)】|\[Theme[：:](.*?)\]', result_text, re.IGNORECASE)
    if match:
        detected = (match.group(1) or match.group(2)).strip()
        if '赛博' in detected or 'cyber' in detected: final_theme = 'cyberpunk'
        elif '历史' in detected or '古风' in detected: final_theme = 'historical'
        elif '克苏鲁' in detected or 'cthulhu' in detected: final_theme = 'cthulhu'
        elif '废土' in detected or '末日' in detected: final_theme = 'wasteland'
        elif '二次元' in detected or '萌' in detected: final_theme = 'anime'
        elif '终端' in detected or '黑客' in detected: final_theme = 'terminal'
        elif '经典' in detected or '原版' in detected: final_theme = 'classic'
        elif '简约' in detected or '默认' in detected: final_theme = 'default'
        
        result_text = re.sub(r'【主题[：:](.*?)】\n*|\[Theme[：:](.*?)\]\n*', '', result_text, count=1, flags=re.IGNORECASE).strip()
    return result_text, final_theme

def load_markdown_fonts(font_type='sans'):
    """动态加载各级标题的字体，支持衬线与非衬线切换"""
    font_path = FONT_PATH if os.path.exists(FONT_PATH) else "./fonts/SimHei.ttf"
    bold_path = font_path.replace("msyh.ttc", "msyhbd.ttc")
    
    # 历史风格尝试加载楷体或宋体
    if font_type == 'serif':
        if os.path.exists("C:/Windows/Fonts/simkai.ttf"):
            font_path = bold_path = "C:/Windows/Fonts/simkai.ttf"
        elif os.path.exists("C:/Windows/Fonts/simsun.ttc"):
            font_path = bold_path = "C:/Windows/Fonts/simsun.ttc"
            
    if not os.path.exists(bold_path): bold_path = font_path

    try:
        return {
            'normal': ImageFont.truetype(font_path, 26), 'bold': ImageFont.truetype(bold_path, 26),
            'h1': ImageFont.truetype(bold_path, 38), 'h2': ImageFont.truetype(bold_path, 32),
            'h3': ImageFont.truetype(bold_path, 28), 'h4': ImageFont.truetype(bold_path, 24),
            'h5': ImageFont.truetype(bold_path, 22), 'title': ImageFont.truetype(bold_path, 42),
            'small': ImageFont.truetype(font_path, 20)
        }
    except:
        df = ImageFont.load_default()
        return {k: df for k in['normal','bold','h1','h2','h3','h4','h5','title','small']}

def parse_markdown_layout(text, width, padding, fonts, colors):
    text = text.replace('\t', '    ').replace('\r', '')
    text = re.sub(r'[\u2600-\u27BF\U0001F300-\U0001FAFF]', '', text)
    text = re.sub(r'^```.*$', '', text, flags=re.MULTILINE)
    text = text.replace('<think>', '\n>[AI 思考过程]：\n> ').replace('</think>', '\n---\n')

    latex_reps = {r'$\rightarrow$':'→', r'\rightarrow':'→', r'$\leftarrow$':'←', r'\leftarrow':'←', r'$\Rightarrow$':'⇒', r'\Rightarrow':'⇒', r'$\leftrightarrow$':'↔', r'\leftrightarrow':'↔', r'$\uparrow$':'↑', r'\uparrow':'↑', r'$\downarrow$':'↓', r'\downarrow':'↓', r'$\times$':'×', r'\times':'×', r'$\div$':'÷', r'\div':'÷', r'$\ge$':'≥', r'\ge':'≥', r'$\geq$':'≥', r'\geq':'≥', r'$\le$':'≤', r'\le':'≤', r'$\leq$':'≤', r'\leq':'≤', r'$\neq$':'≠', r'\neq':'≠', r'$\approx$':'≈', r'\approx':'≈', r'$\pm$':'±', r'\pm':'±', r'$\cdot$':'·', r'\cdot':'·', r'$\dots$':'...', r'\dots':'...'}
    for old_s, new_s in latex_reps.items(): text = text.replace(old_s, new_s)
    text = re.sub(r'(?<!\*)\*(?!\s)(.*?)(?<!\s)\*(?!\*)', '\x02\\1\x02', text)

    lines = text.split('\n')
    layout =[]
    max_w = width - padding
    global_bold = False; global_italic = False

    i = 0
    while i < len(lines):
        line = lines[i].strip(); raw_line = lines[i]
        if not line:
            layout.append({'type': 'spacing', 'height': 20}); i += 1
            continue

        if line.startswith('|') and line.endswith('|'):
            table_lines =[]
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip()); i += 1
            rows =[]
            for t_line in table_lines: rows.append([c.strip() for c in t_line.strip('|').split('|')])
            if len(rows) > 1 and all(re.match(r'^[\s\-:]+$', c) for c in rows[1]): rows.pop(1)
            if not rows: continue
            
            num_cols = max(len(r) for r in rows)
            for r in rows:
                while len(r) < num_cols: r.append("")
            col_widths = [0] * num_cols; font = fonts['normal']
            for r in rows:
                for j, cell in enumerate(r):
                    clean_cell = cell.replace('**', '').replace('\x02', '')
                    col_widths[j] = max(col_widths[j], font.getlength(clean_cell) + 20)
            total_w = sum(col_widths); max_table_w = width - 2 * padding
            if total_w > max_table_w: col_widths =[max(40, int(w / total_w * max_table_w)) for w in col_widths]
            elif total_w < max_table_w:
                extra = max_table_w - total_w
                for j in range(num_cols): col_widths[j] += extra // num_cols
                    
            layout.append({'type': 'table_border', 'height': 2})
            for row_idx, r in enumerate(rows):
                row_cells_wrapped =[]; max_lines = 1
                for j, cell in enumerate(r):
                    cell_lines = []; curr_line =[]; curr_x = 0
                    tokens = re.split(r'(\*\*|\x02)', cell)
                    is_bold = (row_idx == 0); is_italic = False
                    for part in tokens:
                        if part == '**': is_bold = not is_bold; continue
                        elif part == '\x02': is_italic = not is_italic; continue
                        if not part: continue
                        c_font_style = 'bold' if is_bold else 'normal'
                        c_color = colors['text']
                        if is_italic and c_font_style == 'normal': c_font_style = 'bold'; c_color = colors['italic']
                        c_font = fonts[c_font_style]
                        curr_chunk = ""
                        for char in part:
                            char_w = c_font.getlength(char)
                            if curr_x + char_w > col_widths[j] - 10:
                                if curr_chunk: curr_line.append((curr_chunk, c_font_style, c_color))
                                cell_lines.append(curr_line); curr_line =[]; curr_x = char_w; curr_chunk = char
                            else: curr_chunk += char; curr_x += char_w
                        if curr_chunk: curr_line.append((curr_chunk, c_font_style, c_color))
                    if curr_line: cell_lines.append(curr_line)
                    if not cell_lines: cell_lines = [[("", 'normal', colors['text'])]]
                    row_cells_wrapped.append(cell_lines); max_lines = max(max_lines, len(cell_lines))
                layout.append({'type': 'table_row', 'height': max_lines * 34 + 20, 'cells': row_cells_wrapped, 'col_widths': col_widths, 'is_header': row_idx == 0})
                layout.append({'type': 'table_border', 'height': 2})
            layout.append({'type': 'spacing', 'height': 20})
            continue

        if re.match(r'^[-*_]{3,}$', line):
            layout.append({'type': 'hr', 'height': 30}); i += 1; continue

        is_quote = False; base_color = colors['text']; font_style = 'normal'; start_x = padding
        header_match = re.match(r'^(#{1,5})\s+(.*)', line)
        list_match = re.match(r'^(\s*)[*+-]\s+(.*)', raw_line)
        
        if header_match:
            level = len(header_match.group(1)); line = header_match.group(2).replace('**', '').replace('\x02', '')
            font_style = f'h{level}'; base_color = colors['title']
        elif line.startswith('> '):
            is_quote = True; line = line[2:]; base_color = colors['quote']; start_x += 20
        elif list_match:
            indent = len(list_match.group(1)) // 2; line = '• ' + list_match.group(2); start_x += indent * 20
        elif line.startswith('【') and '】' in line:
            base_color = colors['highlight']; font_style = 'bold'; line = line.replace('**', '').replace('\x02', '')

        tokens = re.split(r'(\*\*|\x02)', line)
        current_line_elements =[]; current_x = start_x; line_max_h = 0
        for part in tokens:
            if part == '**': global_bold = not global_bold; continue
            elif part == '\x02': global_italic = not global_italic; continue
            if not part: continue
            
            curr_style = font_style; curr_color = base_color
            if curr_style == 'normal':
                if global_bold: curr_style = 'bold'
                if global_italic: curr_style = 'bold'; curr_color = colors['italic']
                    
            font = fonts[curr_style]
            chunk_h = {'h1':50, 'h2':42, 'h3':36, 'h4':32, 'h5':30, 'bold':34, 'normal':34}.get(curr_style, 34)
            curr_chunk = ""
            for char in part:
                char_w = font.getlength(char)
                if current_x + char_w > max_w:
                    if curr_chunk: current_line_elements.append((curr_chunk, font, curr_color, current_x - font.getlength(curr_chunk)))
                    line_max_h = max(line_max_h, chunk_h)
                    layout.append({'type': 'text_line', 'height': line_max_h, 'elements': current_line_elements, 'is_quote': is_quote, 'start_x': start_x})
                    current_line_elements =[]; current_x = start_x + (20 if is_quote else 0)
                    curr_chunk = char; current_x += char_w; line_max_h = chunk_h
                else: curr_chunk += char; current_x += char_w
            if curr_chunk:
                current_line_elements.append((curr_chunk, font, curr_color, current_x - font.getlength(curr_chunk)))
                line_max_h = max(line_max_h, chunk_h)
        if current_line_elements: layout.append({'type': 'text_line', 'height': line_max_h, 'elements': current_line_elements, 'is_quote': is_quote, 'start_x': start_x})
        layout.append({'type': 'spacing', 'height': 10})
        i += 1
    return layout

def text_to_images(text, file_title, report_title="TRPG 模组解析报告", theme='default', token_usage=''):
    """支持多风格模板与 Tokens 显示的终极排版引擎"""
    parts =[p.strip() for p in text.split('【分页符】') if p.strip()]
    if not parts: parts = [text]
        
    width, padding = 900, 50
    colors = get_theme_config(theme)
    fonts = load_markdown_fonts(colors['font_type'])
    images_bytes =[]
    
    for idx, part in enumerate(parts):
        layout = parse_markdown_layout(part, width, padding, fonts, colors)
        total_h = 110 + sum(line['height'] for line in layout) + 60 + padding
        
        img = Image.new('RGB', (width, total_h), colors['bg'])
        draw = ImageDraw.Draw(img)
        
        draw.rectangle([(0, 0), (width, 110)], fill=colors['header_bg'])
        draw.text((padding, 30), report_title, font=fonts['title'], fill=colors['header_text'])
        page_text = f"Page {idx+1}/{len(parts)} - {file_title[:15]}..." if len(parts) > 1 else f"{file_title[:20]}..."
        draw.text((padding+520, 48), page_text, font=fonts['small'], fill=colors['header_text'])
        
        y = 140
        for line in layout:
            if line['type'] == 'text_line':
                if line['is_quote']:
                    draw.rectangle([(padding, y + 4), (padding + 4, y + line['height'] - 4)], fill=colors['quote'])
                for txt, font_obj, color, x_pos in line['elements']:
                    draw.text((x_pos, y), txt, font=font_obj, fill=color)
            elif line['type'] == 'table_border':
                draw.line([(padding, y), (width - padding, y)], fill=colors['border'], width=2)
            elif line['type'] == 'table_row':
                col_x = padding
                for j, cell_lines in enumerate(line['cells']):
                    draw.line([(col_x, y), (col_x, y + line['height'])], fill=colors['border'], width=1)
                    text_y = y + 10
                    for c_line in cell_lines:
                        cell_x = col_x + 10
                        for txt, f_style, f_color in c_line:
                            draw.text((cell_x, text_y), txt, font=fonts[f_style], fill=f_color)
                            cell_x += fonts[f_style].getlength(txt)
                        text_y += 34
                    col_x += line['col_widths'][j]
                draw.line([(col_x, y), (col_x, y + line['height'])], fill=colors['border'], width=1)
            elif line['type'] == 'hr':
                draw.line([(padding, y + 15), (width - padding, y + 15)], fill=colors['border'], width=2)
            y += line['height']
            
        footer_text = f"AI 来自 Air {token_usage}"
        draw.text((padding, total_h-40), footer_text, font=fonts['small'], fill=colors['quote'])
        
        buf = BytesIO()
        img.save(buf, 'PNG')
        buf.seek(0)
        images_bytes.append(buf.getvalue())
        
    return images_bytes

def text_to_image(text, key_id):
    """向下兼容的单图模式"""
    images_bytes = text_to_images(text, key_id, "TRPG 跑团日志评分")
    return Image.open(BytesIO(images_bytes[0]))

# --- 百度网盘 OAuth 2.0 鉴权管理 ---
def get_valid_access_token():
    """获取有效的 access_token，如果过期则自动刷新"""
    token_data = {}
    
    # 1. 尝试从本地加载已保存的 Token
    if os.path.exists(BAIDU_TOKEN_FILE):
        try:
            with open(BAIDU_TOKEN_FILE, 'r', encoding='utf-8') as f:
                token_data = json.load(f)
        except Exception as e:
            print(f"读取 Token 文件失败: {e}")

    # 2. 检查 access_token 是否有效（预留 300 秒的缓冲时间）
    if token_data and 'access_token' in token_data:
        if time.time() < token_data.get('expires_at', 0) - 300:
            return token_data['access_token']
        
        # 3. 如果已过期，尝试使用 refresh_token 刷新
        refresh_token = token_data.get('refresh_token')
        if refresh_token:
            print("Access Token 已过期，正在自动刷新...")
            try:
                return refresh_baidu_token(refresh_token)
            except Exception as e:
                print(f"Token 刷新失败: {e}，将尝试使用 Authorization Code 重新获取。")

    # 4. 如果没有 Token，或者刷新失败，尝试使用配置中的 BAIDU_AUTH_CODE 获取
    if BAIDU_AUTH_CODE:
        print("正在使用 Authorization Code 首次获取 Token...")
        return fetch_new_token_with_code(BAIDU_AUTH_CODE)
    
    raise Exception("无法获取有效的百度网盘 access_token，请检查 AppKey、SecretKey 以及 Auth Code 配置。")

def fetch_new_token_with_code(code):
    """使用授权码(Code)换取首次的 Access Token"""
    url = "https://openapi.baidu.com/oauth/2.0/token"
    params = {
        "grant_type": "authorization_code",
        "code": code,
        "client_id": BAIDU_APP_KEY,
        "client_secret": BAIDU_SECRET_KEY,
        "redirect_uri": "oob"
    }
    resp = requests.get(url, params=params).json()
    if "access_token" in resp:
        # 计算过期时间戳 (当前时间 + 有效期秒数)
        resp['expires_at'] = time.time() + resp.get('expires_in', 2592000)
        with open(BAIDU_TOKEN_FILE, 'w', encoding='utf-8') as f:
            json.dump(resp, f)
        return resp['access_token']
    else:
        raise Exception(f"使用 Code 获取 Token 失败: {resp.get('error_description', resp)}")

def refresh_baidu_token(refresh_token):
    """使用 refresh_token 刷新 Access Token"""
    url = "https://openapi.baidu.com/oauth/2.0/token"
    params = {
        "grant_type": "refresh_token",
        "refresh_token": refresh_token,
        "client_id": BAIDU_APP_KEY,
        "client_secret": BAIDU_SECRET_KEY
    }
    resp = requests.get(url, params=params).json()
    if "access_token" in resp:
        resp['expires_at'] = time.time() + resp.get('expires_in', 2592000)
        with open(BAIDU_TOKEN_FILE, 'w', encoding='utf-8') as f:
            json.dump(resp, f)
        return resp['access_token']
    else:
        # 如果 refresh 也失效了（通常是10年过期或被用户手动撤销），建议删掉文件重新用code获取
        if os.path.exists(BAIDU_TOKEN_FILE):
            os.remove(BAIDU_TOKEN_FILE)
        raise Exception(f"刷新 Token 失败，授权可能已失效，请重新获取 Code: {resp.get('error_description', resp)}")


# --- 百度网盘 API 业务函数 ---
def baidu_search_files(keyword, limit=15):
    """在指定目录搜索，返回匹配的多个文件/文件夹（默认最多限制15个以防过大）"""
    token = get_valid_access_token()
    url = f"https://pan.baidu.com/rest/2.0/xpan/file?method=search&access_token={token}"
    params = {'key': keyword, 'dir': BAIDU_TARGET_DIR, 'recursion': 1}
    resp = requests.get(url, params=params).json()
    if resp.get('errno') == 0 and resp.get('list'):
        # 截取前 limit 个结果返回
        return resp['list'][:limit]
    return []

def baidu_create_share(fs_ids):
    """创建包含多个文件/文件夹的合并分享链接"""
    import random
    import string
    
    token = get_valid_access_token()
    url = f"https://pan.baidu.com/rest/2.0/xpan/share?method=set&access_token={token}"
    
    random_pwd = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
    
    # 将多个 fs_id 转换成 "[id1,id2,id3]" 的 JSON 数组格式
    fid_list_str = "[" + ",".join(map(str, fs_ids)) + "]"
    
    data = {
        'fid_list': fid_list_str,
        'schannel': 4,
        'channel_list': '[]',
        'period': 1,
        'pwd': random_pwd
    }
    
    headers: Dict[str, str | bytes] = {"User-Agent": "pan.baidu.com"}
    resp = requests.post(url, data=data, headers=headers).json()
    
    if resp.get('errno') == 0:
        return f"链接: {resp.get('link')}\n提取码: {resp.get('pwd', random_pwd)}"
    else:
        raise Exception(f"创建合并分享失败: {resp}")

import posixpath
import concurrent.futures

# --- 重构的高效下载组件 ---
def baidu_get_dlinks_batch(fs_ids, token):
    """批量获取多个文件的 dlink 下载链接"""
    if not fs_ids: return {}
    url = f"https://pan.baidu.com/rest/2.0/xpan/multimedia?method=filemetas&access_token={token}"
    dlinks_map = {}
    
    # 百度要求一次请求最多约100个fs_id，我们分块处理
    for i in range(0, len(fs_ids), 100):
        chunk = fs_ids[i:i+100]
        params = {'fsids': '[' + ','.join(map(str, chunk)) + ']', 'dlink': 1}
        try:
            resp = requests.get(url, params=params, timeout=15).json()
            if resp.get('errno') == 0:
                for item in resp.get('list', []):
                    if item.get('dlink'):
                        dlinks_map[item['fs_id']] = item['dlink']
        except Exception as e:
            print(f"批量获取 dlink 失败: {e}")
    return dlinks_map

def download_baidu_file_to_temp(dlink, save_path, token):
    """使用特定的请求头并启用长连接下载文件"""
    download_url = f"{dlink}&access_token={token}"
    # 官方文档强烈要求必须带的 User-Agent
    headers: Dict[str, str | bytes] = {"User-Agent": "pan.baidu.com"}
    
    try:
        # 使用 Session 处理 302 跳转更稳定，加入超时防卡死
        with requests.Session() as s:
            resp = s.get(download_url, headers=headers, stream=True, timeout=(10, 60))
            resp.raise_for_status()
            with open(save_path, 'wb') as f:
                for chunk in resp.iter_content(chunk_size=1024*64): # 64KB一块，提高写入效率
                    if chunk: 
                        f.write(chunk)
    except Exception as e:
        print(f"文件下载失败 [{save_path}]: {e}")

def baidu_collect_files_recursive(target_dir, base_save_path, token, collected_tasks):
    """递归遍历文件夹，但不立即下载，只收集需要下载的任务列表"""
    url = f"https://pan.baidu.com/rest/2.0/xpan/file?method=list&access_token={token}"
    params = {'dir': target_dir, 'limit': 1000} # 单页最多拉取1000条
    
    try:
        resp = requests.get(url, params=params, timeout=15).json()
        if resp.get('errno') != 0: return

        for item in resp.get('list', []):
            if item['isdir'] == 1:
                new_dir = os.path.join(base_save_path, item['server_filename'])
                os.makedirs(new_dir, exist_ok=True)
                # 递归进下一层
                baidu_collect_files_recursive(item['path'], new_dir, token, collected_tasks)
            else:
                file_save_path = os.path.join(base_save_path, item['server_filename'])
                collected_tasks.append((item['fs_id'], file_save_path))
    except Exception as e:
        print(f"获取目录列表失败 [{target_dir}]: {e}")


# --- 优化后的搜索及打包上传模块 ---
def background_search_module(job_id, keyword, is_local, group_id, upload_baseurl):
    print(f"[{job_id}] 开始搜索模组: {keyword}, 本地模式: {is_local}")
    try:
        token = get_valid_access_token()
        
        # 1. 扩大初始拉取数量，供我们在本地进行精准过滤 (拉取前100个)
        raw_targets = baidu_search_files(keyword, limit=100)
        if not raw_targets:
            raise Exception(f"在网盘库中未找到包含 '{keyword}' 的内容")
        
        # 2. 核心优化 1：严格精准匹配与安全拦截
        # 将用户输入的关键字按空格拆分，确保每一个词都必须在文件名中出现
        keywords_list =[k.lower() for k in keyword.split()]
        strict_targets =[]
        
        for t in raw_targets:
            filename_lower = t['server_filename'].lower()
            
            # 强制拦截：文件名必须包含所有搜索关键词，剔除百度胡乱推荐的无关文件！
            if not all(k in filename_lower for k in keywords_list):
                continue
                
            # 安全拦截：过滤掉所有带有【】标识的保护级大合集文件夹
            if t['isdir'] == 1 and '{' in t['server_filename'] and '}' in t['server_filename']:
                print(f"[{job_id}] 安全拦截：丢弃集合文件夹 {t['server_filename']}")
                continue
                
            strict_targets.append(t)

        if not strict_targets:
            raise Exception(f"搜索词 '{keyword}' 未能精准匹配到有效模组（或命中了被保护的合集文件夹）。")

        # 3. 核心优化 2：路径折叠去重（只保留主文件夹，踢出多余的子文件）
        # 将所有结果按路径长度从小到大排序，这样父文件夹一定排在子文件前面
        strict_targets.sort(key=lambda x: len(x['path']))
        
        dedup_targets =[]
        accepted_dirs = set()
        
        for t in strict_targets:
            path = t['path']
            is_sub = False
            
            # 检查当前文件是否已经被包含在某个已被采纳的父文件夹中
            for ad in accepted_dirs:
                if path.startswith(ad + '/'):
                    is_sub = True
                    break
            
            # 如果不是任何已记录文件夹的子文件，才把它加入最终列表
            if not is_sub:
                dedup_targets.append(t)
                # 如果这个本身就是一个文件夹，记录它的路径
                if t['isdir'] == 1:
                    accepted_dirs.add(path)
        
        # 4. 按所属父级目录进行智能分组
        grouped_targets = {}
        for t in dedup_targets:
            parent_dir = posixpath.dirname(t['path'])
            if parent_dir not in grouped_targets:
                grouped_targets[parent_dir] = []
            grouped_targets[parent_dir].append(t)
        
        # 5. 网盘分享模式
        if not is_local:
            share_messages =[]
            # 最多处理前 5 个不同的文件夹区，防刷屏
            for parent_dir, targets_in_dir in list(grouped_targets.items())[:5]:
                fs_ids =[t['fs_id'] for t in targets_in_dir]
                names = [t['server_filename'] for t in targets_in_dir]
                
                try:
                    share_text = baidu_create_share(fs_ids)
                    folder_name = posixpath.basename(parent_dir) if parent_dir != '/' else '根目录'
                    names_str = "、".join(names[:4]) + ("..." if len(names)>4 else "")
                    
                    share_messages.append(f"📁 来自【{folder_name}】:\n📄 {names_str}\n{share_text}")
                except Exception as e:
                    print(f"[{job_id}] 创建分享失败 {parent_dir}: {e}")
            
            if not share_messages:
                raise Exception("网盘分享链接生成失败，可能是接口限制。")
            
            final_msg = f"🔍 找到了分散在不同文件夹的相关模组：\n\n" + "\n\n".join(share_messages)
            if len(grouped_targets) > 5:
                final_msg += f"\n\n(为防刷屏，已折叠其余 {len(grouped_targets)-5} 个文件夹的结果)"
                
            JOB_CACHE[job_id]['status'] = 'done'
            JOB_CACHE[job_id]['msg'] = final_msg
            return

        # 6. 本地下载打包模式
        temp_dir = tempfile.mkdtemp()
        try:
            print(f"[{job_id}] 正在分析目标结构...")
            group_folder_name = f"搜索结果_{keyword}"
            group_folder_path = os.path.join(temp_dir, group_folder_name)
            os.makedirs(group_folder_path, exist_ok=True)
            
            all_download_tasks = []
            total_files_limit = 30 # 限制总下载数
            names_display_lines =[]
            
            for parent_dir, targets_in_dir in list(grouped_targets.items())[:10]: 
                folder_name = posixpath.basename(parent_dir) if parent_dir != '/' else '根目录'
                names = [t['server_filename'] for t in targets_in_dir]
                names_display_lines.append(f"📁 【{folder_name}】: " + "、".join(names[:3]) + ("..." if len(names)>3 else ""))
                
                for t in targets_in_dir:
                    if len(all_download_tasks) >= total_files_limit: break
                    
                    server_filename = t['server_filename']
                    if t['isdir'] == 0:
                        save_path = os.path.join(group_folder_path, folder_name, server_filename)
                        os.makedirs(os.path.dirname(save_path), exist_ok=True)
                        all_download_tasks.append((t['fs_id'], save_path))
                    else:
                        sub_folder = os.path.join(group_folder_path, folder_name, server_filename)
                        os.makedirs(sub_folder, exist_ok=True)
                        baidu_collect_files_recursive(t['path'], sub_folder, token, all_download_tasks)
            
            names_display = "\n".join(names_display_lines)
            if len(grouped_targets) > 10:
                names_display += f"\n...等共 {len(grouped_targets)} 个文件夹的匹配结果"

            # 批量换取下载直链并多线程下载
            print(f"[{job_id}] 总计 {len(all_download_tasks)} 个文件，正在批量换取下载直链...")
            task_fs_ids = [task[0] for task in all_download_tasks]
            dlinks_map = baidu_get_dlinks_batch(task_fs_ids, token)
            
            def _worker(task):
                fs_id, save_path = task
                dlink = dlinks_map.get(fs_id)
                if dlink: download_baidu_file_to_temp(dlink, save_path, token)

            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                list(executor.map(_worker, all_download_tasks))
            
            # 打包并上传
            final_filename = f"{group_folder_name}.zip"
            final_upload_path = os.path.join(temp_dir, final_filename)
            
            with zipfile.ZipFile(final_upload_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(group_folder_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.join(group_folder_name, os.path.relpath(file_path, group_folder_path))
                        zipf.write(file_path, arcname)

            sess = get_session()
            upload_url = f"{upload_baseurl}/upload_group_file?group_id={group_id}&file=file://{final_upload_path}&name={urllib.parse.quote(final_filename)}"
            upload_resp = sess.get(upload_url, timeout=300).json() 

            if upload_resp.get('status') == 'ok':
                JOB_CACHE[job_id]['status'] = 'done'
                JOB_CACHE[job_id]['msg'] = f"✅ 已将内容归档为【{final_filename}】并上传至群文件！\n包含以下内容：\n{names_display}"
            else:
                raise Exception(f"群文件上传失败: {upload_resp}")
                
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    except Exception as e:
        print(f"[{job_id}] 搜索处理失败: {e}")
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['msg'] = f"模组获取失败：{str(e)}"

# --- 新增 Flask API 路由 ---
@app.route('/api/search_module', methods=['GET'])
def search_module_task():
    """网盘模组搜索任务"""
    if len(JOB_CACHE) > 100: JOB_CACHE.clear()
    
    keyword = request.args.get('keyword')
    is_local = request.args.get('local', 'false').lower() == 'true'
    group_id = request.args.get('group_id', '')
    upload_baseurl = request.args.get('upload_url', '')

    if not keyword:
        return jsonify({'status': 'error', 'msg': '缺少搜索关键字'})

    job_id = str(uuid.uuid4())
    JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time()}
    
    executor.submit(background_search_module, job_id, keyword, is_local, group_id, upload_baseurl)
    
    return jsonify({'status': 'ok', 'id': job_id})

# --- API 接口 ---

def resolve_bridge_file_for_analysis(group_id, public_base="", index=0):
    cleanup_expired()
    public_base = str(public_base or "").rstrip('/')

    with STATE_LOCK:
        state = snapshot_item(UPLOAD_STATES.get(group_id))
        file_list = LATEST_FILES.get(group_id, [])
        resolved_index = index if index >= 0 else max(0, len(file_list) - 1)
        item = snapshot_item(file_list[resolved_index]) if file_list and 0 <= resolved_index < len(file_list) else None
        last_err = LAST_ERROR_BY_GROUP.get(group_id, '')

    if state:
        status = str(state.get('status', 'processing'))
        if status in ('queued', 'processing'):
            raise RuntimeError(f"群文件仍在处理中: {state.get('name', '') or state.get('file_id', '')}")
        if status == 'error' and not item:
            raise RuntimeError(str(state.get('last_error') or last_err or '群文件处理失败'))

    remote_info = None
    if REFRESH_LATEST_ON_READ:
        remote_info = get_latest_group_file_info(group_id)
        if remote_info:
            remote_info['public_base'] = public_base
        if should_refresh_cached_item(remote_info, item):
            item = process_group_upload(remote_info or {})
            with STATE_LOCK:
                LAST_ERROR_BY_GROUP.pop(group_id, None)

    if not item and PULL_LATEST_ON_EMPTY and not REFRESH_LATEST_ON_READ:
        item = pull_latest_from_napcat(group_id, public_base=public_base)
        if item:
            with STATE_LOCK:
                LAST_ERROR_BY_GROUP.pop(group_id, None)

    if not item:
        raise RuntimeError(last_err or '未找到可供分析的群文件')

    return hydrate_bridge_item(item, public_base=public_base) or dict(item)

@app.route('/api/halt', methods=['POST'])
def api_halt():
    """v4.4.0: 强制停止指定群内所有进行中的AI生成任务。"""
    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id', 0), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400

    halted = []
    with STATE_LOCK:
        for jid, job in list(JOB_CACHE.items()):
            j_group = safe_int(job.get('group_id', 0), 0)
            if j_group == group_id and job.get('status') == 'processing':
                if jid in CANCEL_FLAGS:
                    CANCEL_FLAGS[jid].set()
                job['status'] = 'cancelled'
                job['text'] = '任务已被停止。'
                halted.append(jid)

    # Clean up stale cancel flags (已取消超过3600秒的)
    now = time.time()
    stale = []
    for jid, flag in list(CANCEL_FLAGS.items()):
        if flag.is_set():
            job = JOB_CACHE.get(jid)
            if job:
                job_created = job.get('created', 0)
                if now - job_created > 3600:
                    stale.append(jid)
            else:
                stale.append(jid)
    for jid in stale:
        CANCEL_FLAGS.pop(jid, None)

    return jsonify({'status': 'ok', 'halted': halted, 'count': len(halted)})


@app.route('/api/submit', methods=['GET', 'POST'])
def submit_task():
    """提交日志分析任务 (统一且安全的参数提取)"""
    if len(JOB_CACHE) > 100: JOB_CACHE.clear()
    
    # 统一接收参数，无论是 GET 还是 POST 都转为字典提取，彻底消灭变量未定义异常！
    req_data = request.get_json(silent=True) or {} if request.method == 'POST' else request.args

    key = req_data.get('key') or req_data.get('url')
    password = req_data.get('password')
    source = req_data.get('source')
    keys = req_data.get('keys') or req_data.get('urls')
    passwords = req_data.get('passwords')
    sources = req_data.get('sources')
    is_pro = str(req_data.get('pro', 'false')).lower() == 'true'
    is_kind = str(req_data.get('kind', 'false')).lower() == 'true'
    mode = req_data.get('mode', 'analyze')
    persona = req_data.get('persona', '')
    custom_prompt = req_data.get('custom_prompt', '')
    theme = req_data.get('theme', 'default')
    get_text = str(req_data.get('get_text', 'false')).lower() == 'true'
    direct_text = req_data.get('text', '')  # v4.2: 直接文本输入（用于.ai无文件模式）

    # 规范化并优先使用 parse_log_target_entry 的推断结果（URL 主机名优先）
    if direct_text:
        # v4.2: 直接文本模式 — 跳过URL抓取，直接使用提供的文本
        job_id = str(uuid.uuid4())
        group_id_submit = safe_int(req_data.get('group_id', 0), 0)
        CANCEL_FLAGS[job_id] = threading.Event()
        JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time(), 'group_id': group_id_submit}
        if group_id_submit > 0:
            BRIDGE_POLL_GROUPS.add(group_id_submit)
            ensure_poll_worker_started()
        executor.submit(background_process_direct_text, job_id, direct_text, is_pro, is_kind, mode, persona, custom_prompt, theme, get_text, group_id_submit)
        return jsonify({'status': 'ok', 'id': job_id})

    if isinstance(keys, list) and len(keys) > 0:
        # 多项提交：逐项解析并构建 key/source/password 列表
        parsed_keys = []
        parsed_sources = []
        parsed_passwords = []
        for i, k in enumerate(keys):
            pw = (passwords[i] if isinstance(passwords, list) and i < len(passwords) else None) if passwords else None
            src_hint = (sources[i] if isinstance(sources, list) and i < len(sources) else (sources if isinstance(sources, str) else None))
            tgt = parse_log_target_entry(k, password=pw, source=src_hint)
            if tgt:
                parsed_keys.append(tgt.get('key') or k)
                parsed_sources.append(tgt.get('source') or (src_hint or infer_source_by_key(tgt.get('key') or k)))
                parsed_passwords.append(tgt.get('password') or pw)
            else:
                parsed_keys.append(k)
                parsed_sources.append(src_hint or infer_source_by_key(k))
                parsed_passwords.append(pw)
        key = parsed_keys
        password = parsed_passwords
        source = parsed_sources
    else:
        # 单项提交：若为 URL 或包含主机信息，parse_log_target_entry 返回的 source 优先
        tgt = parse_log_target_entry(key, password=password, source=source)
        if tgt:
            key = tgt.get('key') or key
            # 如果 parse 返回了明确来源，优先使用它
            if tgt.get('source'):
                source = tgt.get('source')
            if tgt.get('password'):
                password = tgt.get('password')
            else:
                if not source:
                    source = infer_source_by_key(key)
        else:
            if not source:
                source = infer_source_by_key(key)

    job_id = str(uuid.uuid4())
    group_id_submit = safe_int(req_data.get('group_id', 0), 0)
    CANCEL_FLAGS[job_id] = threading.Event()
    JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time(), 'group_id': group_id_submit}

    # Register for polling if using bridge services
    if group_id_submit > 0:
        BRIDGE_POLL_GROUPS.add(group_id_submit)
        ensure_poll_worker_started()
    # 将所有参数（包括 theme, group_id）传入后台线程
    executor.submit(background_process, job_id, key, password, source, is_pro, is_kind, mode, persona, custom_prompt, theme, group_id_submit, get_text)
    return jsonify({'status': 'ok', 'id': job_id})


@app.route('/api/submit_file', methods=['GET', 'POST'])
def submit_file_task():
    """提交本地文件分析任务 (统一且安全的参数提取)"""
    if len(JOB_CACHE) > 100: JOB_CACHE.clear()
    
    req_data = request.get_json(silent=True) or {} if request.method == 'POST' else request.args

    file_url = req_data.get('url')
    filename = req_data.get('filename')
    file_urls = req_data.get('urls')
    filenames = req_data.get('filenames')
    group_id = safe_int(req_data.get('group_id', 0), 0)
    mode = req_data.get('mode', 'analyze')
    is_pro = str(req_data.get('pro', 'false')).lower() == 'true'
    is_kind = str(req_data.get('kind', 'false')).lower() == 'true'
    persona = req_data.get('persona', '')
    custom_prompt = req_data.get('custom_prompt', '')
    theme = req_data.get('theme', 'default')
    get_text = str(req_data.get('get_text', 'false')).lower() == 'true'

    if isinstance(file_urls, list) and len(file_urls) > 0:
        file_url = [str(u or '').strip() for u in file_urls if str(u or '').strip()]
        if isinstance(filenames, list) and len(filenames) > 0:
            filename = [str(n or '').strip() for n in filenames]
        else:
            filename = [f"file_{i+1}.txt" for i in range(len(file_url))]

    if isinstance(file_url, list) and len(file_url) == 1:
        file_url = file_url[0]
        if isinstance(filename, list):
            filename = filename[0] if filename else 'file_1.txt'

    bridge_item = None
    if group_id > 0 and not file_url:
        try:
            bridge_item = resolve_bridge_file_for_analysis(group_id, public_base=request.host_url.rstrip('/'))
        except Exception as exc:
            return jsonify({'status': 'error', 'msg': f'bridge resolve failed: {str(exc)}'}), 400

        file_url = bridge_item.get('content_url')
        filename = bridge_item.get('text_filename') or bridge_item.get('name') or f'group_{group_id}.txt'

    if (not file_url) or (not filename): 
        return jsonify({'status': 'error', 'msg': 'Missing url or filename'})

    job_id = str(uuid.uuid4())
    CANCEL_FLAGS[job_id] = threading.Event()
    JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time(), 'group_id': group_id}

    # 将所有参数（包括 theme）传入后台线程
    executor.submit(background_file_process, job_id, file_url, filename, mode, is_pro, is_kind, persona, custom_prompt, theme, get_text, group_id)
    if bridge_item:
        return jsonify({
            'status': 'ok',
            'id': job_id,
            'bridge_file': {
            'group_id': bridge_item.get('group_id'),
            'name': bridge_item.get('name'),
            'text_filename': bridge_item.get('text_filename'),
            'content_url': bridge_item.get('content_url'),
            },
        })
    return jsonify({'status': 'ok', 'id': job_id})

@app.route('/api/status', methods=['GET'])
def check_status():
    """查询任务状态，附带图像数量"""
    job_id = request.args.get('id')
    job = JOB_CACHE.get(job_id)
    if not job: return jsonify({'status': 'not_found'})
    
    # 兼容老版只返回单图的逻辑以及新版的多图逻辑
    img_count = len(job.get('images', [])) if 'images' in job else (1 if 'image' in job else 0)
    
    resp = {
        'status': job['status'],
        'msg': job.get('msg', ''),
        'image_count': img_count
    }
    if job.get('text_key'):
        resp['text_key'] = job.get('text_key', '')
        resp['text_filename'] = job.get('text_filename', 'ai_analysis.txt')
    return jsonify(resp)

@app.route('/api/result', methods=['GET'])
def get_result():
    """获取最终图片 (支持index下标获取指定分页)"""
    job_id = request.args.get('id')
    index = int(request.args.get('index', 0))
    job = JOB_CACHE.get(job_id)
    
    if not job: return "Result not found", 404
    
    # 获取图像数据（兼容新老字段）
    img_data = None
    if 'images' in job and index < len(job['images']):
        img_data = job['images'][index]
    elif 'image' in job and index == 0:
        img_data = job['image']
        
    if not img_data: return "Index out of range or not ready", 404
    
    return send_file(
        BytesIO(img_data),
        mimetype='image/png',
        download_name=f'log_analysis_{job_id}_{index}.png'
    )

@app.route('/api/result_text', methods=['GET'])
def get_result_text():
    """获取AI分析的原始文本结果（get_text模式）"""
    job_id = request.args.get('id')
    job = JOB_CACHE.get(job_id)

    if not job: return "Result not found", 404

    text_data = job.get('text', '')
    if not text_data: return "Text not available or not ready", 404

    return jsonify({
        'status': 'ok',
        'text': text_data,
        'id': job_id,
    })

# --- 文本生成图片任务 ---
import random

def background_generate_image(job_id, prompt, size="1024x1024"):
    print(f"[{job_id}] 开始生成图片(NovelAI V4.5 官方复刻模式): {prompt}, 尺寸: {size}")
    try:
        # NovelAI V4.5 最优分辨率映射
        width, height = 1024, 1024
        if size == "1792x1024":
            width, height = 1216, 832   # 横图
        elif size == "1024x1792":
            width, height = 832, 1216   # 竖图

        url = "https://image.novelai.net/ai/generate-image"
        headers: Dict[str, str | bytes] = {
            "Authorization": f"Bearer {IMAGE_API_KEY.strip()}",
            "Content-Type": "application/json"
        }
        
        # 强制负面提示词：首位死锁 nsfw，并融合官方 V4.5 的默认负面起手式
        negative_prompt = "nsfw, lowres, artistic error, film grain, scan artifacts, worst quality, bad quality, jpeg artifacts, very displeasing, chromatic aberration, dithering, halftone, screentone, multiple views, logo, too many watermarks, negative space, blank page, bad anatomy, bad hands, text, error, missing fingers, extra digit, fewer digits, cropped, normal quality, signature, watermark, username, blurry"
        
        # 像素级复刻官方 Payload 结构，填补引发 500 报错的所有缺失字段
        payload = {
            "input": prompt,
            "model": "nai-diffusion-4-5-full",
            "action": "generate",
            "parameters": {
                "params_version": 3,
                "width": width,
                "height": height,
                "scale": 5.0,
                "sampler": "k_euler_ancestral",  # 官方默认采样器
                "steps": 28,
                "seed": random.randint(1, 999999999), # 必须带上随机种子
                "n_samples": 1,
                "qualityToggle": True,
                "dynamic_thresholding": False,
                "controlnet_strength": 1.0,
                "legacy": False,
                "add_original_image": False,
                "cfg_rescale": 0,
                "noise_schedule": "karras",
                "legacy_v3_extend": False,
                "use_coords": False,
                "legacy_uc": False,
                "characterPrompts": [],
                # 【修复核心】完整的 V4 提示词块
                "v4_prompt": {
                    "caption": {
                        "base_caption": prompt,
                        "char_captions": []
                    },
                    "use_coords": False,
                    "use_order": True
                },
                # 【修复核心】完整的 V4 负面提示词块（不传这个必定报 500）
                "v4_negative_prompt": {
                    "caption": {
                        "base_caption": negative_prompt,
                        "char_captions": []
                    },
                    "legacy_uc": False
                },
                "negative_prompt": negative_prompt,
                "deliberate_euler_ancestral_bug": False,
                "prefer_brownian": True,
                "image_format": "png"
            }
        }
        
        # 发送请求
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
        
        if resp.status_code != 200:
            raise Exception(f"NovelAI 接口返回错误: {resp.status_code} - {resp.text}")
            
        img_bytes = None
        
        # 按照 NovelAI 规范解压 ZIP（或者直接接收 PNG）
        import zipfile
        from io import BytesIO
        try:
            with zipfile.ZipFile(BytesIO(resp.content)) as zip_ref:
                img_bytes = zip_ref.read("image_0.png")
        except zipfile.BadZipFile:
            # V4.5 API 在指定 image_format: png 时可能直接返回图片文件
            img_bytes = resp.content
            
        if not img_bytes:
            raise Exception("未能从 NovelAI 返回的数据中提取出图片。")
        
        # 保存图像字节流
        JOB_CACHE[job_id]['status'] = 'done'
        JOB_CACHE[job_id]['image'] = img_bytes
        print(f"[{job_id}] 图片生成完成 (V4.5 完美版)")

    except Exception as e:
        print(f"[{job_id}] 图片生成失败: {e}")
        err_img_bytes = text_to_images(f"NovelAI 绘图失败：\n{str(e)}", "绘图错误")[0]
        JOB_CACHE[job_id]['status'] = 'error'
        JOB_CACHE[job_id]['images'] = [err_img_bytes]

@app.route('/api/submit_image_gen', methods=['GET'])
def submit_image_gen_task():
    """提交生成图片任务"""
    if len(JOB_CACHE) > 100: JOB_CACHE.clear()
    
    prompt = request.args.get('prompt')
    size = request.args.get('size', '1024x1024')
    
    if not prompt:
        return jsonify({'status': 'error', 'msg': '缺少 prompt'})

    job_id = str(uuid.uuid4())
    JOB_CACHE[job_id] = {'status': 'processing', 'created': time.time()}
    
    executor.submit(background_generate_image, job_id, prompt, size)
    
    return jsonify({'status': 'ok', 'id': job_id})


@app.route('/health', methods=['GET'])
def health():
    cleanup_expired()
    with STATE_LOCK:
        last_event = dict(LAST_EVENT_SUMMARY)
    return jsonify({
        'status': 'ok',
        'groups': len(LATEST_FILES),
        'pending_groups': len(UPLOAD_STATES),
        'poll_groups': len(BRIDGE_POLL_GROUPS),
        'poll_interval_sec': BRIDGE_POLL_INTERVAL_SEC,
        'queue_size': UPLOAD_QUEUE.qsize(),
        'cache_dir': BRIDGE_CACHE_DIR,
        'napcat_base_candidates': NAPCAT_BASE_CANDIDATES,
        'napcat_last_ok_base': LAST_NAPCAT_BASE,
        'napcat_last_error': LAST_NAPCAT_ERROR,
        'last_event': last_event,
    })


@app.route('/napcat/event', methods=['POST'])
def on_napcat_event():
    if not check_auth(request):
        bridge_log('auth', 'reject /napcat/event unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    with STATE_LOCK:
        LAST_EVENT_SUMMARY.clear()
        LAST_EVENT_SUMMARY.update({
            'post_type': str(payload.get('post_type', '')),
            'notice_type': str(payload.get('notice_type', '')),
            'group_id': safe_int(payload.get('group_id', 0), 0),
            'user_id': safe_int(payload.get('user_id', 0), 0),
            'ts': now_ts(),
        })

    bridge_log(
        'event',
        f"recv post_type={LAST_EVENT_SUMMARY['post_type']} notice_type={LAST_EVENT_SUMMARY['notice_type']} group={LAST_EVENT_SUMMARY['group_id']}",
    )

    info = extract_group_upload(payload)
    if not info:
        bridge_log('event', 'ignored: not a supported upload event')
        return jsonify({'status': 'ignored'})

    # Register this group for periodic polling (fallback if webhook misses events)
    gid_event = safe_int(payload.get('group_id', 0), 0) or safe_int(info.get('group_id', 0), 0)
    if gid_event > 0:
        BRIDGE_POLL_GROUPS.add(gid_event)
        ensure_poll_worker_started()

    # Dedup check: same (group_id, file_id) within RECENT_FILE_DEDUPE_WINDOW_SEC gets skipped
    info_group_id = safe_int(info.get('group_id', 0), 0)
    info_file_id = str(info.get('file_id', ''))
    dedup_key = (info_group_id, info_file_id)
    now_ts_dedup = time.time()
    # Purge stale entries (>60s old)
    stale = [k for k, v in RECENT_UPLOAD_EVENTS.items() if now_ts_dedup - v > 60]
    for k in stale:
        RECENT_UPLOAD_EVENTS.pop(k, None)
    last_seen = RECENT_UPLOAD_EVENTS.get(dedup_key)
    if last_seen and (now_ts_dedup - last_seen) <= RECENT_FILE_DEDUPE_WINDOW_SEC:
        bridge_log('event', f'dedup skipped group={info_group_id} file_id={info_file_id} (seen {now_ts_dedup - last_seen:.2f}s ago)')
        return jsonify({'status': 'dedup_skipped', 'file_id': info_file_id})
    RECENT_UPLOAD_EVENTS[dedup_key] = now_ts_dedup

    info['public_base'] = request.host_url.rstrip('/')
    cleanup_expired()

    try:
        queued = enqueue_group_upload(info)
        with STATE_LOCK:
            LAST_EVENT_SUMMARY['file_id'] = queued['file_id']
            LAST_EVENT_SUMMARY['name'] = queued['name']
            LAST_EVENT_SUMMARY['queue_size'] = queued.get('queue_size', 0)

        bridge_log('queue', f"queued group={queued['group_id']} file={queued['name']} qsize={queued.get('queue_size', 0)}")
        return jsonify({
            'status': 'queued',
            'file': queued,
            'queue_size': queued.get('queue_size', 0),
        })
    except Exception as exc:
        gid = safe_int(info.get('group_id', 0), 0)
        if gid > 0:
            with STATE_LOCK:
                LAST_ERROR_BY_GROUP[gid] = str(exc)
        bridge_log('error', f"group={gid} {str(exc)}")
        return jsonify({'status': 'error', 'message': str(exc)}), 500


@app.route('/bridge/latest', methods=['POST', 'GET'])
def bridge_latest():
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /bridge/latest unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    cleanup_expired()
    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id', request.args.get('group_id', 0)), 0)
    index = safe_int(payload.get('index', request.args.get('index', -1)), -1)
    if group_id <= 0:
        bridge_log('latest', f'invalid group_id payload={payload}')
        return jsonify({'status': 'error', 'message': 'invalid group_id'}), 400

    # Register this group for periodic polling
    BRIDGE_POLL_GROUPS.add(group_id)
    ensure_poll_worker_started()

    public_base = request.host_url.rstrip('/')

    with STATE_LOCK:
        state = snapshot_item(UPLOAD_STATES.get(group_id))
        file_list = LATEST_FILES.get(group_id, [])
        resolved_index = index if index >= 0 else max(0, len(file_list) - 1)
        item = snapshot_item(file_list[resolved_index]) if file_list and 0 <= resolved_index < len(file_list) else None
        last_err = LAST_ERROR_BY_GROUP.get(group_id, '')

    if state:
        hydrated_state = hydrate_bridge_item(state, public_base=public_base) or dict(state)
        bridge_log('latest', f"{hydrated_state.get('status', 'processing')} group={group_id} file={hydrated_state.get('name', '')}")
        return jsonify({'status': hydrated_state.get('status', 'processing'), 'file': hydrated_state})

    remote_info = None
    if REFRESH_LATEST_ON_READ:
        try:
            remote_info = get_latest_group_file_info(group_id)
            if remote_info:
                remote_info['public_base'] = public_base
            if should_refresh_cached_item(remote_info, item):
                pulled = process_group_upload(remote_info or {})
                with STATE_LOCK:
                    LAST_ERROR_BY_GROUP.pop(group_id, None)
                bridge_log(
                    'latest',
                    f"refresh success group={group_id} file={pulled.get('name', '')} chars={pulled.get('text_chars', 0)} source_ts={pulled.get('source_ts', 0)}",
                )
                # Auto-import to logutil if recording
                try:
                    auto_import_bridge_file_to_logutil(pulled)
                except Exception:
                    pass
                return jsonify({'status': 'ok', 'file': hydrate_bridge_item(pulled, public_base=public_base), 'source': 'refresh'})
        except Exception as exc:
            with STATE_LOCK:
                LAST_ERROR_BY_GROUP[group_id] = str(exc)
            bridge_log('latest', f'refresh failed group={group_id} err={exc}')

    if not item:
        if PULL_LATEST_ON_EMPTY and not REFRESH_LATEST_ON_READ:
            try:
                bridge_log('latest', f'empty group={group_id}, try pull from napcat')
                pulled = pull_latest_from_napcat(group_id, public_base=public_base)
                if pulled:
                    with STATE_LOCK:
                        LAST_ERROR_BY_GROUP.pop(group_id, None)
                    bridge_log('latest', f"pull success group={group_id} file={pulled.get('name', '')} chars={pulled.get('text_chars', 0)}")
                    # Auto-import to logutil if recording
                    try:
                        auto_import_bridge_file_to_logutil(pulled)
                    except Exception:
                        pass
                    return jsonify({'status': 'ok', 'file': hydrate_bridge_item(pulled, public_base=public_base), 'source': 'pull'})
            except Exception as exc:
                with STATE_LOCK:
                    LAST_ERROR_BY_GROUP[group_id] = str(exc)
                bridge_log('latest', f'pull failed group={group_id} err={exc}')

        bridge_log('latest', f'empty group={group_id} last_err={last_err}')
        return jsonify({
            'status': 'empty',
            'group_id': group_id,
            'last_error': last_err,
            'hint': '请先确认 /napcat/event 是否收到 group_upload 事件，或检查 get_group_root_files/get_group_file_url/解析步骤日志',
        })

    hydrated_item = hydrate_bridge_item(item, public_base=public_base) or dict(item)
    bridge_log('latest', f"ok group={group_id} index={index} file={hydrated_item.get('name', '')} chars={hydrated_item.get('text_chars', 0)}")
    return jsonify({
        'status': 'ok',
        'file': hydrated_item,
        'file_count': len(file_list),
        'index': index,
    })


@app.route('/bridge/list', methods=['POST', 'GET'])
def bridge_list():
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /bridge/list unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    cleanup_expired()
    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id', request.args.get('group_id', 0)), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'message': 'invalid group_id'}), 400

    # Register this group for periodic polling
    BRIDGE_POLL_GROUPS.add(group_id)
    ensure_poll_worker_started()

    public_base = request.host_url.rstrip('/')
    with STATE_LOCK:
        file_list = [snapshot_item(f) for f in LATEST_FILES.get(group_id, [])]
        link_list = [snapshot_item(l) for l in LINK_CACHE.get(group_id, [])]

    files = []
    for idx, item in enumerate(file_list):
        h = hydrate_bridge_item(item, public_base=public_base)
        if h:
            h['_index'] = idx
            # Use the real file name in output
            h['display_name'] = h.get('name', '')
            h['_type'] = 'file'
            ck = h.get('content_key', '')
            h['preview'] = get_content_preview(ck, 12)
            files.append(h)

    links = []
    for idx, item in enumerate(link_list):
        h = hydrate_bridge_item(item, public_base=public_base)
        if h:
            h['_index'] = idx
            h['display_name'] = h.get('url', h.get('name', ''))
            h['_type'] = 'link'
            ck = h.get('content_key', '')
            h['preview'] = get_content_preview(ck, 12)
            links.append(h)

    # v4.4.5: include history items for this group (matching instance backend)
    with STATE_LOCK:
        hist_all = [snapshot_item(h) for h in HISTORY]
    history = []
    filtered_idx = 0
    for item in hist_all:
        if not item:
            continue
        item_gid = safe_int(item.get('group_id', 0), 0)
        if item_gid != group_id:
            continue
        hydrated = hydrate_bridge_item(item, public_base=public_base)
        if hydrated:
            hydrated['_index'] = filtered_idx
            hydrated['display_name'] = hydrated.get('name', hydrated.get('url', ''))
            hydrated['_type'] = hydrated.get('_type', 'file')
            ck = hydrated.get('content_key', '')
            hydrated['preview'] = get_content_preview(ck, 12)
            history.append(hydrated)
            filtered_idx += 1

    bridge_log('list', f"group={group_id} files={len(files)} links={len(links)} history={len(history)}")
    return jsonify({'status': 'ok', 'group_id': group_id, 'files': files, 'links': links, 'history': history, 'count': len(files)})


@app.route('/bridge/content/<content_key>', methods=['GET'])
def bridge_content(content_key):
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /bridge/content unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    cleanup_expired()
    with STATE_LOCK:
        path = CONTENT_INDEX.get(content_key, '')
    if not path or not os.path.exists(path):
        bridge_log('content', f'not_found key={content_key}')
        abort(404)

    bridge_log('content', f'serve key={content_key} path={path}')
    return send_file(path, mimetype='text/plain')


# ====== v4.4.0: Web GUI ======

_BRIDGE_GUI_HTML = r'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>LogAI Bridge Manager</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:system-ui,sans-serif;background:#1a1a2e;color:#e0e0e0;padding:20px;max-width:1200px;margin:0 auto}
h1{color:#e94560;margin-bottom:10px}
h2{color:#e94560;margin:20px 0 10px}
.group-select{margin-bottom:20px}
.group-select select,.group-select input{padding:8px;font-size:14px;border-radius:4px;border:1px solid #333;background:#16213e;color:#e0e0e0}
table{width:100%;border-collapse:collapse;margin-bottom:30px}
th,td{padding:8px 10px;text-align:left;border-bottom:1px solid #333;font-size:13px}
th{background:#16213e;color:#e94560;position:sticky;top:0}
td a{color:#4fc3f7;text-decoration:none}
td a:hover{text-decoration:underline}
tr:hover{background:#16213e}
.cmd-box{margin:20px 0;display:flex;gap:8px}
.cmd-box input[type=text]{flex:1;padding:10px;font-size:14px;border-radius:4px;border:1px solid #333;background:#16213e;color:#e0e0e0}
.cmd-box button{padding:10px 20px;font-size:14px;border-radius:4px;border:none;background:#e94560;color:#fff;cursor:pointer}
.cmd-box button:hover{background:#c73a52}
.result{background:#16213e;padding:12px;border-radius:4px;margin:10px 0;white-space:pre-wrap;font-size:13px;max-height:300px;overflow-y:auto}
preview{color:#888;font-size:12px}
</style>
</head>
<body>
<h1>LogAI Bridge Manager v4.5.3</h1>
<div class="group-select">
  <b>选择群组:</b>
  <select id="groupSelect" onchange="loadData()" style="padding:8px;font-size:14px;border-radius:4px;border:1px solid #333;background:#16213e;color:#e0e0e0;min-width:200px;">
    <option value="">-- 选择群组 --</option>
  </select>
  <button onclick="loadData()">刷新</button>
</div>

<h2>【文件】Files</h2>
<div id="fileTable"></div>

<h2>【链接】Links</h2>
<div id="linkTable"></div>

<h2>【历史】History</h2>
<div id="historyTable"></div>

<h2>命令输入</h2>
<div class="cmd-box">
  <input type="text" id="cmdInput" placeholder="输入命令，如: .logai [file]-0 pro">
  <select id="cmdGroup" style="padding:10px;background:#16213e;color:#e0e0e0;border:1px solid #333;border-radius:4px;"></select>
  <button onclick="sendCommand()">发送</button>
</div>
<div class="result" id="cmdResult"></div>

<script>
const defaultGroup='{{GROUP_ID}}';

async function initGroups(){
  const resp=await fetch('/api/bridge_groups');
  const data=await resp.json();
  const groups=data.groups||[];
  let opts='<option value="">-- 选择群组 --</option>';
  groups.forEach(g=>{
    const sel=g===defaultGroup?' selected':'';
    opts+=`<option value="${g}"${sel}>${g}</option>`;
  });
  const gs=document.getElementById('groupSelect');
  gs.innerHTML=opts;
  // Sync command group selector
  let sel2='';
  groups.forEach(g=>{const s=g===defaultGroup?' selected':'';sel2+=`<option value="${g}"${s}>${g}</option>`});
  document.getElementById('cmdGroup').innerHTML=sel2||'<option>--</option>';
  if(defaultGroup) loadData();
}

async function loadData(){
  const gid=document.getElementById('groupSelect').value;
  if(!gid){document.getElementById('fileTable').innerHTML='';document.getElementById('linkTable').innerHTML='';document.getElementById('historyTable').innerHTML='';return}
  const resp=await fetch(`/api/bridge_gui_data?group_id=${gid}`);
  const data=await resp.json();
  renderTable('fileTable',data.files||[],['#','名称','字数','保存时间','开头'],(r)=>`<a href="/bridge/content/${r.content_key}">${esc(r.name||'?')}</a>`);
  renderTable('linkTable',data.links||[],['#','URL','字数','保存时间','开头'],(r)=>`<a href="/bridge/content/${r.content_key}">${esc((r.url||'?')).slice(0,60)}</a>`);
  renderTable('historyTable',data.history||[],['#','类型','名称','字数','开头'],(r)=>`${r._type||'?'} ${esc((r.name||r.url||'?')).slice(0,60)}`);
}
function renderTable(id,items,cols,nameFn){
  if(!items.length){document.getElementById(id).innerHTML='<p style="color:#888">（无数据）</p>';return}
  let h=`<table><tr>${cols.map(c=>`<th>${c}</th>`).join('')}</tr>`;
  items.forEach((r,i)=>{
    const preview=(r.preview||r.name||r.url||'').slice(0,12);
    const ts=r.ts?new Date(r.ts*1000).toLocaleString():'-';
    h+=`<tr><td>${i}</td><td>${nameFn?nameFn(r):esc(r.name||'?')}</td><td>${r.text_chars||0}</td><td>${ts}</td>${cols.length>4?`<td><preview>${esc(preview)}</preview></td>`:''}</tr>`;
  });
  h+='</table>';
  document.getElementById(id).innerHTML=h;
}
function esc(s){return String(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')}
async function sendCommand(){
  const gid=document.getElementById('cmdGroup').value;
  const cmd=document.getElementById('cmdInput').value.trim();
  if(!gid||!cmd){return}
  document.getElementById('cmdResult').textContent='处理中...';
  const resp=await fetch('/api/bridge_gui_command',{
    method:'POST',
    headers:{'Content-Type':'application/json'},
    body:JSON.stringify({group_id:parseInt(gid),command:cmd})
  });
  const data=await resp.json();
  document.getElementById('cmdResult').textContent=JSON.stringify(data,null,2);
}
initGroups();
</script>
</body>
</html>'''


@app.route('/api/bridge_master', methods=['POST', 'GET'])
def api_bridge_master():
    """v4.4.1: .bridge master 命令 - 返回 Web UI 地址。"""
    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id') or request.args.get('group_id', 0), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    public_base = resolve_public_base_or_fallback()
    gui_url = f"{public_base}/bridge/gui/{group_id}"
    return jsonify({'status': 'ok', 'gui_url': gui_url})


@app.route('/bridge/gui', methods=['GET'])
def bridge_gui():
    """v4.4.0: 桥接管理 Web GUI（通用入口，需手动输入群号）。"""
    return _BRIDGE_GUI_HTML.replace('{{GROUP_ID}}', '')


@app.route('/bridge/gui/<int:group_id>', methods=['GET'])
def bridge_gui_group(group_id):
    """v4.4.0: 桥接管理 Web GUI（指定群号）。"""
    return _BRIDGE_GUI_HTML.replace('{{GROUP_ID}}', str(group_id))


@app.route('/api/bridge_groups', methods=['GET'])
def api_bridge_groups():
    """v4.4.5: 返回所有存在桥接数据的群号列表。"""
    groups = set()
    with STATE_LOCK:
        for gid in LATEST_FILES:
            if LATEST_FILES.get(gid):
                groups.add(str(gid))
        for gid in LINK_CACHE:
            if LINK_CACHE.get(gid):
                groups.add(str(gid))
        for item in HISTORY:
            gid = str(item.get('group_id', ''))
            if gid:
                groups.add(gid)
        for gid in BRIDGE_POLL_GROUPS:
            groups.add(str(gid))
    return jsonify({'status': 'ok', 'groups': sorted(groups, key=lambda x: int(x) if x.isdigit() else 0)})


@app.route('/api/bridge_gui_data', methods=['GET'])
def api_bridge_gui_data():
    """v4.4.0: Web GUI 数据接口。"""
    group_id = safe_int(request.args.get('group_id', 0), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400

    with STATE_LOCK:
        files = [snapshot_item(f) for f in LATEST_FILES.get(group_id, [])]
        links = [snapshot_item(l) for l in LINK_CACHE.get(group_id, [])]
        # v4.4.5: filter history by group_id, use sequential numbering matching api_bridge_list
        hist_all = [snapshot_item(h) for h in HISTORY]
        history = []
        filtered_idx = 0
        for h in hist_all:
            if not h:
                continue
            if safe_int(h.get('group_id', 0), 0) != group_id:
                continue
            h['_index'] = filtered_idx
            history.append(h)
            filtered_idx += 1

    return jsonify({'status': 'ok', 'files': files, 'links': links, 'history': history})


@app.route('/api/bridge_gui_command', methods=['POST'])
def api_bridge_gui_command():
    """v4.4.0: Web GUI 命令输入接口（转发给后端处理）。"""
    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id', 0), 0)
    command = str(payload.get('command', '')).strip()
    if group_id <= 0 or not command:
        return jsonify({'status': 'error', 'msg': 'missing group_id or command'}), 400

    # 简单解析命令并模拟处理
    # 此处仅提供基本桥接，完整功能仍需通过QQ/SealDice
    return jsonify({'status': 'ok', 'msg': f'命令已接收: {command[:100]}', 'note': 'Web GUI命令功能为辅助性质，完整功能请通过QQ使用。'})


@app.route('/api/bridge_poll_on', methods=['POST'])
def api_bridge_poll_on():
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /api/bridge_poll_on unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id') or request.args.get('group_id', 0), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400

    BRIDGE_POLL_GROUPS.add(group_id)
    ensure_poll_worker_started()
    bridge_log('poll', f"bridge on: group={group_id} total_groups={len(BRIDGE_POLL_GROUPS)}")
    return jsonify({'status': 'ok', 'group_id': group_id, 'poll_active': True, 'groups_count': len(BRIDGE_POLL_GROUPS)})


@app.route('/api/bridge_poll_off', methods=['POST'])
def api_bridge_poll_off():
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /api/bridge_poll_off unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id') or request.args.get('group_id', 0), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400

    BRIDGE_POLL_GROUPS.discard(group_id)
    bridge_log('poll', f"bridge off: group={group_id} remaining_groups={len(BRIDGE_POLL_GROUPS)}")
    return jsonify({'status': 'ok', 'group_id': group_id, 'poll_active': False, 'groups_count': len(BRIDGE_POLL_GROUPS)})


@app.route('/api/bridge_poll_status', methods=['GET', 'POST'])
def api_bridge_poll_status():
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /api/bridge_poll_status unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id') or request.args.get('group_id', 0), 0)
    group_active = group_id in BRIDGE_POLL_GROUPS if group_id > 0 else False
    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'poll_active': group_active,
        'total_groups': len(BRIDGE_POLL_GROUPS),
    })


@app.route('/api/bridge_list', methods=['GET', 'POST'])
def api_bridge_list():
    """返回指定群组桥接缓存中的全部文件列表。"""
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /api/bridge_list unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id') or request.args.get('group_id', 0), 0)
    filter_type = str(payload.get('filter') or request.args.get('filter', 'all')).lower()
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400

    show_files = filter_type in ('all', 'file')
    show_links = filter_type in ('all', 'link')
    show_history = filter_type in ('all', 'history')
    show_history_only = filter_type == 'history'

    file_items = []
    link_items = []
    history_items = []

    with STATE_LOCK:
        if show_files:
            file_list = list(LATEST_FILES.get(group_id, []))
        else:
            file_list = []
        if show_links:
            link_list = list(LINK_CACHE.get(group_id, []))
        else:
            link_list = []
        if show_history:
            hist_list = list(HISTORY)
        else:
            hist_list = []

    for idx, item in enumerate(file_list):
        ck = item.get('content_key', '')
        path_cached = CONTENT_INDEX.get(ck, '')
        file_items.append({
            'index': idx,
            'file_id': str(item.get('file_id', '')),
            'name': str(item.get('name', '')),
            'text_chars': safe_int(item.get('text_chars', 0), 0),
            'text_bytes': safe_int(item.get('text_bytes', 0), 0),
            'ts': safe_int(item.get('ts', 0), 0),
            'cached': bool(path_cached and os.path.exists(path_cached)),
            'content_key': ck,
            'content_url': item.get('content_url', ''),
            'preview': get_content_preview(ck, 12),
            '_type': 'file',
        })

    for idx, item in enumerate(link_list):
        ck = item.get('content_key', '')
        path_cached = CONTENT_INDEX.get(ck, '')
        link_items.append({
            'index': idx,
            'url': str(item.get('url', '')),
            'name': str(item.get('name', '')),
            'text_chars': safe_int(item.get('text_chars', 0), 0),
            'text_bytes': safe_int(item.get('text_bytes', 0), 0),
            'ts': safe_int(item.get('ts', 0), 0),
            'cached': bool(path_cached and os.path.exists(path_cached)),
            'content_key': ck,
            'content_url': item.get('content_url', ''),
            'preview': get_content_preview(ck, 12),
            '_type': 'link',
        })

    filtered_idx = 0
    for item in hist_list:
        # v4.4.4: always filter history by group_id for data isolation
        if safe_int(item.get('group_id', 0), 0) != group_id:
            continue
        ck = item.get('content_key', '')
        path_cached = CONTENT_INDEX.get(ck, '')
        history_items.append({
            'index': filtered_idx,
            'name': str(item.get('name', '')),
            'url': str(item.get('url', '')),
            'text_chars': safe_int(item.get('text_chars', 0), 0),
            'ts': safe_int(item.get('_evicted_ts', item.get('ts', 0)), 0),
            'cached': bool(path_cached and os.path.exists(path_cached)),
            'content_key': ck,
            'content_url': item.get('content_url', ''),
            '_type': item.get('_type', 'file'),
        })
        filtered_idx += 1

    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'total': len(file_items),
        'files': file_items,
        'links': link_items,
        'link_count': len(link_items),
        'history': history_items,
        'history_count': len(history_items),
        'poll_active': group_id in BRIDGE_POLL_GROUPS,
    })


# Per-group custom poll interval overrides (in seconds)
BRIDGE_POLL_INTERVAL_OVERRIDE: Dict[int, float] = {}


@app.route('/api/bridge_rate', methods=['POST'])
def api_bridge_rate():
    """设置指定群组的桥接轮询间隔（秒）。
    请求体: {group_id, interval_sec}"""
    if not check_auth_with_query(request):
        bridge_log('auth', 'reject /api/bridge_rate unauthorized')
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id') or request.args.get('group_id', 0), 0)
    interval_sec = safe_int(payload.get('interval_sec') or payload.get('rate', 0), 0)
    if group_id <= 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    if interval_sec <= 0:
        # Reset to default
        BRIDGE_POLL_INTERVAL_OVERRIDE.pop(group_id, None)
        bridge_log('poll', f"bridge rate reset: group={group_id} -> default ({BRIDGE_POLL_INTERVAL_SEC}s)")
        return jsonify({
            'status': 'ok',
            'group_id': group_id,
            'interval_sec': BRIDGE_POLL_INTERVAL_SEC,
            'is_custom': False,
        })

    BRIDGE_POLL_INTERVAL_OVERRIDE[group_id] = float(interval_sec)
    ensure_poll_worker_started()
    bridge_log('poll', f"bridge rate set: group={group_id} -> {interval_sec}s")
    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'interval_sec': int(interval_sec),
        'is_custom': True,
    })


# ====== 新增：del_paren 助手与 logutil API ======

def segments_to_text(message):
    """Convert a NapCat/OneBot Array-format message to plain text.
    Handles text/image/at/forward/file segment types.
    Mirrors fwlog's segments_to_text exactly."""
    if isinstance(message, str):
        return message
    if not isinstance(message, list):
        return str(message or "")
    parts = []
    for seg in message:
        if not isinstance(seg, dict):
            continue
        t = seg.get("type")
        d = seg.get("data") or {}
        if t == "text":
            parts.append(d.get("text", ""))
        elif t == "image":
            file_val = d.get("file", "")
            url_val = d.get("url") or d.get("file_url") or ""
            if file_val and url_val:
                parts.append(f"[CQ:image,file={file_val},url={url_val}]")
            elif file_val:
                parts.append(f"[CQ:image,file={file_val}]")
            elif url_val:
                parts.append(f"[CQ:image,url={url_val}]")
            else:
                parts.append("[图片]")
        elif t == "at":
            qq_val = d.get("qq", "")
            if qq_val:
                parts.append(f"[CQ:at,qq={qq_val}]")
        elif t == "forward":
            fid = d.get("id") or d.get("res_id") or d.get("message_id")
            if fid:
                parts.append(f"[CQ:forward,id={fid}]")
        else:
            parts.append(f"[{t}]")
    if not parts:
        return "[空消息]"
    return "".join(parts)

def parse_cq_params(segment_text):
    """Parse CQ code parameters from a segment like 'CQ:file,id=xxx,name=yyy'.
    Mirrors fwlog's parse_cq_params exactly."""
    params = {}
    for part in str(segment_text or "").split(","):
        if "=" not in part:
            continue
        key, value = part.split("=", 1)
        key = key.strip()
        value = value.strip()
        if key:
            params[key] = value
    return params

def strip_paren_text(text: str) -> str:
    """删除所有括号包裹的行和括号开头的段落。
    增强版：正确处理多行括号内容（如方括号管道格式产生的（多行内容）），
    避免括号段落后紧跟的非括号内容被误删。"""
    if not text:
        return text
    lines = text.splitlines()
    out_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            out_lines.append(line)
            i += 1
            continue
        # 1) 完全由括号包裹的一行（中文或英文括号）：整行删除
        if ((stripped.startswith('(') and stripped.endswith(')')) or
            (stripped.startswith('（') and stripped.endswith('）'))):
            i += 1
            continue
        # 2) 段落以左括号开头：尝试找到右括号闭合或空行为止
        if stripped.startswith('(') or stripped.startswith('（'):
            is_cn = stripped.startswith('（')
            close_char = '）' if is_cn else ')'
            j = i
            found_close = False
            while j < len(lines):
                cur = lines[j].strip()
                if not cur:  # 空行：段落结束
                    break
                if cur.endswith(close_char):
                    j += 1  # 包含闭合行
                    found_close = True
                    break
                j += 1
            if found_close:
                # 括号段落已闭合，跳过整个段落
                i = j
                continue
            else:
                # 未找到闭合括号，跳过直到空行（保留旧行为作为兜底）
                while j < len(lines) and lines[j].strip() != '':
                    j += 1
                i = j
                continue
        out_lines.append(line)
        i += 1
    return '\n'.join(out_lines)

def parse_structured_text_to_items(text, sender_name, sender_id, ts, raw_msg_id):
    """
    将文本解析为结构化 log_item 列表。
    支持 6 种 fwlog 格式 + 3 种新增方括号/星号格式。
    无结构化内容时回退为单条消息。
    """
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip("\n")
    if not normalized.strip():
        return []

    parsed_items = []
    prefix_lines = []
    current_name = None
    current_user_id = sender_id
    current_ts = safe_int(ts, int(time.time()))
    current_lines = []
    structured_found = False
    item_index = 0
    base_raw_id = raw_msg_id or f"parsed-{safe_int(ts, int(time.time()))}"
    lines = normalized.split("\n")
    pending_meta = None

    def flush_current():
        nonlocal item_index, current_name, current_user_id, current_ts, current_lines
        if current_name is None:
            return
        content = "\n".join(current_lines).strip("\n")
        if content.strip():
            parsed_items.append(
                make_log_item(current_name, current_user_id, current_ts, content, f"{base_raw_id}#{item_index}")
            )
            item_index += 1
        current_name = None
        current_user_id = sender_id
        current_ts = safe_int(ts, int(time.time()))
        current_lines = []

    line_index = 0
    while line_index < len(lines):
        line = lines[line_index]
        matched = match_speaker_line(line, ts)
        next_index = line_index + 1
        if not matched:
            multiline_matched = match_multiline_angle_speaker(
                lines,
                line_index,
                pending_meta["time"] if pending_meta else ts,
                pending_meta,
            )
            if multiline_matched:
                matched = multiline_matched[0]
                next_index = multiline_matched[1]

        if not matched and current_name is None:
            meta_only = parse_metadata_only_line(line, ts)
            if meta_only:
                if pending_meta:
                    prefix_lines.append(pending_meta["raw"])
                pending_meta = meta_only
                line_index += 1
                continue

        if matched:
            if pending_meta:
                matched["time"] = pending_meta["time"]
            structured_found = True
            flush_current()
            current_name = matched["name"]
            current_user_id = matched.get("user_id") or ""
            current_ts = safe_int(matched.get("time"), safe_int(ts, int(time.time())))
            current_lines = []
            if matched.get("content"):
                current_lines.append(matched["content"])
            pending_meta = None
            line_index = next_index
            continue

        if current_name is None:
            if pending_meta:
                prefix_lines.append(pending_meta["raw"])
                pending_meta = None
            if line.strip():
                prefix_lines.append(line)
        else:
            current_lines.append(line)

        line_index += 1

    flush_current()

    if pending_meta:
        prefix_lines.append(pending_meta["raw"])

    if structured_found:
        prefix_text = "\n".join(prefix_lines).strip()
        if prefix_text:
            parsed_items.insert(
                0,
                make_log_item(sender_name, sender_id, ts, prefix_text, f"{base_raw_id}#preface"),
            )
        return parsed_items

    return [make_log_item(sender_name, sender_id, ts, normalized.strip(), base_raw_id)]


async def extract_items_from_text_chunk(text, sender_name, sender_id, ts, raw_msg_id, group_id=None):
    """Extract log items from a text chunk.  URLs are resolved via
    fetch_log_text_by_source and parsed into structured items; plain text
    portions (before/after/between URLs) are also parsed.

    Mirrors fwlog's extract_items_from_text_chunk faithfully, with extended
    group_id support for bridge_file / bridge_file_name resolution."""
    text = str(text or "")
    if not text.strip():
        return []

    items = []
    plain_parts = []
    cursor = 0
    url_index = 0

    for match in URL_RE.finditer(text):
        plain_parts.append(text[cursor:match.start()])
        url = match.group(0)
        target = parse_log_target_entry(url)
        extracted_text = ""
        if target:
            try:
                extracted_text = await asyncio.to_thread(
                    fetch_log_text_by_source,
                    target["key"],
                    target.get("password", ""),
                    target.get("source"),
                    group_id,
                )
            except Exception:
                extracted_text = ""

        if extracted_text and extracted_text.strip():
            plain_text = "".join(plain_parts).strip()
            if plain_text:
                items.extend(
                    parse_structured_text_to_items(
                        plain_text,
                        sender_name,
                        sender_id,
                        ts,
                        f"{raw_msg_id}:text:{url_index}",
                    )
                )
            plain_parts = []
            items.extend(
                parse_structured_text_to_items(
                    extracted_text,
                    sender_name,
                    sender_id,
                    ts,
                    f"{raw_msg_id}:url:{url_index}",
                )
            )
            url_index += 1
        else:
            plain_parts.append(url)

        cursor = match.end()

    plain_parts.append(text[cursor:])
    plain_text = "".join(plain_parts).strip()
    if plain_text:
        items.extend(
            parse_structured_text_to_items(
                plain_text,
                sender_name,
                sender_id,
                ts,
                f"{raw_msg_id}:text:tail",
            )
        )

    return items


# ====== 文件去重 ======

def cleanup_recent_file_captures(now=None):
    """清理过期的文件捕获记录"""
    now = now or time.time()
    expired = [key for key, value in RECENT_FILE_CAPTURES.items() if value[1] <= now]
    for key in expired:
        RECENT_FILE_CAPTURES.pop(key, None)


def remember_file_capture(session_id, log_name, event_ts):
    """
    文件去重检查。同一 session+log 在 1 秒内不重复记录。
    返回 True 表示首次捕获，False 表示重复。
    """
    now = time.time()
    cleanup_recent_file_captures(now)
    cache_key = (str(session_id), str(log_name), "file-upload-window")
    current_event_ts = safe_int(event_ts, int(now))
    previous = RECENT_FILE_CAPTURES.get(cache_key)
    if previous:
        previous_event_ts, _ = previous
        if abs(current_event_ts - previous_event_ts) <= RECENT_FILE_DEDUPE_WINDOW_SEC:
            return False
    RECENT_FILE_CAPTURES[cache_key] = (current_event_ts, now + RECENT_FILE_CAPTURE_TTL_SEC)
    return True


# ====== fwlog 移植：文件提取辅助函数 ======

def extract_file_payloads(event):
    """Extract file payloads from various event shapes.
    Mirrors fwlog's extract_file_payloads exactly.
    Handles: notice.group_upload, message-segment type=file, CQ:file string."""
    payloads = []
    post_type = str(event.get('post_type') or '').lower()
    notice_type = str(event.get('notice_type') or '').lower()

    if post_type == 'notice' and notice_type == 'group_upload':
        file_info = event.get('file') or {}
        file_id = file_info.get('id') or file_info.get('file_id')
        if file_id:
            payloads.append({
                'group_id': str(event.get('group_id') or ''),
                'user_id': str(event.get('user_id') or ''),
                'file_id': str(file_id),
                'name': str(file_info.get('name') or file_info.get('file_name') or '未知文件'),
                'busid': safe_int(file_info.get('busid'), 0),
                'url': str(file_info.get('url') or ''),
            })
        return payloads

    message = event.get('message')
    if isinstance(message, list):
        for seg in message:
            if not isinstance(seg, dict) or str(seg.get('type') or '').lower() != 'file':
                continue
            data = seg.get('data') or {}
            file_id = data.get('id') or data.get('file_id')
            if not file_id:
                continue
            payloads.append({
                'group_id': str(event.get('group_id') or ''),
                'user_id': str(event.get('user_id') or ''),
                'file_id': str(file_id),
                'name': str(data.get('name') or data.get('file') or '未知文件'),
                'busid': safe_int(data.get('busid'), 0),
                'url': str(data.get('url') or data.get('file_url') or ''),
            })
        return payloads

    if isinstance(message, str) and '[CQ:file' in message:
        idx = 0
        while True:
            start = message.find('[CQ:file', idx)
            if start == -1:
                break
            end = message.find(']', start)
            if end == -1:
                break
            segment = message[start + 1:end]
            param_text = segment.split(',', 1)[1] if ',' in segment else ''
            params = parse_cq_params(param_text)
            file_id = params.get('id') or params.get('file_id')
            if file_id:
                payloads.append({
                    'group_id': str(event.get('group_id') or ''),
                    'user_id': str(event.get('user_id') or ''),
                    'file_id': str(file_id),
                    'name': str(params.get('name') or params.get('file') or '未知文件'),
                    'busid': safe_int(params.get('busid'), 0),
                    'url': str(params.get('url') or params.get('file_url') or ''),
                })
            idx = end + 1

    return payloads


def resolve_group_file_url_napcat(group_id, file_id, busid):
    """Resolve a group file download URL via NapCat HTTP API.
    Used when NC_FILE_BRIDGE_MODE != 0 (i.e. polling mode or no WS)."""
    try:
        return get_group_file_url(group_id, file_id, busid)
    except Exception:
        raise RuntimeError('未获取到文件下载地址')


async def resolve_group_file_url_ws(group_id, file_id, busid):
    """Resolve a group file download URL via WebSocket API (fwlog-style).
    Used when NC_FILE_BRIDGE_MODE == 0 and WS_CLIENT is connected."""
    global WS_CLIENT
    if WS_CLIENT is None or WS_CLIENT.ws_conn is None:
        raise RuntimeError('[logutil-ws] WebSocket 未连接，无法获取文件下载地址')
    response = await WS_CLIENT.send_api(
        "get_group_file_url",
        {
            "group_id": str(group_id),
            "file_id": str(file_id),
            "busid": safe_int(busid, 0),
        },
    )
    data = (response.get("data") or {}) if isinstance(response, dict) else {}
    url = str(data.get("url") or "")
    if not url:
        raise RuntimeError('未获取到文件下载地址')
    return url


async def extract_items_from_file_payload_napcat(payload, sender_name, sender_id, ts, raw_msg_id):
    """Download a file via NapCat, extract its text, and parse into log items.
    Mirrors fwlog's extract_items_from_file_payload.
    Uses WebSocket get_group_file_url when NC_FILE_BRIDGE_MODE == 0 (fwlog-style),
    falls back to HTTP API otherwise."""
    file_url = str(payload.get('url') or '')
    if not file_url:
        group_id = payload.get('group_id')
        file_id = payload.get('file_id')
        if group_id and file_id:
            try:
                if NC_FILE_BRIDGE_MODE == 0 and WS_CLIENT is not None:
                    # fwlog-style: resolve via WebSocket
                    file_url = await resolve_group_file_url_ws(
                        group_id, file_id, payload.get('busid', 0),
                    )
                else:
                    file_url = await asyncio.to_thread(
                        resolve_group_file_url_napcat,
                        group_id,
                        file_id,
                        payload.get('busid', 0),
                    )
            except Exception:
                raise RuntimeError('未获取到文件下载地址')

    if not file_url:
        raise RuntimeError('未获取到文件下载地址')

    file_bytes = await asyncio.to_thread(download_file_bytes, file_url)
    extracted_text = await asyncio.to_thread(
        extract_text_from_group_file,
        payload.get('name') or '未知文件',
        file_bytes,
    )
    return parse_structured_text_to_items(
        extracted_text, sender_name, sender_id, ts, raw_msg_id,
    )




# ====== 合并转发展开 ======

def extract_forward_ids_from_text(text):
    """从文本中提取 [CQ:forward,id=...] 的转发 ID 列表"""
    ids = []
    if not text:
        return ids
    prefix = "[CQ:forward"
    if prefix in text:
        idx = 0
        while True:
            start = text.find(prefix, idx)
            if start == -1:
                break
            end = text.find("]", start)
            if end == -1:
                break
            segment = text[start:end]
            value = ""
            for key in ["id=", "res_id=", "message_id="]:
                pos = segment.find(key)
                if pos != -1:
                    pos += len(key)
                    j = pos
                    while j < len(segment) and segment[j] not in ",]":
                        j += 1
                    value = segment[pos:j]
                    if value:
                        break
            if value:
                ids.append(value)
            idx = end + 1
    return ids


def fetch_forward_messages_napcat(forward_id):
    """
    通过 NapCat HTTP API 获取合并转发消息的原始内容。
    返回 nodes 列表，每个 node 含 sender, time, message 等字段。
    """
    try:
        response = napcat_json_post(
            "get_forward_msg",
            {"id": str(forward_id)},
            DOWNLOAD_TIMEOUT_SEC,
        )
    except Exception:
        # 用 message_id 重试
        try:
            response = napcat_json_post(
                "get_forward_msg",
                {"message_id": str(forward_id)},
                DOWNLOAD_TIMEOUT_SEC,
            )
        except Exception:
            return []

    data = response.get("data")
    if response.get("status") != "ok" or not data:
        return []

    if isinstance(data, dict) and "messages" in data:
        nodes = data["messages"]
    elif isinstance(data, list):
        nodes = data
    else:
        nodes = []
    return nodes


async def extract_items_from_forward(forward_id, group_id, fallback_name, fallback_user_id, ts, raw_msg_id):
    """Expand a merged-forward message via NapCat HTTP API.
    Each inner node is converted to plain text and then parsed via
    parse_structured_text_to_items (to handle embedded URLs, etc.).

    Mirrors fwlog's extract_items_from_forward, adapted for HTTP-based
    NapCat access instead of WebSocket."""
    nodes = await asyncio.to_thread(fetch_forward_messages_napcat, forward_id)
    if not nodes:
        return []

    items = []
    for index, node in enumerate(nodes):
        if not isinstance(node, dict):
            continue
        sender = node.get("sender") or {}
        node_sender_id = str(sender.get("user_id") or "")
        node_sender_name = sender.get("card") or sender.get("nickname") or (f"QQ:{node_sender_id}" if node_sender_id else fallback_name)
        node_ts = safe_int(node.get("time"), safe_int(ts, int(time.time())))
        node_message = node.get("message") or node.get("content") or ""

        # Convert message segments to plain text using shared helper
        node_text = segments_to_text(node_message)

        if not node_text.strip():
            continue

        node_raw_id = str(node.get("message_id") or f"forward:{forward_id}:{index}")
        node_items = parse_structured_text_to_items(
            node_text, node_sender_name, node_sender_id, node_ts, node_raw_id
        )
        items.extend(node_items)

    return items


def get_logutil_db_connection():
    db_dir = os.path.dirname(LOGUTIL_DB_FILE)
    if db_dir:
        os.makedirs(db_dir, exist_ok=True)
    conn = sqlite3.connect(LOGUTIL_DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_logutil_db():
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS groups (
            group_id TEXT PRIMARY KEY,
            current_log_name TEXT,
            recording INTEGER DEFAULT 0,
            baseline_file_id TEXT DEFAULT '',
            created_at INTEGER,
            updated_at INTEGER
        )
    ''')
    # Migration: add baseline_file_id if column missing (upgrade from older schema)
    try:
        c.execute('ALTER TABLE groups ADD COLUMN baseline_file_id TEXT DEFAULT \'\'')
    except sqlite3.OperationalError:
        pass  # column already exists
    # v4.3.1: add raw_recording column
    try:
        c.execute('ALTER TABLE groups ADD COLUMN raw_recording INTEGER DEFAULT 0')
    except sqlite3.OperationalError:
        pass
    c.execute('''
        CREATE TABLE IF NOT EXISTS logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            group_id TEXT,
            name TEXT,
            ended INTEGER DEFAULT 0,
            created_at INTEGER,
            updated_at INTEGER,
            upload_url TEXT,
            upload_time INTEGER DEFAULT 0,
            UNIQUE(group_id, name)
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            log_id INTEGER,
            nickname TEXT,
            im_userid TEXT,
            time INTEGER,
            message TEXT,
            raw_msg_id TEXT,
            FOREIGN KEY(log_id) REFERENCES logs(id)
        )
    ''')
    conn.commit()
    conn.close()


def ensure_logutil_group_state(group_id):
    if group_id is None:
        group_id = ''
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT * FROM groups WHERE group_id = ?', (str(group_id),))
    row = c.fetchone()
    if not row:
        now = int(time.time() * 1000)
        c.execute(
            '''
                INSERT INTO groups (group_id, current_log_name, recording, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?)
            ''',
            (str(group_id), '', 0, now, now),
        )
        conn.commit()
        c.execute('SELECT * FROM groups WHERE group_id = ?', (str(group_id),))
        row = c.fetchone()
    conn.close()
    return dict(row)


def update_logutil_group_state(group_id, **kwargs):
    if group_id is None:
        group_id = ''
    conn = get_logutil_db_connection()
    c = conn.cursor()
    updates = []
    values = []
    for k, v in kwargs.items():
        updates.append(f"{k} = ?")
        values.append(v)
    values.append(str(group_id))
    sql = f"UPDATE groups SET {', '.join(updates)} WHERE group_id = ?"
    c.execute(sql, values)
    conn.commit()
    conn.close()


def ensure_logutil_log(group_id, name):
    if group_id is None:
        group_id = ''
    name = str(name or '').strip() or generate_default_log_name()
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT * FROM logs WHERE group_id = ? AND name = ?', (str(group_id), name))
    row = c.fetchone()
    if not row:
        now = int(time.time() * 1000)
        c.execute(
            '''
                INSERT INTO logs (group_id, name, ended, created_at, updated_at)
                VALUES (?, ?, 0, ?, ?)
            ''',
            (str(group_id), name, now, now),
        )
        conn.commit()
        c.execute('SELECT * FROM logs WHERE group_id = ? AND name = ?', (str(group_id), name))
        row = c.fetchone()
    conn.close()
    return dict(row)


def update_logutil_log_meta(log_id, **kwargs):
    conn = get_logutil_db_connection()
    c = conn.cursor()
    updates = []
    values = []
    for k, v in kwargs.items():
        updates.append(f"{k} = ?")
        values.append(v)
    values.append(log_id)
    sql = f"UPDATE logs SET {', '.join(updates)} WHERE id = ?"
    c.execute(sql, values)
    conn.commit()
    conn.close()


def add_logutil_items(log_id, items):
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT COUNT(*) FROM items WHERE log_id = ?', (log_id,))
    old_count = c.fetchone()[0]
    for item in items:
        c.execute(
            '''
                INSERT INTO items (log_id, nickname, im_userid, time, message, raw_msg_id)
                VALUES (?, ?, ?, ?, ?, ?)
            ''',
            (
                log_id,
                str(item.get('nickname') or ''),
                str(item.get('im_userid') or ''),
                safe_int(item.get('time'), int(time.time())),
                str(item.get('message') or ''),
                str(item.get('raw_msg_id') or ''),
            ),
        )
    now = int(time.time() * 1000)
    c.execute('UPDATE logs SET updated_at = ? WHERE id = ?', (now, log_id))
    conn.commit()
    conn.close()
    return old_count, old_count + len(items)


def clear_logutil_items(log_id):
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('DELETE FROM items WHERE log_id = ?', (log_id,))
    conn.commit()
    conn.close()


def get_logutil_log_full(group_id, name):
    if group_id is None:
        group_id = ''
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT * FROM logs WHERE group_id = ? AND name = ?', (str(group_id), str(name)))
    row = c.fetchone()
    if not row:
        conn.close()
        return None
    log_obj = dict(row)
    c.execute('SELECT * FROM items WHERE log_id = ? ORDER BY id', (log_obj['id'],))
    log_obj['items'] = [dict(item) for item in c.fetchall()]
    conn.close()
    return log_obj


def get_logutil_logs_list(group_id):
    if group_id is None:
        group_id = ''
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT * FROM logs WHERE group_id = ? ORDER BY created_at DESC', (str(group_id),))
    logs = [dict(row) for row in c.fetchall()]
    for l in logs:
        c.execute('SELECT COUNT(*) FROM items WHERE log_id = ?', (l['id'],))
        l['item_count'] = c.fetchone()[0]
    conn.close()
    return logs


def delete_logutil_log(group_id, name):
    if group_id is None:
        group_id = ''
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT id FROM logs WHERE group_id = ? AND name = ?', (str(group_id), str(name)))
    row = c.fetchone()
    if row:
        log_id = row['id']
        c.execute('DELETE FROM items WHERE log_id = ?', (log_id,))
        c.execute('DELETE FROM logs WHERE id = ?', (log_id,))
        conn.commit()
        # Clear imported files tracking for this log
        with STATE_LOCK:
            LOG_IMPORTED_FILES.pop(str(log_id), None)
    conn.close()


def generate_log_text(log_obj):
    items = log_obj.get('items', [])
    blocks = []
    for item in items:
        ts = item.get('time', 0)
        dt = format_time(ts)
        name = item.get('nickname', 'Unknown')
        uid = item.get('im_userid', '')
        msg = item.get('message', '')
        header = f"{name}({uid}) {dt}" if uid else f"{name} {dt}"
        if msg is None:
            msg = ''
        msg = str(msg)
        content_lines = [f" {line}" for line in msg.splitlines()]
        content_text = '\n'.join(content_lines)
        blocks.append(f"{header}\n{content_text}")
    return '\n\n'.join(blocks)


def normalize_uniform_id(value, fallback_prefix="LOGAI"):
    """标准化 uniform_id 格式为 prefix:id (如 QQ:123456)。
    与 fwlog 的 normalize_uniform_id 完全一致。"""
    text = str(value or "").strip()
    if not text:
        return f"{fallback_prefix}:unknown"
    if re.fullmatch(r"[^:]+:\S+", text):
        return text
    prefix = "QQ" if re.fullmatch(r"\d+", text) else fallback_prefix
    return f"{prefix}:{text}"


def build_story_painter_items(log_obj):
    items = []
    for index, item in enumerate(log_obj.get('items', []), start=1):
        im_userid = str(item.get('im_userid') or '')
        items.append(
            {
                'id': index,
                'nickname': str(item.get('nickname') or ''),
                'IMUserId': im_userid,
                'time': safe_int(item.get('time'), 0),
                'message': str(item.get('message') or ''),
                'isDice': False,
                'commandId': 0,
                'commandInfo': None,
                'rawMsgId': str(item.get('raw_msg_id') or ''),
                'uniformId': normalize_uniform_id(item.get('uniform_id') or im_userid),
                'channel': '',
            }
        )
    return items


def upload_log_to_story_painter(log_obj, uniform_id):
    upload_url = str(log_obj.get('upload_url') or '')
    upload_time = safe_int(log_obj.get('upload_time'), 0)
    update_time = safe_int(log_obj.get('updated_at'), 0)
    if upload_url and upload_time > update_time:
        return upload_url, True
    payload = {
        'version': STORY_PAINTER_VERSION,
        'items': build_story_painter_items(log_obj),
    }
    encoded_payload = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    compressed_payload = zlib.compress(encoded_payload)
    headers = {}
    if STORY_PAINTER_TOKEN:
        headers['Authorization'] = f"Bearer {STORY_PAINTER_TOKEN}"
    session = requests.Session()
    session.trust_env = False
    response = session.put(
        STORY_PAINTER_UPLOAD_URL,
        data={
            'name': str(log_obj.get('name') or ''),
            'uniform_id': normalize_uniform_id(uniform_id or log_obj.get('group_id') or 'fwlog'),
            'client': STORY_PAINTER_CLIENT,
            'version': str(STORY_PAINTER_VERSION),
        },
        files={
            'file': ('log-zlib-compressed', compressed_payload, 'application/octet-stream'),
        },
        headers=headers,
        timeout=STORY_PAINTER_TIMEOUT_SEC,
    )
    response.raise_for_status()
    try:
        result = response.json()
    except Exception:
        result = {}
    url = str(result.get('url') or '')
    if not url:
        raise RuntimeError(f"染色器返回异常: {response.text[:200]}")
    upload_ts = int(time.time() * 1000)
    update_logutil_log_meta(log_obj['id'], upload_url=url, upload_time=upload_ts)
    log_obj['upload_url'] = url
    log_obj['upload_time'] = upload_ts
    return url, False


def auto_import_bridge_file_to_logutil(bridge_item):
    """If logutil is recording for this group, auto-add bridge file text as log items."""
    group_id = bridge_item.get('group_id', 0)
    if not group_id:
        return
    group_id_str = str(group_id)
    state = ensure_logutil_group_state(group_id_str)
    if not state.get('recording'):
        return
    current_name = state.get('current_log_name')
    if not current_name:
        return

    file_id = str(bridge_item.get('file_id', ''))

    # Basline check: skip if this file is the same one that existed when recording started
    baseline_file_id = str(state.get('baseline_file_id', '')).strip()
    if baseline_file_id and file_id and baseline_file_id == file_id:
        bridge_log("autoimport", f"baseline skipped: file_id={file_id} matches baseline for log={current_name} group={group_id_str}")
        return

    # Get or create the log object
    log_obj = ensure_logutil_log(group_id_str, current_name)
    log_id_str = str(log_obj['id'])

    # Dedup layer 1: in-memory set (fast, no DB hit)
    if file_id:
        with STATE_LOCK:
            imported_set = LOG_IMPORTED_FILES.get(log_id_str, set())
            if file_id in imported_set:
                bridge_log("autoimport", f"dedup skipped (mem): file_id={file_id} already in log={current_name}")
                return

    # Dedup layer 2: DB check with unified prefix
    if file_id:
        conn = get_logutil_db_connection()
        c = conn.cursor()
        import_pattern = f"file:{file_id}%"
        c.execute('SELECT COUNT(*) FROM items WHERE log_id = ? AND raw_msg_id LIKE ?', (log_obj['id'], import_pattern))
        existing_count = c.fetchone()[0]
        conn.close()
        if existing_count > 0:
            bridge_log("autoimport", f"dedup skipped (db): file_id={file_id} already imported into log={current_name} group={group_id_str} ({existing_count} items exist)")
            return

    ck = bridge_item.get('content_key', '')
    with STATE_LOCK:
        path = CONTENT_INDEX.get(ck, '')
    if not path or not os.path.exists(path):
        return
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if not text.strip():
        return

    file_name = bridge_item.get('name', 'BridgeFile')
    sender_name = f"[BridgeFile] {file_name}"
    file_ts = bridge_item.get('ts', int(time.time()))
    raw_msg_id = f"file:{bridge_item.get('file_id', '')}"

    items = parse_structured_text_to_items(
        text,
        sender_name=sender_name,
        sender_id=str(bridge_item.get('user_id', '')),
        ts=file_ts,
        raw_msg_id=raw_msg_id,
    )
    if items:
        old_count, new_count = add_logutil_items(log_obj['id'], items)
        # Mark as imported
        if file_id:
            with STATE_LOCK:
                if log_id_str not in LOG_IMPORTED_FILES:
                    LOG_IMPORTED_FILES[log_id_str] = set()
                LOG_IMPORTED_FILES[log_id_str].add(file_id)
        bridge_log("autoimport", f"added {len(items)} items to logutil log={current_name} group={group_id} (total={new_count})")


def capture_baseline_file_id(group_id):
    """Capture the current latest cached file's file_id as baseline for the group.
    Returns the baseline file_id string, or empty string if no cached file exists."""
    group_id_int = safe_int(group_id, 0)
    baseline = ''
    if group_id_int > 0:
        with STATE_LOCK:
            file_list = LATEST_FILES.get(group_id_int, [])
        if file_list:
            baseline = str(file_list[-1].get('file_id', ''))
    update_logutil_group_state(str(group_id), baseline_file_id=baseline)
    if baseline:
        bridge_log("autoimport", f"baseline set: group={group_id} baseline_file_id={baseline}")
    return baseline


def generate_default_log_name():
    return f"fwlog_{time.strftime('%Y%m%d_%H%M%S', time.localtime())}"

def pad2(n):
    """将数字格式化为两位字符串，不足两位前面补零"""
    return f"{n:02d}"


def format_time(ts):
    try:
        t = int(ts)
    except Exception:
        t = int(time.time())
    # 如果是毫秒单位，则转换为秒
    if t > 9999999999:
        t = t // 1000
    d = time.localtime(t)
    y = d.tm_year
    m = pad2(d.tm_mon)
    day = pad2(d.tm_mday)
    hh = pad2(d.tm_hour)
    mm = pad2(d.tm_min)
    ss = pad2(d.tm_sec)
    return f"{y}/{m}/{day} {hh}:{mm}:{ss}"


def export_log_text(log_obj, del_paren=False):
    if not log_obj or not log_obj.get('items'):
        return None, ''
    full_text = generate_log_text(log_obj)
    if del_paren:
        full_text = strip_paren_text(full_text)
    ensure_bridge_cache_dir()
    key = uuid.uuid4().hex
    base_name = os.path.splitext(str(log_obj.get('name') or 'log'))[0] or 'log'
    text_filename = f"{base_name}.txt"
    out_path = os.path.join(BRIDGE_CACHE_DIR, f"{key}.txt")
    with open(out_path, 'w', encoding='utf-8') as fw:
        fw.write(full_text)
    public_base = (request.host_url or '').rstrip('/') if request else resolve_public_base('')
    content_url = build_content_url(key, public_base=public_base)
    cached_ts = int(time.time())
    group_id_str = str(log_obj.get('group_id') or '').strip()
    group_id_int = safe_int(group_id_str, 0) if group_id_str else 0
    with STATE_LOCK:
        CONTENT_INDEX[key] = out_path
    # v4.4.0: export_log_text 不再注册到 LATEST_FILES 中（避免与 send_log_via_napcat 重复）
    return text_filename, content_url


# Initialize persistent logutil state storage
init_logutil_db()


def _update_runtime_config(key, value):
    """Update a module-level runtime config variable.  Safe whitelist only."""
    allowed = {
        'NAPCAT_WS_URL',
        'NAPCAT_WS_TOKEN',
        'LOGUTIL_WS_ENABLED',
        'NC_FILE_BRIDGE_MODE',
    }
    if key not in allowed:
        return False
    import builtins
    mod = builtins.__import__(__name__)
    setattr(mod, key, value)
    # If WS was enabled/disabled, restart the worker
    if key == 'LOGUTIL_WS_ENABLED':
        global WS_WORKER
        if value and (WS_WORKER is None or not WS_WORKER.is_alive()):
            ensure_ws_worker_started()
    return True


@app.route('/api/logutil_config', methods=['POST'])
def api_logutil_config():
    """Set runtime WS config from frontend."""
    payload = request.get_json(silent=True) or {}
    updated = {}
    for key in ('NAPCAT_WS_URL', 'NAPCAT_WS_TOKEN', 'LOGUTIL_WS_ENABLED', 'NC_FILE_BRIDGE_MODE'):
        if key in payload:
            val = payload[key]
            if key == 'LOGUTIL_WS_ENABLED':
                val = bool(val)
            elif key == 'NC_FILE_BRIDGE_MODE':
                val = int(val)
            else:
                val = str(val) if val else ''
            if _update_runtime_config(key, val):
                updated[key] = val
    current = {
        'NAPCAT_WS_URL': str(getattr(__import__(__name__), 'NAPCAT_WS_URL', '')),
        'NAPCAT_WS_TOKEN': '***' if NAPCAT_WS_TOKEN else '',
        'LOGUTIL_WS_ENABLED': bool(LOGUTIL_WS_ENABLED),
        'NC_FILE_BRIDGE_MODE': int(NC_FILE_BRIDGE_MODE),
    }
    return jsonify({'status': 'ok', 'updated': updated, 'current': current})


@app.route('/api/logutil_config', methods=['GET'])
def api_logutil_config_get():
    """Return current WS config (token masked)."""
    return jsonify({
        'status': 'ok',
        'NAPCAT_WS_URL': NAPCAT_WS_URL,
        'NAPCAT_WS_TOKEN': '***' if NAPCAT_WS_TOKEN else '',
        'LOGUTIL_WS_ENABLED': bool(LOGUTIL_WS_ENABLED),
        'NC_FILE_BRIDGE_MODE': int(NC_FILE_BRIDGE_MODE),
        'MAX_BRIDGE_FILES_PER_GROUP': int(MAX_BRIDGE_FILES_PER_GROUP),
    })


@app.route('/api/logutil_toggle', methods=['POST'])
def api_logutil_toggle():
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id', '')).strip()
    action = str(payload.get('action') or request.args.get('action') or '').lower()
    if not group_id or action not in ('on', 'off'):
        return jsonify({'status': 'error', 'msg': 'invalid params'}), 400
    state = ensure_logutil_group_state(group_id)
    if action == 'on':
        name = payload.get('name') or state.get('current_log_name') or generate_default_log_name()
        log_obj = ensure_logutil_log(group_id, name)
        update_logutil_group_state(group_id, current_log_name=log_obj['name'], recording=1)
        update_logutil_log_meta(log_obj['id'], ended=0)
    else:
        update_logutil_group_state(group_id, recording=0)
    state = ensure_logutil_group_state(group_id)
    logs = get_logutil_logs_list(group_id)
    current_log = next((l for l in logs if l['name'] == state.get('current_log_name')), None)
    saved = current_log['item_count'] if current_log else 0
    return jsonify({'status': 'ok', 'group_id': group_id, 'enabled': bool(state.get('recording')), 'saved': saved})


@app.route('/api/logutil_status', methods=['GET'])
def api_logutil_status():
    group_id = str(request.args.get('group_id') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    logs = get_logutil_logs_list(group_id)
    current_log = next((l for l in logs if l['name'] == state.get('current_log_name')), None)
    saved = current_log['item_count'] if current_log else 0
    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'enabled': bool(state.get('recording')),
        'current_log_name': state.get('current_log_name'),
        'saved': saved,
    })


@app.route('/api/logutil_new', methods=['POST'])
def api_logutil_new():
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    name = str(payload.get('name') or '').strip() or generate_default_log_name()
    raw_mode = str(payload.get('raw') or '').lower() in ('1', 'true', 'yes', 'on')
    print(f"[logutil_new] payload={payload} group_id={group_id!r} name={name!r}")
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    gid_int = safe_int(group_id, 0)
    if gid_int > 0:
        BRIDGE_POLL_GROUPS.add(gid_int)
        ensure_poll_worker_started()
    log_obj = ensure_logutil_log(group_id, name)
    clear_logutil_items(log_obj['id'])
    # Clear imported file tracking for this log (fresh start)
    log_id_str = str(log_obj['id'])
    with STATE_LOCK:
        LOG_IMPORTED_FILES.pop(log_id_str, None)
    update_logutil_log_meta(log_obj['id'], ended=0)
    update_logutil_group_state(group_id, current_log_name=log_obj['name'], recording=1,
                                raw_recording=1 if raw_mode else 0)
    # Capture the currently latest file as baseline — it won't be auto-imported
    # until a different (newer) file appears
    capture_baseline_file_id(group_id)
    print(f"[logutil_new] created log={log_obj['name']!r} raw={raw_mode} current_log_name set for group={group_id}")
    return jsonify({'status': 'ok', 'group_id': group_id, 'name': log_obj['name'], 'saved': 0})


@app.route('/api/logutil_on', methods=['POST'])
def api_logutil_on():
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    name = str(payload.get('name') or '').strip()
    raw_mode = str(payload.get('raw') or '').lower() in ('1', 'true', 'yes', 'on')
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    gid_int = safe_int(group_id, 0)
    if gid_int > 0:
        BRIDGE_POLL_GROUPS.add(gid_int)
        ensure_poll_worker_started()
    state = ensure_logutil_group_state(group_id)
    if not name:
        name = state.get('current_log_name') or generate_default_log_name()
    log_obj = ensure_logutil_log(group_id, name)
    update_logutil_group_state(group_id, current_log_name=log_obj['name'], recording=1,
                                raw_recording=1 if raw_mode else 0)
    update_logutil_log_meta(log_obj['id'], ended=0)
    # Capture the currently latest file as baseline — it won't be auto-imported
    # until a different (newer) file appears
    capture_baseline_file_id(group_id)
    conn = get_logutil_db_connection()
    c = conn.cursor()
    c.execute('SELECT COUNT(*) FROM items WHERE log_id = ?', (log_obj['id'],))
    saved = c.fetchone()[0]
    conn.close()
    return jsonify({'status': 'ok', 'group_id': group_id, 'name': log_obj['name'], 'saved': saved})


@app.route('/api/logutil_off', methods=['POST'])
def api_logutil_off():
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    update_logutil_group_state(group_id, recording=0, baseline_file_id='')
    current_name = state.get('current_log_name')
    saved = 0
    if current_name:
        logs = get_logutil_logs_list(group_id)
        current = next((l for l in logs if l['name'] == current_name), None)
        saved = current['item_count'] if current else 0
    return jsonify({'status': 'ok', 'group_id': group_id, 'enabled': False, 'saved': saved})


@app.route('/api/logutil_list', methods=['GET'])
def api_logutil_list():
    group_id = str(request.args.get('group_id') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    logs = get_logutil_logs_list(group_id)
    return jsonify({'status': 'ok', 'group_id': group_id, 'current_log_name': state.get('current_log_name'), 'recording': bool(state.get('recording')), 'logs': logs})


@app.route('/api/logutil_clear', methods=['POST'])
def api_logutil_clear():
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    name = str(payload.get('name') or request.args.get('name') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    if not name:
        name = state.get('current_log_name')
    if not name:
        return jsonify({'status': 'error', 'msg': 'missing log name'}), 400
    delete_logutil_log(group_id, name)
    if state.get('current_log_name') == name:
        update_logutil_group_state(group_id, current_log_name='', recording=0)
    return jsonify({'status': 'ok', 'group_id': group_id, 'name': name})


@app.route('/api/logutil_add', methods=['POST'])
def api_logutil_add():
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    name = str(payload.get('name') or state.get('current_log_name') or '').strip()
    if not name:
        return jsonify({'status': 'error', 'msg': 'missing log name'}), 400
    log_obj = ensure_logutil_log(group_id, name)
    items_data = payload.get('items') or []
    items = []
    if isinstance(items_data, list) and items_data:
        for item in items_data:
            items.append({
                'nickname': item.get('nickname') or item.get('name') or 'Unknown',
                'im_userid': item.get('im_userid') or item.get('user_id') or '',
                'time': safe_int(item.get('time'), int(time.time())),
                'message': item.get('message') or item.get('text') or '',
                'raw_msg_id': item.get('raw_msg_id') or item.get('id') or '',
            })
    else:
        sender_name = str(payload.get('nickname') or payload.get('sender_name') or 'Unknown')
        sender_id = str(payload.get('im_userid') or payload.get('user_id') or '')
        raw_msg_id = str(payload.get('raw_msg_id') or f"manual-{int(time.time())}")
        ts = safe_int(payload.get('time'), int(time.time()))
        message_array = payload.get('message')

        # --- Array segment 格式 (OneBot v11 standard) ---
        if isinstance(message_array, list):
            for index, seg in enumerate(message_array):
                if not isinstance(seg, dict):
                    continue
                seg_type = str(seg.get('type') or '').lower()
                data = seg.get('data') or {}
                if seg_type == 'text':
                    seg_text = str(data.get('text') or '')
                    if seg_text.strip():
                        text_items = asyncio.run(
                            extract_items_from_text_chunk(
                                seg_text, sender_name, sender_id, ts,
                                f"{raw_msg_id}:text:{index}", group_id
                            )
                        )
                        items.extend(text_items)
                elif seg_type == 'forward':
                    forward_id = data.get('id') or data.get('res_id') or data.get('message_id')
                    if forward_id:
                        try:
                            fwd_items = asyncio.run(
                                extract_items_from_forward(
                                    str(forward_id), group_id, sender_name, sender_id,
                                    ts, f"{raw_msg_id}:fwd:{forward_id}"
                                )
                            )
                            items.extend(fwd_items)
                        except Exception:
                            pass
                elif seg_type == 'file':
                    payloads = extract_file_payloads({
                        'post_type': 'message',
                        'message': [seg],
                        'group_id': group_id,
                        'user_id': sender_id,
                    })
                    for file_payload in payloads:
                        try:
                            file_items = asyncio.run(
                                extract_items_from_file_payload_napcat(
                                    file_payload, sender_name, sender_id, ts,
                                    f"{raw_msg_id}:file:{file_payload.get('file_id') or f'message-{index}'}"
                                )
                            )
                            items.extend(file_items)
                        except Exception:
                            pass
        else:
            text = str(payload.get('text') or '').strip()
            if not text:
                return jsonify({'status': 'error', 'msg': 'missing text or items'}), 400

            # --- 合并转发展开 (CQ code string) ---
            forward_ids = extract_forward_ids_from_text(text)
            for fid in forward_ids:
                try:
                    fwd_items = asyncio.run(
                        extract_items_from_forward(fid, group_id, sender_name, sender_id, ts, f"{raw_msg_id}:fwd:{fid}")
                    )
                    items.extend(fwd_items)
                except Exception:
                    pass

            # --- 清除 CQ 标签后继续文本解析 ---
            cleaned_text = re.sub(r"\[CQ:(?:forward|file)[^\]]*\]", "", text).strip()
            if cleaned_text:
                text_items = asyncio.run(
                    extract_items_from_text_chunk(cleaned_text, sender_name, sender_id, ts, raw_msg_id, group_id)
                )
                items.extend(text_items)
    if not items:
        return jsonify({'status': 'error', 'msg': 'no items parsed'}), 400
    old_count, new_count = add_logutil_items(log_obj['id'], items)
    milestone = new_count // 1000 > old_count // 1000
    result = {
        'status': 'ok',
        'group_id': group_id,
        'name': log_obj['name'],
        'saved': new_count,
        'added': len(items),
        'milestone': milestone,
    }
    if milestone:
        result['milestone_text'] = f"当前日志 {log_obj['name']} 已记录 {new_count} 条消息。"
    return jsonify(result)


@app.route('/api/logutil_compound', methods=['POST'])
def api_logutil_compound():
    """Atomically execute a compound logutil command:
    [new [title]] <op1> <op2> ... [end] [logai]
    All ops resolved against current bridge cache state before any modification."""
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400

    do_new = bool(payload.get('new'))
    do_end = bool(payload.get('end'))
    do_logai = bool(payload.get('logai'))
    raw_mode = bool(payload.get('raw'))
    title = str(payload.get('title') or '').strip()
    ops = payload.get('ops') or []
    if not isinstance(ops, list):
        ops = []

    results = {'ops': [], 'total_added': 0}

    # Step 1: start new recording if requested
    if do_new:
        name = title or generate_default_log_name()
        log_obj = ensure_logutil_log(group_id, name)
        clear_logutil_items(log_obj['id'])
        log_id_str = str(log_obj['id'])
        with STATE_LOCK:
            LOG_IMPORTED_FILES.pop(log_id_str, None)
        update_logutil_log_meta(log_obj['id'], ended=0)
        ensure_logutil_group_state(group_id)  # ensure row exists before UPDATE
        update_logutil_group_state(group_id, current_log_name=log_obj['name'], recording=1,
                                   raw_recording=1 if raw_mode else 0)
        capture_baseline_file_id(group_id)
        # Register for polling
        gid_int = safe_int(group_id, 0)
        if gid_int > 0:
            BRIDGE_POLL_GROUPS.add(gid_int)
            ensure_poll_worker_started()
        results['name'] = log_obj['name']
    else:
        # Use current recording, or start unnamed if none
        state = ensure_logutil_group_state(group_id)
        current_name = state.get('current_log_name') or ''
        if state.get('recording') and current_name:
            log_obj = ensure_logutil_log(group_id, current_name)
            results['name'] = current_name
        else:
            name = generate_default_log_name()
            log_obj = ensure_logutil_log(group_id, name)
            clear_logutil_items(log_obj['id'])
            log_id_str = str(log_obj['id'])
            with STATE_LOCK:
                LOG_IMPORTED_FILES.pop(log_id_str, None)
            update_logutil_log_meta(log_obj['id'], ended=0)
            update_logutil_group_state(group_id, current_log_name=name, recording=1,
                                       raw_recording=1 if raw_mode else 0)
            capture_baseline_file_id(group_id)
            gid_int = safe_int(group_id, 0)
            if gid_int > 0:
                BRIDGE_POLL_GROUPS.add(gid_int)
                ensure_poll_worker_started()
            results['name'] = name

    # Step 2: resolve and add all ops atomically
    sender_name = str(payload.get('nickname') or 'User')
    sender_id = str(payload.get('user_id') or '')
    ts = int(time.time())

    for i, op in enumerate(ops):
        op_text = str(op or '').strip()
        if not op_text:
            results['ops'].append({'op': op_text, 'status': 'skipped', 'added': 0})
            continue
        try:
            # Resolve: is this a bridge file ref, URL, weizaima link, or raw text?
            target = parse_log_target_entry(op_text)
            resolved_text = None
            if target:
                source = target.get('source', '')
                if source in ('bridge_file', 'bridge_file_name', 'bridge_link', 'bridge_history', 'raw_url',
                              'weizaima', 'kokona', 'dice_zone', 'trpgbot'):
                    try:
                        resolved_text = fetch_log_text_by_source(
                            target['key'],
                            password=target.get('password', ''),
                            source=source,
                            group_id=group_id,
                        )
                    except Exception:
                        resolved_text = None
            if not resolved_text:
                resolved_text = op_text  # use as raw text

            # Parse into items — v4.4.4.1: respect raw_mode
            if raw_mode:
                raw_text = str(resolved_text or '').strip()
                if raw_text:
                    items = [make_log_item(sender_name, sender_id, ts, raw_text,
                                           f"compound-raw:{i}:{int(time.time())}")]
                else:
                    items = []
            else:
                items = parse_structured_text_to_items(
                    resolved_text,
                    sender_name=sender_name,
                    sender_id=sender_id,
                    ts=ts,
                    raw_msg_id=f"compound:{i}:{int(time.time())}",
                )
            if items:
                _, new_count = add_logutil_items(log_obj['id'], items)
                results['ops'].append({
                    'op': op_text[:80],
                    'status': 'ok',
                    'added': len(items),
                })
                results['total_added'] += len(items)
            else:
                results['ops'].append({
                    'op': op_text[:80],
                    'status': 'empty',
                    'added': 0,
                })
        except Exception as e:
            results['ops'].append({
                'op': op_text[:80],
                'status': 'error',
                'error': str(e),
                'added': 0,
            })

    results['status'] = 'ok'
    results['saved'] = results['total_added']

    # Step 3: end recording if requested
    if do_end:
        log_obj = get_logutil_log_full(group_id, results['name'])
        if log_obj and log_obj.get('items'):
            del_paren_compound = bool(payload.get('del_paren'))
            text_filename, content_url = export_log_text(log_obj, del_paren=del_paren_compound)
            now_ts = int(time.time() * 1000)
            try:
                update_logutil_log_meta(log_obj['id'], ended=1, updated_at=now_ts)
                update_logutil_group_state(group_id, recording=0, baseline_file_id='')
                with STATE_LOCK:
                    LOG_IMPORTED_FILES.pop(str(log_obj['id']), None)
            except Exception:
                pass
            log_obj['ended'] = 1
            log_obj['updated_at'] = now_ts

            # Send file via NapCat — 传递 del_paren 以保持与 export_log_text 一致
            file_sent = False
            napcat_error = ''
            try:
                full_text = generate_log_text(log_obj)
                if del_paren_compound:
                    full_text = strip_paren_text(full_text)
                file_sent, _, public_url = send_log_via_napcat(log_obj, group_id, full_text=full_text)
                results['content_url'] = public_url
                results['file_sent'] = file_sent
            except Exception as exc:
                napcat_error = str(exc)
                results['napcat_error'] = napcat_error
                results['content_url'] = content_url

            # Upload to dye
            try:
                dye_link, _ = upload_log_to_story_painter(log_obj, group_id)
                if dye_link:
                    results['dye_link'] = dye_link
            except Exception:
                pass
        else:
            results['end_error'] = 'empty log'

    # Return results (logai flag is handled by JS frontend via /api/submit polling)
    return jsonify(results)


@app.route('/api/logutil_get', methods=['GET'])
def api_logutil_get():
    """获取日志：通过 NapCat 发送 txt 文件到群并上传染色器（fwlog 风格）。
    不会标记为 ended，只发送文件。"""
    group_id = str(request.args.get('group_id') or '').strip()
    name = str(request.args.get('name') or '').strip()
    print(f"[logutil_get] group_id={group_id!r} requested_name={name!r}")
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    if not name:
        name = state.get('current_log_name') or ''
    print(f"[logutil_get] state current_log_name={state.get('current_log_name')!r} resolved_name={name!r}")
    if not name:
        return jsonify({'status': 'error', 'msg': 'missing log name'}), 400
    log_obj = get_logutil_log_full(group_id, name)
    if not log_obj:
        return jsonify({'status': 'error', 'msg': 'log not found'}), 404
    if not log_obj.get('items'):
        return jsonify({'status': 'error', 'msg': 'empty log'}), 400

    # Export to text
    del_paren = str(request.args.get('del_paren') or '').lower() in ('1', 'true', 'yes', 'on')
    text_filename, content_url = export_log_text(log_obj, del_paren=del_paren)

    # Send txt via NapCat (fwlog-style)
    file_sent = False
    public_url = content_url
    napcat_error = ''
    try:
        # 传递 del_paren 以保持与 export_log_text 一致
        full_text = generate_log_text(log_obj)
        if del_paren:
            full_text = strip_paren_text(full_text)
        file_sent, _, public_url = send_log_via_napcat(log_obj, group_id, full_text=full_text)
    except Exception as exc:
        napcat_error = str(exc)

    # Upload to dye server
    dye_link = ''
    try:
        dye_link, _ = upload_log_to_story_painter(log_obj, group_id)
    except Exception:
        dye_link = ''

    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'name': name,
        'text_filename': text_filename,
        'content_url': public_url or content_url,
        'dye_link': dye_link,
        'file_sent': file_sent,
        'napcat_error': napcat_error,
    })


@app.route('/api/logutil_end', methods=['POST'])
def api_logutil_end():
    """结束记录并发送日志：fwlog 风格 — 通过 NapCat 上传 txt 文件到群，
    再上传染色器并返回链接。"""
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    name = str(payload.get('name') or '').strip()
    del_paren = str(payload.get('del_paren') or request.args.get('del_paren') or '').lower() in ('1', 'true', 'yes', 'on')
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    state = ensure_logutil_group_state(group_id)
    if not name:
        name = state.get('current_log_name') or ''
    if not name:
        return jsonify({'status': 'error', 'msg': 'missing log name'}), 400
    log_obj = get_logutil_log_full(group_id, name)
    if not log_obj:
        return jsonify({'status': 'error', 'msg': 'log not found'}), 404
    if not log_obj.get('items'):
        return jsonify({'status': 'error', 'msg': 'empty log'}), 400

    # Export to text (also registers in bridge cache for content_url)
    text_filename, content_url = export_log_text(log_obj, del_paren=del_paren)

    # Mark as ended
    now_ts = int(time.time() * 1000)
    try:
        update_logutil_log_meta(log_obj['id'], ended=1, updated_at=now_ts)
        update_logutil_group_state(group_id, recording=0, baseline_file_id='')
        # Clean up imported files tracking for this log
        with STATE_LOCK:
            LOG_IMPORTED_FILES.pop(str(log_obj['id']), None)
    except Exception:
        pass
    log_obj['ended'] = 1
    log_obj['updated_at'] = now_ts

    # 1. Send txt file via NapCat (fwlog-style)
    file_sent = False
    napcat_result = None
    public_url = content_url
    napcat_error = ''
    try:
        # 传递 del_paren 以保持与 export_log_text 一致
        full_text = generate_log_text(log_obj)
        if del_paren:
            full_text = strip_paren_text(full_text)
        file_sent, napcat_result, public_url = send_log_via_napcat(log_obj, group_id, full_text=full_text)
    except Exception as exc:
        napcat_error = str(exc)

    # 2. Upload to dye server (story painter)
    dye_link = ''
    dye_reused = False
    try:
        dye_link, dye_reused = upload_log_to_story_painter(log_obj, group_id)
    except Exception:
        dye_link = ''

    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'name': name,
        'text_filename': text_filename,
        'content_url': public_url or content_url,
        'dye_link': dye_link,
        'file_sent': file_sent,
        'napcat_error': napcat_error,
    })


# ====== NapCat 消息/文件发送 (fwlog-style logutil end) ======

def _parse_bridge_ref(ref_str):
    """v4.4.0: 解析 '[file]-3', '[link]-1', '[history]-5' 为 ('file', 3) 等。
    v4.4.4.1: 支持跨群 '[file]-3-123456' → ('file', 3, 123456)。"""
    m = re.match(r'^\[(file|link|history)\]-(\d+)(?:-(\d+))?$', str(ref_str).strip(), re.I)
    if m:
        cross_gid = int(m.group(3)) if m.group(3) else None
        return m.group(1).lower(), int(m.group(2)), cross_gid
    return None, -1, None


@app.route('/api/bridge_get', methods=['POST'])
def api_bridge_get():
    """v4.4.0: .bridge get 后端：支持 [file]-N/[link]-N/[history]-N 格式。
    接受 {group_id, ref 或 index, type}，返回上传结果。"""
    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id', 0), 0)
    ref_str = str(payload.get('ref') or '').strip()
    index = safe_int(payload.get('index', -1), -1)
    ref_type = str(payload.get('type', 'file')).lower()

    # Parse ref if provided (e.g. "[file]-3" or "[file]-3-123456" for cross-group)
    cross_group_id = None
    if ref_str:
        parsed_type, parsed_idx, cross_gid = _parse_bridge_ref(ref_str)
        if parsed_type:
            ref_type = parsed_type
            index = parsed_idx
            cross_group_id = cross_gid

    if group_id <= 0 or index < 0:
        return jsonify({'status': 'error', 'msg': 'missing group_id or ref'}), 400

    # v4.4.4.1: cross-group access — use target group for lookup
    lookup_group_id = cross_group_id if cross_group_id else group_id

    if ref_type == 'history':
        with STATE_LOCK:
            hist_list = [h for h in HISTORY if safe_int(h.get('group_id', 0), 0) == lookup_group_id]
        if index >= len(hist_list):
            return jsonify({'status': 'error', 'msg': f'index {index} out of range (0~{len(hist_list)-1})'}), 404
        item = hist_list[index]
    elif ref_type == 'link':
        with STATE_LOCK:
            link_list = list(LINK_CACHE.get(lookup_group_id, []))
        if index >= len(link_list):
            return jsonify({'status': 'error', 'msg': f'index {index} out of range (0~{len(link_list)-1})'}), 404
        item = link_list[index]
    else:
        with STATE_LOCK:
            file_list = list(LATEST_FILES.get(lookup_group_id, []))
        if index >= len(file_list):
            return jsonify({'status': 'error', 'msg': f'index {index} out of range (0~{len(file_list)-1})'}), 404
        item = file_list[index]

    ck = str(item.get('content_key', ''))
    with STATE_LOCK:
        path = CONTENT_INDEX.get(ck, '')
    if not path or not os.path.exists(path):
        return jsonify({'status': 'error', 'msg': 'cached file not found on disk'}), 404

    # v4.4.4: for link-type items, build clean name (not URL-named)
    item_name = str(item.get('name', 'file'))
    if ref_type in ('link', 'history') and item.get('_type') == 'link':
        # Use content_key for a clean short name
        item_name = f"link_{ck[:8]}"
    base_name = os.path.splitext(item_name)[0]
    txt_name = f"{base_name}.txt"
    file_sent, result = napcat_upload_group_file(group_id, path, txt_name)
    if file_sent:
        bridge_log("bridge_get", f"uploaded type={ref_type} group={group_id} file={txt_name}")
        return jsonify({'status': 'ok', 'file_sent': True, 'filename': txt_name})
    return jsonify({'status': 'error', 'file_sent': False,
                    'msg': str(result.get('error', str(result)) if isinstance(result, dict) else result)})


@app.route('/api/bridge_del', methods=['POST'])
def api_bridge_del():
    """v4.4.0: .bridge del 后端：删除指定桥接项。
    接受 {group_id, targets: ["[file]-3", "[link]-1", "[history]-5"]}。
    从后向前处理目标以保证索引正确。"""
    if not check_auth_with_query(request):
        return jsonify({'status': 'unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    group_id = safe_int(payload.get('group_id', 0), 0)
    targets = payload.get('targets', [])
    if group_id <= 0 or not targets:
        return jsonify({'status': 'error', 'msg': 'missing group_id or targets'}), 400

    deleted = []
    errors = []

    with STATE_LOCK:
        file_list = LATEST_FILES.get(group_id, [])
        link_list = LINK_CACHE.get(group_id, [])
        hist_list = list(HISTORY)

        # Sort targets by index descending (highest first) for safe deletion
        parsed_targets = []
        for t in targets:
            ptype, pidx, _ = _parse_bridge_ref(str(t).strip())
            if ptype:
                parsed_targets.append((ptype, pidx, str(t).strip()))
            else:
                errors.append(f'{t}: 无效的引用格式（应为 [file]-N / [link]-N / [history]-N）')
        parsed_targets.sort(key=lambda x: x[1], reverse=True)

        for ptype, pidx, raw_ref in parsed_targets:
            if ptype == 'file':
                if 0 <= pidx < len(file_list):
                    item = file_list.pop(pidx)
                    ck = str(item.get('content_key', ''))
                    CONTENT_INDEX.pop(ck, '')
                    deleted.append(raw_ref)
                else:
                    errors.append(f'{raw_ref}: 索引超出范围 (0~{len(file_list)-1})')
            elif ptype == 'link':
                if 0 <= pidx < len(link_list):
                    item = link_list.pop(pidx)
                    ck = str(item.get('content_key', ''))
                    CONTENT_INDEX.pop(ck, '')
                    deleted.append(raw_ref)
                else:
                    errors.append(f'{raw_ref}: 索引超出范围 (0~{len(link_list)-1})')
            elif ptype == 'history':
                # v4.4.4: filter history by group_id for isolation
                group_hist = [h for h in hist_list if safe_int(h.get('group_id', 0), 0) == group_id]
                if 0 <= pidx < len(group_hist):
                    item = group_hist[pidx]
                    # Remove from actual hist_list by matching content_key
                    ck = str(item.get('content_key', ''))
                    for i, h in enumerate(hist_list):
                        if str(h.get('content_key', '')) == ck:
                            hist_list.pop(i)
                            break
                    CONTENT_INDEX.pop(ck, '')
                    deleted.append(raw_ref)
                else:
                    errors.append(f'{raw_ref}: 索引超出范围 (0~{len(group_hist)-1})')

        # Apply changes back to HISTORY
        HISTORY[:] = hist_list

    # Clean up disk files (outside lock)
    for ptype, pidx, raw_ref in parsed_targets:
        if raw_ref in deleted:
            # disk cleanup already handled by CONTENT_INDEX removal
            pass

    return jsonify({'status': 'ok', 'deleted': deleted, 'errors': errors})


def napcat_send_group_msg(group_id, text):
    """通过 NapCat HTTP API 发送群消息。"""
    return napcat_json_post(
        "/send_group_msg",
        {"group_id": int(group_id), "message": str(text)},
        timeout_sec=20,
    )


def napcat_upload_group_file(group_id, file_path, name):
    """通过 NapCat HTTP API 上传文件到群。
    将文件内容 base64 编码后作为 base64:// 参数发送。
    返回 (file_sent, result) 二元组，失败时 file_sent=False。"""
    try:
        with open(file_path, 'rb') as f:
            file_bytes = f.read()
        b64_content = base64.b64encode(file_bytes).decode('utf-8')
        result = napcat_json_post(
            "/upload_group_file",
            {
                "group_id": int(group_id),
                "file": f"base64://{b64_content}",
                "name": str(name),
            },
            timeout_sec=60,
        )
        bridge_log("send", f"upload_group_file ok group={group_id} name={name} size={len(file_bytes)}")
        return True, result
    except Exception as e:
        bridge_log("send", f"upload_group_file FAILED group={group_id} name={name} err={e}")
        return False, {"error": str(e)}


def resolve_public_base_or_fallback():
    """返回当前请求的 public_base，若不可用或为 loopback 则返回 LAN IP 默认值。"""
    try:
        from flask import has_request_context
        if has_request_context() and request:
            host_url = (request.host_url or '').rstrip('/')
            if host_url and not is_loopback_base(host_url):
                return host_url
    except Exception:
        pass
    return BRIDGE_PUBLIC_BASE


def send_log_via_napcat(log_obj, group_id, full_text=None):
    """将 log_obj 生成为 txt 并通过 NapCat 上传到群。
    若提供 full_text 则直接使用（避免重复生成，保证与 export_log_text 一致）。
    返回 (file_sent, napcat_result, public_url) 三元组。"""
    if full_text is None:
        full_text = generate_log_text(log_obj)
    if not full_text or not str(full_text).strip():
        raise RuntimeError("日志内容为空")

    # Write to temp file
    ensure_bridge_cache_dir()
    tmp_key = uuid.uuid4().hex
    tmp_path = os.path.join(BRIDGE_CACHE_DIR, f"{tmp_key}_send.txt")
    txt_name = f"{str(log_obj.get('name') or 'log')}.txt"

    try:
        with open(tmp_path, 'w', encoding='utf-8') as fw:
            fw.write(str(full_text))

        # 1. Send via NapCat upload_group_file
        file_sent, napcat_result = napcat_upload_group_file(group_id, tmp_path, txt_name)

        # 2. Also register in bridge content index so the file is accessible via URL
        public_base = resolve_public_base_or_fallback()
        with open(tmp_path, 'r', encoding='utf-8') as f:
            content = f.read()
        out_path = os.path.join(BRIDGE_CACHE_DIR, f"{tmp_key}.txt")
        os.rename(tmp_path, out_path)
        tmp_path = ""  # prevent deletion below
        with STATE_LOCK:
            CONTENT_INDEX[tmp_key] = out_path
        public_url = build_content_url(tmp_key, public_base=public_base)

        # Register in bridge file list
        group_id_int = safe_int(group_id, 0)
        if group_id_int > 0:
            bridge_item = {
                "group_id": group_id_int,
                "file_id": f"logutil-send-{tmp_key[:8]}",
                "busid": 0,
                "name": txt_name,
                "user_id": 0,
                "ts": int(time.time()),
                "source_ts": int(time.time()),
                "cached_ts": int(time.time()),
                "content_key": tmp_key,
                "content_url": public_url,
                "text_filename": txt_name,
                "text_chars": len(content),
                "text_bytes": os.path.getsize(out_path),
                "_type": "file",
            }
            with STATE_LOCK:
                if group_id_int not in LATEST_FILES:
                    LATEST_FILES[group_id_int] = []
                LATEST_FILES[group_id_int].append(bridge_item)
                # v4.4.0: evict overflow to history
                while len(LATEST_FILES[group_id_int]) > MAX_BRIDGE_FILES_PER_GROUP:
                    removed = LATEST_FILES[group_id_int].pop(0)
                    _evict_to_history(removed, 'file')

        return file_sent, napcat_result, public_url

    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


# ====== 发出消息记录 ======

def should_ignore_self_log_message(text):
    """过滤无需记录的自身消息（空消息/CQ:file/系统已记录提醒）"""
    normalized = str(text or "").strip()
    if not normalized:
        return True
    if normalized.startswith("[CQ:file"):
        return True
    if SELF_LOG_IGNORE_RE.search(normalized):
        return True
    return False


def record_logutil_outgoing_message(group_id, nickname, user_id, text):
    """将 bot 发出的消息记入当前活跃日志"""
    if should_ignore_self_log_message(text):
        return None
    group_id = str(group_id or "").strip()
    if not group_id:
        return None
    state = ensure_logutil_group_state(group_id)
    current_log_name = state.get("current_log_name") or ""
    if not state.get("recording") or not current_log_name:
        return None
    log_obj = ensure_logutil_log(group_id, current_log_name)
    items = [
        make_log_item(
            nickname or "Bot",
            str(user_id or ""),
            int(time.time()),
            str(text),
            f"self:{uuid.uuid4().hex}",
        )
    ]
    old_count, new_count = add_logutil_items(log_obj["id"], items)
    return {"old_count": old_count, "new_count": new_count}


@app.route('/api/logutil_record_outgoing', methods=['POST'])
def api_logutil_record_outgoing():
    """记录 bot 发出消息的 API 端点"""
    payload = request.get_json(silent=True) or {}
    group_id = str(payload.get('group_id') or request.args.get('group_id') or '').strip()
    if not group_id:
        return jsonify({'status': 'error', 'msg': 'missing group_id'}), 400
    nickname = str(payload.get('nickname') or 'Bot')
    user_id = str(payload.get('user_id') or '')
    text = str(payload.get('text') or payload.get('message') or '')
    if not text.strip():
        return jsonify({'status': 'error', 'msg': 'missing text'}), 400
    result = record_logutil_outgoing_message(group_id, nickname, user_id, text)
    if result is None:
        state = ensure_logutil_group_state(group_id)
        if not state.get("recording"):
            return jsonify({'status': 'ok', 'recorded': False, 'reason': 'not_recording'})
        return jsonify({'status': 'ok', 'recorded': False, 'reason': 'ignored'})
    return jsonify({
        'status': 'ok',
        'group_id': group_id,
        'recorded': True,
        'saved': result['new_count'],
    })


# ====== WebSocket 实时监听 (logutil recording) ======

def ws_log(*args):
    print(f"[logutil-ws][{time.strftime('%Y-%m-%d %H:%M:%S')}]", *args)


def normalize_logutil_prefix(text):
    """Normalize .logutil command prefix (adapted from fwlog's normalize_fwlog_prefix)."""
    if not text:
        return ""
    t = text.lstrip()
    if not t:
        return ""
    prefixes = [".", "。", "/", "、"]  # . 。 / 、
    has_prefix = False
    for p in prefixes:
        if t.startswith(p):
            t = t[len(p):].lstrip()
            has_prefix = True
            break
    if t.lower().startswith("logutil"):
        if has_prefix:
            return ".logutil" + t[7:]
    return text


class LogutilBotClient:
    """Async WebSocket client for NapCat OneBot v11.
    Adapted from fwlog's BotClient for logutil recording use."""

    def __init__(self, ws_url, ws_token):
        self.ws_url = ws_url
        self.ws_token = ws_token
        self.ws_conn = None
        self.pending = {}
        self.bot_user_id = ""
        self._echo_counter = 0

    def _gen_echo(self):
        self._echo_counter += 1
        return f"logutil-ws-{self._echo_counter}"

    async def send_api(self, action, params=None):
        if params is None:
            params = {}
        if self.ws_conn is None or self.ws_conn.closed:
            raise RuntimeError("[logutil-ws] WebSocket 未连接")
        echo = self._gen_echo()
        loop = asyncio.get_running_loop()
        fut = loop.create_future()
        self.pending[echo] = fut
        payload = {"action": action, "params": params, "echo": echo}
        if self.ws_token:
            payload["token"] = self.ws_token
        try:
            await self.ws_conn.send(json.dumps(payload, ensure_ascii=False))
            return await asyncio.wait_for(fut, timeout=20.0)
        except asyncio.TimeoutError:
            self.pending.pop(echo, None)
            raise RuntimeError(f"[logutil-ws] API请求超时: {action}")
        except Exception as e:
            self.pending.pop(echo, None)
            raise e

    def handle_api_response(self, msg):
        echo = msg.get("echo")
        if not echo:
            return
        fut = self.pending.pop(echo, None)
        if fut is None:
            return
        if not fut.done():
            fut.set_result(msg)

    async def send_group_msg(self, group_id, text):
        try:
            await self.send_api(
                "send_group_msg",
                {"group_id": str(group_id), "message": text},
            )
            record_logutil_outgoing_message(str(group_id), self.bot_user_id, self.bot_user_id, text)
        except Exception as e:
            ws_log(f"发送群消息失败: {e}")

    async def run(self):
        """Connect to NapCat WebSocket, receive events, dispatch to message queue."""
        global WS_MESSAGE_QUEUE
        while True:
            try:
                ws_log(f"尝试连接到 NapCat WS: {self.ws_url}")
                extra_headers = (
                    {"Authorization": f"Bearer {self.ws_token}"}
                    if self.ws_token
                    else None
                )
                async with ws_connect(
                    self.ws_url,
                    extra_headers=extra_headers,
                ) as ws:
                    self.ws_conn = ws
                    ws_log("WS 已连接")
                    async for raw_message in ws:
                        try:
                            data = json.loads(raw_message)
                        except Exception:
                            continue

                        if isinstance(data, dict) and "echo" in data:
                            self.handle_api_response(data)
                            continue

                        if isinstance(data, dict):
                            if data.get("self_id"):
                                self.bot_user_id = str(
                                    data.get("self_id") or self.bot_user_id
                                )
                            post_type = str(data.get("post_type") or "").lower()
                            notice_type = str(data.get("notice_type") or "").lower()
                            if post_type == "message" and data.get("message_type") in [
                                "group",
                                "private",
                            ]:
                                WS_MESSAGE_QUEUE.put_nowait(data)
                            elif post_type == "notice" and notice_type == "group_upload":
                                WS_MESSAGE_QUEUE.put_nowait(data)

            except Exception as e:
                ws_log(f"WS 连接出错或关闭: {e}")
            self.ws_conn = None
            await asyncio.sleep(3)


def _auto_cache_urls_from_message(msg, text):
    """v4.4.1: 自动检测消息中的染色器链接，拉取并缓存到 LINK_CACHE。
    仅在链接尚未缓存时执行（按 URL 去重）。"""
    if not text:
        return
    group_id = safe_int(msg.get("group_id", 0), 0)
    if group_id <= 0:
        return
    # 提取消息中所有 URL
    urls = URL_RE.findall(text)
    if not urls:
        return
    for url in urls:
        url = url.strip().rstrip(')').rstrip(']').rstrip('"').rstrip("'")
        if not url or not url.startswith(('http://', 'https://')):
            continue
        target = parse_log_target_entry(url)
        if not target:
            continue
        source = target.get('source', '')
        # 只处理已知的染色器源
        if source not in ('weizaima', 'dice_zone', 'kokona', 'trpgbot', 'raw_url'):
            continue
        # v4.4.4: 黑名单 — QQ多媒体CDN链接不缓存
        if 'multimedia.nt.qq.com.cn' in url:
            continue
        # 去重：已在 LINK_CACHE 中则跳过
        with STATE_LOCK:
            link_list = LINK_CACHE.get(group_id, [])
            already = any(str(item.get('url', '')) == url for item in link_list)
        if already:
            continue
        # 异步拉取并缓存（在线程中执行以免阻塞 WS 循环）
        try:
            executor.submit(_fetch_and_cache_link, group_id, url, target)
        except Exception:
            pass  # 线程池满时静默跳过


def _fetch_and_cache_link(group_id, url, target):
    """后台拉取染色器链接并写入 LINK_CACHE。"""
    try:
        text = fetch_log_text_by_source(
            target['key'],
            password=target.get('password', ''),
            source=target.get('source'),
            group_id=group_id,
        )
        if text:
            ws_log(f"auto-cached link: group={group_id} url={url[:60]}")
    except Exception as e:
        ws_log(f"auto-cache link failed: {e}")


async def handle_ws_recording_event(event):
    """Process a group message or file upload event, extracting items and adding
    them to the current logutil recording.  Mirrors fwlog's handle_recording_event
    exactly in behavior, adapted for logutil's DB + HTTP-based NapCat access."""
    post_type = str(event.get("post_type") or "").lower()
    notice_type = str(event.get("notice_type") or "").lower()

    group_id = str(event.get("group_id") or "")
    if not group_id:
        return

    # Check recording state
    group_state = ensure_logutil_group_state(group_id)
    current_log_name = group_state.get("current_log_name") or ""
    if not group_state.get("recording") or not current_log_name:
        return

    log_obj = ensure_logutil_log(group_id, current_log_name)
    raw_recording = bool(group_state.get("raw_recording"))

    # Sender info
    sender = event.get("sender") or {}
    sender_name = sender.get("card") or sender.get("nickname") or (
        f"QQ:{sender.get('user_id', '')}" if sender.get("user_id") else "Unknown"
    )
    sender_id = str(sender.get("user_id") or event.get("user_id") or "")
    event_ts = safe_int(event.get("time"), int(time.time()))
    items = []

    # raw 修饰符：跳过消息头解析，直接拼接原始文本
    if raw_recording and post_type == "message":
        raw_text = segments_to_text(event.get("message")).strip()
        if raw_text:
            items = [make_log_item(sender_name, sender_id, event_ts, raw_text,
                                   str(event.get("message_id") or f"raw:{event_ts}"))]
        if items:
            old_count, new_count = add_logutil_items(log_obj["id"], items)
            ws_log(f"[raw] 已追加 {len(items)} 条 (当前共 {new_count} 条)")
        return

    if post_type == "notice" and notice_type == "group_upload":
        baseline_file_id = str(group_state.get('baseline_file_id', '')).strip()
        for index, payload in enumerate(extract_file_payloads(event)):
            if not remember_file_capture(group_id, current_log_name, event_ts):
                continue
            payload_file_id = str(payload.get('file_id') or '')
            # Baseline check: skip files that existed when recording started
            if baseline_file_id and payload_file_id == baseline_file_id:
                ws_log(f"baseline skipped: file={payload.get('name')} file_id={payload_file_id}")
                continue
            # In-memory dedup: skip if already imported for this log
            log_id_str = str(log_obj['id'])
            if payload_file_id:
                with STATE_LOCK:
                    if payload_file_id in LOG_IMPORTED_FILES.get(log_id_str, set()):
                        ws_log(f"dedup skipped (mem): file_id={payload_file_id} already in log={current_log_name}")
                        continue
            # DB dedup: skip if this file_id already imported for this log
            if payload_file_id:
                conn = get_logutil_db_connection()
                c = conn.cursor()
                import_pattern = f"file:{payload_file_id}%"
                c.execute(
                    'SELECT COUNT(*) FROM items WHERE log_id = ? AND raw_msg_id LIKE ?',
                    (log_obj['id'], import_pattern),
                )
                already = c.fetchone()[0]
                conn.close()
                if already > 0:
                    ws_log(f"dedup skipped (db): file_id={payload_file_id} already in log={current_log_name}")
                    continue
            try:
                # Read from bridge cache (populated by process_ws_bridge_file which runs first)
                gid_int = safe_int(group_id, 0)
                file_text = None
                cached_name = payload.get('name', '')
                with STATE_LOCK:
                    file_list = list(LATEST_FILES.get(gid_int, []))
                for cached_item in file_list:
                    if str(cached_item.get('file_id', '')) == payload_file_id:
                        ck = cached_item.get('content_key', '')
                        with STATE_LOCK:
                            disk_path = CONTENT_INDEX.get(ck, '')
                        if disk_path and os.path.exists(disk_path):
                            with open(disk_path, 'r', encoding='utf-8') as f:
                                file_text = f.read()
                            cached_name = cached_item.get('name', cached_name)
                        break
                if file_text:
                    # Parse from cache — no redundant download
                    file_items = parse_structured_text_to_items(
                        file_text,
                        sender_name,
                        sender_id,
                        event_ts,
                        f"file:{payload_file_id}",
                    )
                    items.extend(file_items)
                    if payload_file_id:
                        with STATE_LOCK:
                            if log_id_str not in LOG_IMPORTED_FILES:
                                LOG_IMPORTED_FILES[log_id_str] = set()
                            LOG_IMPORTED_FILES[log_id_str].add(payload_file_id)
                else:
                    # Rare fallback: cache miss → download via NapCat
                    ws_log(f"cache miss for file_id={payload_file_id}, falling back to download")
                    file_items = await extract_items_from_file_payload_napcat(
                        payload,
                        sender_name,
                        sender_id,
                        event_ts,
                        f"file:{payload_file_id or payload.get('name') or f'notice-{index}'}",
                    )
                    items.extend(file_items)
                    if payload_file_id and file_items:
                        with STATE_LOCK:
                            if log_id_str not in LOG_IMPORTED_FILES:
                                LOG_IMPORTED_FILES[log_id_str] = set()
                            LOG_IMPORTED_FILES[log_id_str].add(payload_file_id)
            except Exception as e:
                ws_log(f"文档提取失败: {payload.get('name')} err={e}")
    else:
        message = event.get("message")
        message_id = str(event.get("message_id") or f"event:{event_ts}")
        if isinstance(message, list):
            for index, seg in enumerate(message):
                if not isinstance(seg, dict):
                    continue
                seg_type = str(seg.get("type") or "").lower()
                data = seg.get("data") or {}
                if seg_type == "text":
                    text = str(data.get("text") or "")
                    if text.strip():
                        items.extend(
                            await extract_items_from_text_chunk(
                                text,
                                sender_name,
                                sender_id,
                                event_ts,
                                f"{message_id}:text:{index}",
                                group_id,
                            )
                        )
                elif seg_type == "forward":
                    forward_id = (
                        data.get("id")
                        or data.get("res_id")
                        or data.get("message_id")
                    )
                    if not forward_id:
                        continue
                    try:
                        items.extend(
                            await extract_items_from_forward(
                                str(forward_id),
                                group_id,
                                sender_name,
                                sender_id,
                                event_ts,
                                f"{message_id}:fwd:{forward_id}",
                            )
                        )
                    except Exception as e:
                        ws_log(f"获取转发消息异常: forward_id={forward_id} err={e}")
                elif seg_type == "file":
                    payloads = extract_file_payloads(
                        {
                            "post_type": post_type,
                            "message": [seg],
                            "group_id": event.get("group_id"),
                            "user_id": event.get("user_id"),
                        }
                    )
                    for payload in payloads:
                        if not remember_file_capture(
                            group_id, current_log_name, event_ts
                        ):
                            continue
                        try:
                            file_items = await extract_items_from_file_payload_napcat(
                                payload,
                                sender_name,
                                sender_id,
                                event_ts,
                                f"{message_id}:file:{payload.get('file_id') or payload.get('name') or f'message-{index}'}",
                            )
                            items.extend(file_items)
                        except Exception as e:
                            ws_log(
                                f"文档提取失败: {payload.get('name')} err={e}"
                            )
        else:
            text = segments_to_text(message).strip()
            if text:
                # Expand forward CQ codes before cleaning
                forward_ids = extract_forward_ids_from_text(text)
                for index, forward_id in enumerate(forward_ids):
                    try:
                        items.extend(
                            await extract_items_from_forward(
                                str(forward_id),
                                group_id,
                                sender_name,
                                sender_id,
                                event_ts,
                                f"{message_id}:fwd:{forward_id}",
                            )
                        )
                    except Exception as e:
                        ws_log(
                            f"获取转发消息异常: forward_id={forward_id} err={e}"
                        )

                for payload in extract_file_payloads(event):
                    if not remember_file_capture(
                        group_id, current_log_name, event_ts
                    ):
                        continue
                    try:
                        file_items = await extract_items_from_file_payload_napcat(
                            payload,
                            sender_name,
                            sender_id,
                            event_ts,
                            f"{message_id}:file:{payload.get('file_id') or payload.get('name')}",
                        )
                        items.extend(file_items)
                    except Exception as e:
                        ws_log(f"文档提取失败: {payload.get('name')} err={e}")

                cleaned_text = re.sub(
                    r"\[CQ:(?:forward|file)[^\]]*\]", "", text
                ).strip()
                if cleaned_text:
                    items.extend(
                        await extract_items_from_text_chunk(
                            cleaned_text,
                            sender_name,
                            sender_id,
                            event_ts,
                            f"{message_id}:text",
                            group_id,
                        )
                    )

    if not items:
        return

    old_count, new_count = add_logutil_items(log_obj["id"], items)
    ws_log(
        f"已追加 {len(items)} 条 logutil 内容 "
        f"(当前共 {new_count} 条, log={current_log_name}, group={group_id})"
    )

    # Milestone notification every 1000 items
    if new_count // LOGUTIL_MILESTONE_INTERVAL > old_count // LOGUTIL_MILESTONE_INTERVAL:
        try:
            napcat_json_post(
                "send_group_msg",
                {
                    "group_id": int(group_id),
                    "message": (
                        f"【系统提醒】 当前日志 {log_obj['name']} 已记录 {new_count} 条消息。\n"
                        "如果记录完毕，请记得发送 .logutil end 结束记录。"
                    ),
                },
                timeout_sec=10,
            )
        except Exception as e:
            ws_log(f"里程碑通知失败: {e}")


async def process_ws_bridge_file(event):
    """Process a WS group_upload event for the file bridge (mode 0).
    Downloads the file, extracts text, caches it in LATEST_FILES,
    and auto-imports to logutil if recording.
    Returns the cached bridge item dict or None."""
    post_type = str(event.get("post_type") or "").lower()
    notice_type = str(event.get("notice_type") or "").lower()
    if post_type != "notice" or notice_type != "group_upload":
        return None

    group_id = safe_int(event.get("group_id"), 0)
    if group_id <= 0:
        return None

    file_info = event.get("file") or {}
    file_id = str(file_info.get("id") or file_info.get("file_id") or "")
    if not file_id:
        return None

    # Dedup: skip if we already processed this exact file_id recently
    with STATE_LOCK:
        file_list = LATEST_FILES.get(group_id, [])
    for item in file_list:
        if str(item.get("file_id", "")) == file_id:
            bridge_log("ws-bridge", f"dedup skipped: file_id={file_id} already cached")
            return None

    info = {
        "group_id": group_id,
        "user_id": safe_int(event.get("user_id"), 0),
        "file_id": file_id,
        "name": str(file_info.get("name") or file_info.get("file_name") or "未知文件"),
        "busid": safe_int(file_info.get("busid"), 0),
        "url": str(file_info.get("url") or ""),
        "source_ts": safe_int(event.get("time"), int(time.time())),
        "event_type": "ws.group_upload",
        "public_base": "",
    }

    try:
        pulled = await asyncio.to_thread(process_group_upload, info)
    except Exception as exc:
        bridge_log("ws-bridge", f"process failed group={group_id} file={info['name']} err={exc}")
        with STATE_LOCK:
            LAST_ERROR_BY_GROUP[group_id] = str(exc)
        return None

    duplicate = find_duplicate_by_name_and_chars(group_id, pulled)
    if duplicate:
        remove_bridge_cache_item(group_id, pulled)
        bridge_log("ws-bridge", f"dedup skipped (same name+chars): group={group_id} file={pulled.get('name', '')}")
        return None

    with STATE_LOCK:
        LAST_ERROR_BY_GROUP.pop(group_id, None)
    bridge_log("ws-bridge", f"ok group={group_id} file={pulled.get('name', '')} chars={pulled.get('text_chars', 0)}")

    # Auto-import to logutil if recording
    try:
        auto_import_bridge_file_to_logutil(pulled)
    except Exception as e:
        bridge_log("ws-bridge", f"autoimport failed group={group_id} err={e}")

    return pulled


async def process_ws_messages():
    """Consume messages from the WS queue and dispatch.
    Handles both logutil recording and WS-driven file bridge (mode 0).
    Mirrors fwlog's process_messages."""
    global WS_MESSAGE_QUEUE
    ws_log("消息处理循环已启动")
    while True:
        msg = await WS_MESSAGE_QUEUE.get()
        if not isinstance(msg, dict):
            WS_MESSAGE_QUEUE.task_done()
            continue

        try:
            post_type = str(msg.get("post_type") or "").lower()
            notice_type = str(msg.get("notice_type") or "").lower()

            # --- WS-driven file bridge (mode 0) ---
            if post_type == "notice" and notice_type == "group_upload":
                if NC_FILE_BRIDGE_MODE == 0:
                    await process_ws_bridge_file(msg)
                await handle_ws_recording_event(msg)
            else:
                text = segments_to_text(msg.get("message")).strip()
                # v4.4.4: expand short aliases before pattern matching
                text = expand_short_alias(text)
                normalized = normalize_logutil_prefix(text)

                # v4.4.1: 自动缓存消息中的染色器链接
                _auto_cache_urls_from_message(msg, text)

                if normalized.startswith(".logutil"):
                    ws_log(f"检测到 logutil 指令 (WS端跳过，由HTTP处理): {text[:80]}")
                else:
                    # --- [file]-N command: append bridge-cached file to current log ---
                    file_idx_match = re.match(r'^\[file\]-(\d+)(?:-(\d+))?\s*$', text, re.IGNORECASE)
                    if file_idx_match:
                        group_id = str(msg.get("group_id") or "")
                        # v4.4.5: cross-group override
                        xgid = file_idx_match.group(2)
                        lookup_gid = safe_int(xgid, 0) if xgid else safe_int(group_id, 0)
                        if group_id:
                            gs = ensure_logutil_group_state(group_id)
                            if gs.get("recording") and gs.get("current_log_name"):
                                file_idx = int(file_idx_match.group(1))
                                with STATE_LOCK:
                                    file_list = list(LATEST_FILES.get(lookup_gid, []))
                                if 0 <= file_idx < len(file_list):
                                    bridge_item = file_list[file_idx]
                                    ck = bridge_item.get("content_key", "")
                                    with STATE_LOCK:
                                        disk_path = CONTENT_INDEX.get(ck, "")
                                    if disk_path and os.path.exists(disk_path):
                                        try:
                                            with open(disk_path, "r", encoding="utf-8") as f:
                                                file_text = f.read()
                                            sender_name = (
                                                (msg.get("sender") or {}).get("card")
                                                or (msg.get("sender") or {}).get("nickname")
                                                or f"QQ:{msg.get('sender', {}).get('user_id', '')}"
                                                or "Unknown"
                                            )
                                            sender_id = str((msg.get("sender") or {}).get("user_id") or msg.get("user_id") or "")
                                            event_ts = safe_int(msg.get("time"), int(time.time()))
                                            # v4.4.4.1: respect raw_recording
                                            if gs.get("raw_recording"):
                                                raw_text = file_text.strip()
                                                items = [make_log_item(sender_name, sender_id, event_ts, raw_text,
                                                                       f"file-cmd-raw:{file_idx}")] if raw_text else []
                                            else:
                                                items = await extract_items_from_text_chunk(
                                                    file_text,
                                                    sender_name,
                                                    sender_id,
                                                    event_ts,
                                                    f"file-cmd:{file_idx}:{bridge_item.get('file_id', '')}",
                                                    group_id,
                                                )
                                            if items:
                                                log_obj = ensure_logutil_log(group_id, gs["current_log_name"])
                                                _, new_n = add_logutil_items(log_obj["id"], items)
                                                ws_log(
                                                    f"[file]-{file_idx} 追加 {len(items)} 条, "
                                                    f"当前 {new_n} 条, file={bridge_item.get('name', '')}"
                                                )
                                                napcat_json_post(
                                                    "send_group_msg",
                                                    {
                                                        "group_id": int(group_id),
                                                        "message": f"[file]-{file_idx} 已追加 {len(items)} 条: {bridge_item.get('name', '')} (当前共 {new_n} 条)",
                                                    },
                                                    timeout_sec=10,
                                                )
                                            else:
                                                napcat_json_post(
                                                    "send_group_msg",
                                                    {
                                                        "group_id": int(group_id),
                                                        "message": f"[file]-{file_idx} 文件中未提取到可用条目: {bridge_item.get('name', '')}",
                                                    },
                                                    timeout_sec=10,
                                                )
                                        except Exception as e:
                                            ws_log(f"[file]-{file_idx} 读取失败: {e}")
                                            napcat_json_post(
                                                "send_group_msg",
                                                {
                                                    "group_id": int(group_id),
                                                    "message": f"[file]-{file_idx} 读取失败: {e}",
                                                },
                                                timeout_sec=10,
                                            )
                                    else:
                                        napcat_json_post(
                                            "send_group_msg",
                                            {
                                                "group_id": int(group_id),
                                                "message": f"[file]-{file_idx} 缓存文件不存在",
                                            },
                                            timeout_sec=10,
                                        )
                                else:
                                    napcat_json_post(
                                        "send_group_msg",
                                        {
                                            "group_id": int(group_id),
                                            "message": f"[file]-{file_idx} 超出范围 (当前缓存 0~{len(file_list)-1})",
                                        },
                                        timeout_sec=10,
                                    )
                            else:
                                ws_log(f"[file]-N ignored: group not recording, group_id={group_id}")
                        continue
                    # --- v4.4.0: [link]-N command: append bridge-cached link text to current log ---
                    link_idx_match = re.match(r'^\[link\]-(\d+)(?:-(\d+))?\s*$', text, re.IGNORECASE)
                    if link_idx_match:
                        group_id = str(msg.get("group_id") or "")
                        # v4.4.5: cross-group override
                        xgid = link_idx_match.group(2)
                        lookup_gid = safe_int(xgid, 0) if xgid else safe_int(group_id, 0)
                        if group_id:
                            gs = ensure_logutil_group_state(group_id)
                            if gs.get("recording") and gs.get("current_log_name"):
                                link_idx = int(link_idx_match.group(1))
                                with STATE_LOCK:
                                    link_list = list(LINK_CACHE.get(lookup_gid, []))
                                if 0 <= link_idx < len(link_list):
                                    link_item = link_list[link_idx]
                                    ck = link_item.get("content_key", "")
                                    with STATE_LOCK:
                                        disk_path = CONTENT_INDEX.get(ck, "")
                                    if disk_path and os.path.exists(disk_path):
                                        try:
                                            with open(disk_path, "r", encoding="utf-8") as f:
                                                file_text = f.read()
                                            sender_name = (
                                                (msg.get("sender") or {}).get("card")
                                                or (msg.get("sender") or {}).get("nickname")
                                                or f"QQ:{msg.get('sender', {}).get('user_id', '')}"
                                                or "Unknown"
                                            )
                                            sender_id = str((msg.get("sender") or {}).get("user_id") or msg.get("user_id") or "")
                                            event_ts = safe_int(msg.get("time"), int(time.time()))
                                            # v4.4.4.1: respect raw_recording
                                            if gs.get("raw_recording"):
                                                raw_text = file_text.strip()
                                                items = [make_log_item(sender_name, sender_id, event_ts, raw_text,
                                                                       f"link-cmd-raw:{link_idx}")] if raw_text else []
                                            else:
                                                items = await extract_items_from_text_chunk(
                                                    file_text,
                                                    sender_name,
                                                    sender_id,
                                                    event_ts,
                                                    f"link-cmd:{link_idx}:{link_item.get('url', '')}",
                                                    group_id,
                                                )
                                            if items:
                                                log_obj = ensure_logutil_log(group_id, gs["current_log_name"])
                                                _, new_n = add_logutil_items(log_obj["id"], items)
                                                ws_log(
                                                    f"[link]-{link_idx} 追加 {len(items)} 条, "
                                                    f"当前 {new_n} 条, url={link_item.get('url', '')[:60]}"
                                                )
                                                napcat_json_post(
                                                    "send_group_msg",
                                                    {
                                                        "group_id": int(group_id),
                                                        "message": f"[link]-{link_idx} 已追加 {len(items)} 条: {link_item.get('url', '')[:50]} (当前共 {new_n} 条)",
                                                    },
                                                    timeout_sec=10,
                                                )
                                            else:
                                                napcat_json_post(
                                                    "send_group_msg",
                                                    {
                                                        "group_id": int(group_id),
                                                        "message": f"[link]-{link_idx} 文件中未提取到可用条目: {link_item.get('url', '')[:50]}",
                                                    },
                                                    timeout_sec=10,
                                                )
                                        except Exception as e:
                                            ws_log(f"[link]-{link_idx} 读取失败: {e}")
                                            napcat_json_post(
                                                "send_group_msg",
                                                {
                                                    "group_id": int(group_id),
                                                    "message": f"[link]-{link_idx} 读取失败: {e}",
                                                },
                                                timeout_sec=10,
                                            )
                                    else:
                                        napcat_json_post(
                                            "send_group_msg",
                                            {
                                                "group_id": int(group_id),
                                                "message": f"[link]-{link_idx} 缓存文件不存在",
                                            },
                                            timeout_sec=10,
                                        )
                                else:
                                    napcat_json_post(
                                        "send_group_msg",
                                        {
                                            "group_id": int(group_id),
                                            "message": f"[link]-{link_idx} 超出范围 (当前缓存 0~{len(link_list)-1})",
                                        },
                                        timeout_sec=10,
                                    )
                            else:
                                ws_log(f"[link]-N ignored: group not recording, group_id={group_id}")
                        continue
                    # --- v4.4.3: [history]-N command: append evicted bridge item to current log ---
                    history_idx_match = re.match(r'^\[history\]-(\d+)(?:-(\d+))?\s*$', text, re.IGNORECASE)
                    if history_idx_match:
                        group_id = str(msg.get("group_id") or "")
                        # v4.4.5: cross-group override
                        xgid = history_idx_match.group(2)
                        lookup_gid = safe_int(xgid, 0) if xgid else safe_int(group_id, 0)
                        if group_id:
                            gs = ensure_logutil_group_state(group_id)
                            if gs.get("recording") and gs.get("current_log_name"):
                                hist_idx = int(history_idx_match.group(1))
                                with STATE_LOCK:
                                    hist_list = [h for h in HISTORY if safe_int(h.get('group_id', 0), 0) == lookup_gid]
                                if 0 <= hist_idx < len(hist_list):
                                    hist_item = hist_list[hist_idx]
                                    ck = hist_item.get("content_key", "")
                                    with STATE_LOCK:
                                        disk_path = CONTENT_INDEX.get(ck, "")
                                    if disk_path and os.path.exists(disk_path):
                                        try:
                                            with open(disk_path, "r", encoding="utf-8") as f:
                                                file_text = f.read()
                                            sender_name = (
                                                (msg.get("sender") or {}).get("card")
                                                or (msg.get("sender") or {}).get("nickname")
                                                or f"QQ:{msg.get('sender', {}).get('user_id', '')}"
                                                or "Unknown"
                                            )
                                            sender_id = str((msg.get("sender") or {}).get("user_id") or msg.get("user_id") or "")
                                            event_ts = safe_int(msg.get("time"), int(time.time()))
                                            # v4.4.4.1: respect raw_recording
                                            if gs.get("raw_recording"):
                                                raw_text = file_text.strip()
                                                items = [make_log_item(sender_name, sender_id, event_ts, raw_text,
                                                                       f"history-cmd-raw:{hist_idx}")] if raw_text else []
                                            else:
                                                items = await extract_items_from_text_chunk(
                                                    file_text,
                                                    sender_name,
                                                    sender_id,
                                                    event_ts,
                                                    f"history-cmd:{hist_idx}:{hist_item.get('name', hist_item.get('url', ''))}",
                                                    group_id,
                                                )
                                            if items:
                                                log_obj = ensure_logutil_log(group_id, gs["current_log_name"])
                                                _, new_n = add_logutil_items(log_obj["id"], items)
                                                ws_log(
                                                    f"[history]-{hist_idx} 追加 {len(items)} 条, "
                                                    f"当前 {new_n} 条, name={hist_item.get('name', hist_item.get('url', ''))[:60]}"
                                                )
                                                napcat_json_post(
                                                    "send_group_msg",
                                                    {
                                                        "group_id": int(group_id),
                                                        "message": f"[history]-{hist_idx} 已追加 {len(items)} 条: {hist_item.get('name', hist_item.get('url', ''))[:50]} (当前共 {new_n} 条)",
                                                    },
                                                    timeout_sec=10,
                                                )
                                            else:
                                                napcat_json_post(
                                                    "send_group_msg",
                                                    {
                                                        "group_id": int(group_id),
                                                        "message": f"[history]-{hist_idx} 文件中未提取到可用条目: {hist_item.get('name', hist_item.get('url', ''))[:50]}",
                                                    },
                                                    timeout_sec=10,
                                                )
                                        except Exception as e:
                                            ws_log(f"[history]-{hist_idx} 读取失败: {e}")
                                            napcat_json_post(
                                                "send_group_msg",
                                                {
                                                    "group_id": int(group_id),
                                                    "message": f"[history]-{hist_idx} 读取失败: {e}",
                                                },
                                                timeout_sec=10,
                                            )
                                    else:
                                        napcat_json_post(
                                            "send_group_msg",
                                            {
                                                "group_id": int(group_id),
                                                "message": f"[history]-{hist_idx} 缓存文件不存在",
                                            },
                                            timeout_sec=10,
                                        )
                                else:
                                    napcat_json_post(
                                        "send_group_msg",
                                        {
                                            "group_id": int(group_id),
                                            "message": f"[history]-{hist_idx} 超出范围 (当前缓存 0~{len(hist_list)-1})",
                                        },
                                        timeout_sec=10,
                                    )
                            else:
                                ws_log(f"[history]-N ignored: group not recording, group_id={group_id}")
                        continue
                    await handle_ws_recording_event(msg)
        except Exception as e:
            ws_log(f"处理消息时发生错误: {e}")
        finally:
            WS_MESSAGE_QUEUE.task_done()


def run_ws_event_loop():
    """Run the WebSocket event loop in a background thread.
    Auto-restarts on failure, mirroring fwlog's main_loop resilience."""
    global WS_MESSAGE_QUEUE, WS_CLIENT
    while True:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        WS_MESSAGE_QUEUE = asyncio.Queue()
        WS_CLIENT = LogutilBotClient(NAPCAT_WS_URL, NAPCAT_WS_TOKEN)
        try:
            loop.run_until_complete(
                asyncio.gather(
                    WS_CLIENT.run(),
                    process_ws_messages(),
                )
            )
        except Exception as e:
            ws_log(f"WebSocket 事件循环异常，3秒后重试: {e}")
        finally:
            try:
                loop.close()
            except Exception:
                pass
            WS_CLIENT = None
        ws_log("WebSocket 事件循环将在 3 秒后重新启动...")
        time.sleep(3)


def ensure_ws_worker_started():
    """Start the WebSocket listener thread if enabled and not already running."""
    global WS_WORKER
    if not LOGUTIL_WS_ENABLED:
        return

    with STATE_LOCK:
        if WS_WORKER and WS_WORKER.is_alive():
            return
        WS_WORKER = threading.Thread(
            target=run_ws_event_loop,
            name="logutil-ws-worker",
            daemon=True,
        )
        WS_WORKER.start()


# ====== 继续原来启动入口 ======
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='LogAI 4.4.0 - TRPG Log Analysis Server')
    parser.add_argument('--api-key', type=str, default=None, help='OpenAI / DeepSeek API Key')
    parser.add_argument('--api-base-url', type=str, default=None, help=f'OpenAI-compatible API base URL (default: {AI_BASE_URL})')
    parser.add_argument('--ai-model', type=str, default=None, help=f'Default AI model (default: {AI_MODEL})')
    parser.add_argument('--ai-model-pro', type=str, default=None, help=f'Pro AI model (default: {AI_MODEL_PRO})')
    parser.add_argument('--image-api-key', type=str, default=None, help='NovelAI Image API Key')
    parser.add_argument('--host', type=str, default=None, help=f'Server host (default: {LOGAI_HOST})')
    parser.add_argument('--port', type=int, default=None, help=f'Server port (default: {LOGAI_PORT})')
    parser.add_argument('--napcat-url', type=str, default=None, help=f'NapCat HTTP API base URL (default: {NAPCAT_API_BASE})')
    parser.add_argument('--napcat-token', type=str, default=None, help='NapCat HTTP API access token')
    parser.add_argument('--ws-url', type=str, default=None, help=f'NapCat WebSocket URL (default: {NAPCAT_WS_URL})')
    parser.add_argument('--ws-token', type=str, default=None, help='NapCat WebSocket access token')
    parser.add_argument('--bridge-token', type=str, default=None, help='Bridge HTTP API access token')
    parser.add_argument('--bridge-public-base', type=str, default=None, help=f'Bridge public base URL (default: {BRIDGE_PUBLIC_BASE})')
    parser.add_argument('--ws-enabled', type=str, default=None, choices=['0', '1', 'true', 'false'], help='Enable WS listener (default: 1)')
    parser.add_argument('--bridge-mode', type=int, default=None, choices=[0, 1], help='File bridge mode: 0=WS 1=poll (default: 0)')
    args = parser.parse_args()

    # Apply CLI overrides (CLI > env > default)
    if args.api_key is not None:
        AI_API_KEY = args.api_key
    if args.api_base_url is not None:
        AI_BASE_URL = args.api_base_url
    if args.ai_model is not None:
        AI_MODEL = args.ai_model
    if args.ai_model_pro is not None:
        AI_MODEL_PRO = args.ai_model_pro
    if args.image_api_key is not None:
        IMAGE_API_KEY = args.image_api_key
    if args.host is not None:
        LOGAI_HOST = args.host
    if args.port is not None:
        LOGAI_PORT = args.port
    if args.napcat_url is not None:
        NAPCAT_API_BASE = args.napcat_url.rstrip('/')
    if args.napcat_token is not None:
        NAPCAT_TOKEN = args.napcat_token
    if args.ws_url is not None:
        NAPCAT_WS_URL = args.ws_url
    if args.ws_token is not None:
        NAPCAT_WS_TOKEN = args.ws_token
    if args.bridge_token is not None:
        BRIDGE_TOKEN = args.bridge_token
    if args.bridge_public_base is not None:
        BRIDGE_PUBLIC_BASE = args.bridge_public_base.rstrip('/')
    if args.ws_enabled is not None:
        LOGUTIL_WS_ENABLED = args.ws_enabled.lower() in ('1', 'true')
    if args.bridge_mode is not None:
        NC_FILE_BRIDGE_MODE = args.bridge_mode
    # Re-derive configs that depend on LOGAI_PORT
    if args.port is not None and args.bridge_public_base is None:
        BRIDGE_PUBLIC_BASE = f"http://{get_lan_ip()}:{LOGAI_PORT}"
    # Re-derive NAPCAT_WS_TOKEN if napcat-token changed but ws-token not set
    if args.napcat_token is not None and args.ws_token is None:
        NAPCAT_WS_TOKEN = NAPCAT_TOKEN or ""

    disable_quick_edit()
    if not os.path.exists("./fonts"): os.makedirs("./fonts")
    ensure_bridge_cache_dir()
    # v4.4.0: 恢复历史桥接记录
    load_history()
    # v4.4.0: 注册退出回调 — 将所有缓存移入历史
    atexit.register(shutdown_handler)
    ensure_worker_started()
    ensure_poll_worker_started()
    if LOGUTIL_WS_ENABLED:
        ensure_ws_worker_started()
    print(f"Async Log Server Started (Port: {LOGAI_PORT})")
    print(f"NapCat bridge enabled: {BRIDGE_PUBLIC_BASE}/napcat/event")
    print(f"NapCat latest endpoint: {BRIDGE_PUBLIC_BASE}/bridge/latest")
    print(f"Bridge poll worker: interval={BRIDGE_POLL_INTERVAL_SEC}s (lazy group registration)")
    if LOGUTIL_WS_ENABLED:
        print(f"Logutil WS listener: enabled (url={NAPCAT_WS_URL})")
    else:
        print(f"Logutil WS listener: disabled (set LOGUTIL_WS_ENABLED=1 to enable)")
    app.run(host=LOGAI_HOST, port=LOGAI_PORT, debug=False, threaded=True)
